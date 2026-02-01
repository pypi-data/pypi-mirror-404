"""Multi-format report export for Oscura.

This module provides unified export interface for generating reports in
multiple formats (HTML, PDF, DOCX, Markdown) from a single source.


Example:
    >>> from oscura.reporting.export import export_report
    >>> export_report(report, "output", formats=["pdf", "html", "markdown"])
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any, Literal

if TYPE_CHECKING:
    from oscura.reporting.core import Report


def export_report(
    report: Report,
    output_path: str | Path,
    *,
    formats: list[Literal["pdf", "html", "docx", "markdown"]] | None = None,
    format_options: dict[str, Any] | None = None,
) -> dict[str, Path]:
    """Export report to multiple formats.

    Args:
        report: Report object to export.
        output_path: Base output path (without extension).
        formats: List of formats to generate (default: ["pdf", "html"]).
        format_options: Format-specific options dictionary.

    Returns:
        Dictionary mapping format to output file path.

    Raises:
        ValueError: If unsupported format is specified.

    Example:
        >>> paths = export_report(
        ...     report,
        ...     "report",
        ...     formats=["pdf", "html", "markdown"],
        ...     format_options={
        ...         "pdf": {"dpi": 300, "pdfa_compliance": True},
        ...         "html": {"interactive": True, "dark_mode": True},
        ...     }
        ... )

    References:
        REPORT-010
    """
    if formats is None:
        formats = ["pdf", "html"]

    if format_options is None:
        format_options = {}

    output_path = Path(output_path)
    generated_files = {}

    for fmt in formats:
        fmt_opts = format_options.get(fmt, {})

        if fmt == "pdf":
            file_path = _export_pdf(report, output_path, **fmt_opts)
        elif fmt == "html":
            file_path = _export_html(report, output_path, **fmt_opts)
        elif fmt == "docx":
            file_path = _export_docx(report, output_path, **fmt_opts)
        elif fmt == "markdown":
            file_path = _export_markdown(report, output_path, **fmt_opts)
        else:
            raise ValueError(f"Unsupported format: {fmt}")

        generated_files[fmt] = file_path

    return generated_files  # type: ignore[return-value]


def _export_pdf(
    report: Report,
    output_path: Path,
    **options: Any,
) -> Path:
    """Export report as PDF."""
    from oscura.reporting.pdf import save_pdf_report

    path = output_path.with_suffix(".pdf")
    save_pdf_report(report, path, **options)
    return path


def _export_html(
    report: Report,
    output_path: Path,
    **options: Any,
) -> Path:
    """Export report as HTML."""
    from oscura.reporting.html import save_html_report

    path = output_path.with_suffix(".html")
    save_html_report(report, path, **options)
    return path


def _export_markdown(
    report: Report,
    output_path: Path,
    **options: Any,
) -> Path:
    """Export report as Markdown."""
    path = output_path.with_suffix(".md")
    markdown_content = report.to_markdown()
    path.write_text(markdown_content, encoding="utf-8")
    return path


def _export_docx(
    report: Report,
    output_path: Path,
    **options: Any,
) -> Path:
    """Export report as DOCX.

    Requires python-docx library.

    Args:
        report: Report object to export.
        output_path: Base output path (extension will be changed to .docx).
        **options: Format-specific options (currently unused).

    Returns:
        Path to the created DOCX file.

    Raises:
        ImportError: If python-docx library is not installed.

    References:
        REPORT-019
    """
    doc = _create_docx_document()
    path = output_path.with_suffix(".docx")

    _add_docx_header(doc, report)
    _add_docx_sections(doc, report)

    doc.save(str(path))
    return path


def _create_docx_document() -> Any:
    """Create and configure DOCX document.

    Returns:
        Document object from python-docx.

    Raises:
        ImportError: If python-docx not installed.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. Install with: pip install python-docx"
        )
    return Document()


def _add_docx_header(doc: Any, report: Report) -> None:
    """Add title and metadata to DOCX document.

    Args:
        doc: Document object.
        report: Report to extract metadata from.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_heading(report.config.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if report.config.author:
        doc.add_paragraph(f"Author: {report.config.author}")
    doc.add_paragraph(f"Date: {report.config.created.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()


def _add_docx_sections(doc: Any, report: Report) -> None:
    """Add all sections to DOCX document.

    Args:
        doc: Document object.
        report: Report with sections to add.
    """
    for section in report.sections:
        if not section.visible:
            continue

        doc.add_heading(section.title, level=section.level)
        _add_docx_section_content(doc, section)
        _add_docx_subsections(doc, section)


def _add_docx_section_content(doc: Any, section: Any) -> None:
    """Add section content to DOCX.

    Args:
        doc: Document object.
        section: Section with content to add.
    """
    if isinstance(section.content, str):
        doc.add_paragraph(section.content)
    elif isinstance(section.content, list):
        for item in section.content:
            _add_docx_content_item(doc, item)


def _add_docx_content_item(doc: Any, item: Any) -> None:
    """Add single content item to DOCX.

    Args:
        doc: Document object.
        item: Content item (dict or other).
    """
    if isinstance(item, dict):
        if item.get("type") == "table":
            _add_table_to_docx(doc, item)
        elif item.get("type") == "figure":
            doc.add_paragraph(f"[Figure: {item.get('caption', 'N/A')}]")
    else:
        doc.add_paragraph(str(item))


def _add_docx_subsections(doc: Any, section: Any) -> None:
    """Add subsections to DOCX document.

    Args:
        doc: Document object.
        section: Section with subsections.
    """
    for subsec in section.subsections:
        if not subsec.visible:
            continue
        doc.add_heading(subsec.title, level=subsec.level)
        if isinstance(subsec.content, str):
            doc.add_paragraph(subsec.content)


def _add_table_to_docx(doc: Any, table_dict: dict[str, Any]) -> None:
    """Add table to DOCX document."""
    headers = table_dict.get("headers", [])
    data = table_dict.get("data", [])

    if not headers and not data:
        return

    # Create table
    num_cols = len(headers) if headers else len(data[0]) if data else 0
    num_rows = len(data) + (1 if headers else 0)

    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Light Grid Accent 1"

    # Add headers
    if headers:
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # Add data
    start_row = 1 if headers else 0
    for i, row in enumerate(data):
        row_cells = table.rows[start_row + i].cells
        for j, cell in enumerate(row):
            row_cells[j].text = str(cell)

    # Add caption
    if table_dict.get("caption"):
        doc.add_paragraph(table_dict["caption"], style="Caption")


def export_multiple_reports(
    reports: dict[str, Report],
    output_dir: str | Path,
    *,
    format: Literal["pdf", "html", "docx", "markdown"] = "pdf",
    **options: Any,
) -> dict[str, Path]:
    """Export multiple reports to a directory.

    Args:
        reports: Dictionary mapping name to Report object.
        output_dir: Output directory path.
        format: Export format for all reports.
        **options: Format-specific options.

    Returns:
        Dictionary mapping report name to output path.

    References:
        REPORT-010
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    generated_files = {}

    for name, report in reports.items():
        output_path = output_dir / name
        files = export_report(
            report, output_path, formats=[format], format_options={format: options}
        )
        generated_files[name] = files[format]

    return generated_files


def batch_export_formats(
    report: Report,
    output_dir: str | Path,
    *,
    formats: list[str] | None = None,
    **options: Any,
) -> dict[str, Path]:
    """Export single report to multiple formats in a directory.

    Args:
        report: Report to export.
        output_dir: Output directory.
        formats: List of formats (default: all supported).
        **options: Common options for all formats.

    Returns:
        Dictionary mapping format to output path.

    References:
        REPORT-010
    """
    if formats is None:
        formats = ["pdf", "html", "docx", "markdown"]

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name = report.config.title.lower().replace(" ", "_")
    output_path = output_dir / base_name

    return export_report(report, output_path, formats=formats, format_options=options)  # type: ignore[arg-type]


def create_archive(
    files: dict[str, Path],
    archive_path: str | Path,
    *,
    format: Literal["zip", "tar", "tar.gz"] = "zip",
) -> Path:
    """Create archive of exported report files.

    Args:
        files: Dictionary of files to archive.
        archive_path: Output archive path.
        format: Archive format (zip, tar, tar.gz).

    Returns:
        Path to created archive.

    Raises:
        ValueError: If unsupported archive format is specified.

    References:
        REPORT-010
    """
    from pathlib import Path

    archive_path = Path(archive_path)

    if format == "zip":
        import zipfile

        with zipfile.ZipFile(archive_path.with_suffix(".zip"), "w") as zipf:
            for path in files.values():
                zipf.write(path, arcname=path.name)

        return archive_path.with_suffix(".zip")

    elif format in ("tar", "tar.gz"):
        import tarfile

        mode = "w:gz" if format == "tar.gz" else "w"
        suffix = ".tar.gz" if format == "tar.gz" else ".tar"

        with tarfile.open(archive_path.with_suffix(suffix), mode) as tar:  # type: ignore[call-overload]
            for path in files.values():
                tar.add(path, arcname=path.name)

        return archive_path.with_suffix(suffix)

    else:
        raise ValueError(f"Unsupported archive format: {format}")
