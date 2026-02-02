"""Word document manipulation toolset using python-docx."""

from __future__ import annotations

import json
from typing import Literal, Optional, Any
from pydantic import BaseModel, Field

from openreward.environments import tool, ToolOutput, TextBlock
from openreward.environments.toolset import Toolset


# ===== Pydantic Parameter Models =====

class CreateDocumentParams(BaseModel):
    file_path: str = Field(..., description="Path where the new document will be created")
    title: str | None = Field(None, description="Document title metadata")
    author: str | None = Field(None, description="Document author metadata")


class GetOverviewParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")


class ReadDocumentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    include_tables: bool = Field(True, description="Include table content in output")
    max_paragraphs: int | None = Field(None, description="Limit number of paragraphs returned")


class ReadImageParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    image_index: int | None = Field(None, description="Specific image index to read (if None, returns all images)")
    include_data: bool = Field(True, description="Include base64-encoded image data in response")


class AddContentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    text: str = Field(..., description="Text content to add")
    content_type: Literal["paragraph", "heading"] = Field("paragraph", description="Type of content")
    heading_level: int | None = Field(None, description="Heading level (1-9) if content_type is heading")
    style: str | None = Field(None, description="Paragraph style name")
    position: Literal["end", "start"] = Field("end", description="Where to insert content")


class AddImageParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    image_path: str = Field(..., description="Path to the image file to insert")
    width_inches: float | None = Field(None, description="Image width in inches")
    height_inches: float | None = Field(None, description="Image height in inches")
    position: Literal["end", "start"] = Field("end", description="Where to insert image")


class EditContentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    paragraph_index: int = Field(..., description="Zero-based paragraph index")
    new_text: str = Field(..., description="New text content")
    append: bool = Field(False, description="Append to existing text instead of replacing")


class DeleteContentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    paragraph_index: int = Field(..., description="Zero-based paragraph index to delete")


class ModifyImageParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    image_index: int = Field(..., description="Zero-based image index")
    width_inches: float | None = Field(None, description="New width in inches")
    height_inches: float | None = Field(None, description="New height in inches")


class ApplyFormattingParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    paragraph_index: int = Field(..., description="Zero-based paragraph index")
    bold: bool | None = Field(None, description="Apply bold formatting")
    italic: bool | None = Field(None, description="Apply italic formatting")
    underline: bool | None = Field(None, description="Apply underline formatting")
    font_name: str | None = Field(None, description="Font family name (e.g., 'Arial')")
    font_size: int | None = Field(None, description="Font size in points")
    color: str | None = Field(None, description="Hex color code (e.g., 'FF0000' for red)")


class DeleteDocumentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document to delete")


class SearchDocumentParams(BaseModel):
    file_path: str = Field(..., description="Path to the Word document")
    search_text: str = Field(..., description="Text to search for")
    case_sensitive: bool = Field(False, description="Case-sensitive search")
    search_in_tables: bool = Field(True, description="Include table content in search")
    max_results: int | None = Field(None, description="Limit number of results")


# ===== Word Toolset Class =====

class WordToolset(Toolset):
    """Toolset providing Word document manipulation tools via python-docx executed in sandbox"""

    @tool
    async def word_create_document(self, params: CreateDocumentParams) -> ToolOutput:
        """Create a new Word document with optional metadata"""
        title_escaped = params.title.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n') if params.title else ""
        author_escaped = params.author.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n') if params.author else ""

        script = f'''
import json
from docx import Document

try:
    doc = Document()

    # Set metadata if provided
    if "{title_escaped}":
        doc.core_properties.title = "{title_escaped}"
    if "{author_escaped}":
        doc.core_properties.author = "{author_escaped}"

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "title": "{title_escaped}" if "{title_escaped}" else None,
        "author": "{author_escaped}" if "{author_escaped}" else None,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Document created successfully at {params.file_path}")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_delete_document(self, params: DeleteDocumentParams) -> ToolOutput:
        """Delete a Word document file from sandbox"""
        output, exit_code = await self.sandbox.run(f"rm {params.file_path}")

        if exit_code == 0:
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Document deleted successfully: {params.file_path}")],
                metadata={"file_path": params.file_path, "deleted": True},
                reward=0.0,
                finished=False,
            )
        else:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Failed to delete: {output}")],
                metadata={"file_path": params.file_path, "deleted": False, "error": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_get_document_overview(self, params: GetOverviewParams) -> ToolOutput:
        """Retrieve document structure and metadata including paragraph count, table count, and image count"""
        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    # Count paragraphs
    paragraph_count = len(doc.paragraphs)

    # Count tables
    table_count = len(doc.tables)

    # Count images (inline shapes)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    # Get metadata
    props = doc.core_properties
    metadata = {{
        "title": props.title if props.title else "",
        "author": props.author if props.author else "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
    }}

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "image_count": image_count,
        "metadata": metadata,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # Build summary
            summary = f"Document: {params.file_path}\n"
            summary += f"Paragraphs: {result['paragraph_count']}\n"
            summary += f"Tables: {result['table_count']}\n"
            summary += f"Images: {result['image_count']}\n\n"
            summary += "Metadata:\n"
            for key, value in result['metadata'].items():
                if value:
                    summary += f"  {key.capitalize()}: {value}\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_read_document_content(self, params: ReadDocumentParams) -> ToolOutput:
        """Extract all text content with paragraph structure from the document"""
        max_para = params.max_paragraphs if params.max_paragraphs else 999999

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        if idx >= {max_para}:
            break

        paragraphs.append({{
            "index": idx,
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
        }})

    # Optionally include table content
    tables = []
    if {params.include_tables}:
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({{
                "index": table_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,
            }})

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "table_count": len(tables),
        "tables": tables,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # Build summary
            summary = f"Document content from {params.file_path}\n"
            summary += f"Paragraphs: {result['paragraph_count']}\n"
            if result['table_count'] > 0:
                summary += f"Tables: {result['table_count']}\n"
            summary += "\n"

            # Show first few paragraphs
            for para in result['paragraphs'][:5]:
                text_preview = para['text'][:100] if para['text'] else "(empty)"
                summary += f"[{para['index']}] {para['style']}: {text_preview}\n"

            if result['paragraph_count'] > 5:
                summary += f"... and {result['paragraph_count'] - 5} more paragraphs\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_read_image(self, params: ReadImageParams) -> ToolOutput:
        """Extract image data and metadata from document"""
        script = f'''
import json
import base64
from docx import Document

try:
    doc = Document("{params.file_path}")

    images = []

    # Iterate through all parts to find images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob  # Binary data

            # Extract format from content_type (e.g., 'image/png' -> 'png')
            image_format = image_part.content_type.split('/')[-1]

            image_info = {{
                "index": len(images),
                "format": image_format,
                "size_bytes": len(image_data),
            }}

            # Include base64 data if requested
            if {params.include_data}:
                image_base64 = base64.b64encode(image_data).decode('utf-8')
                image_info["data"] = image_base64

            images.append(image_info)

    # Filter by image_index if specified
    if {params.image_index} is not None:
        if {params.image_index} < 0 or {params.image_index} >= len(images):
            print(json.dumps({{
                "success": False,
                "error": f"Image index {params.image_index} out of range (document has {{len(images)}} images)"
            }}))
        else:
            images = [images[{params.image_index}]]

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "image_count": len(images),
        "images": images,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # Build summary
            summary = f"Found {result['image_count']} image(s) in {params.file_path}\n\n"
            for img in result['images']:
                summary += f"Image {img['index']}: {img['format'].upper()} ({img['size_bytes']} bytes)\n"
                if params.include_data and 'data' in img:
                    summary += f"  Base64 data: {len(img['data'])} characters\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_add_content_text(self, params: AddContentParams) -> ToolOutput:
        """Add paragraphs or headings to document"""
        text_escaped = params.text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
        style_escaped = params.style.replace('\\', '\\\\').replace('"', '\\"') if params.style else ""

        # Build the content addition logic
        if params.content_type == "heading":
            if not params.heading_level or params.heading_level < 1 or params.heading_level > 9:
                return ToolOutput(
                    blocks=[TextBlock(text="❌ Error: heading_level must be between 1 and 9 for heading content")],
                    metadata={"error": "Invalid heading_level"},
                    reward=0.0,
                    finished=False,
                )
            add_code = f'new_para = doc.add_heading("{text_escaped}", level={params.heading_level})'
        else:
            add_code = f'new_para = doc.add_paragraph("{text_escaped}")'
            if params.style:
                add_code += f'\n    new_para.style = "{style_escaped}"'

        # Handle positioning (start vs end)
        if params.position == "start":
            insert_code = f'''
    # Insert at start by inserting before first paragraph
    if len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]._element
        p.getparent().insert(0, new_para._element)
    else:
        {add_code}
'''
        else:
            insert_code = f'    {add_code}'

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

{insert_code}

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "content_type": "{params.content_type}",
        "position": "{params.position}",
        "paragraph_count": len(doc.paragraphs),
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            content_label = f"Heading {params.heading_level}" if params.content_type == "heading" else "Paragraph"
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ {content_label} added at {params.position} of document")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_add_image(self, params: AddImageParams) -> ToolOutput:
        """Insert images into document"""
        width_param = f", width=Inches({params.width_inches})" if params.width_inches else ""
        height_param = f", height=Inches({params.height_inches})" if params.height_inches else ""

        # Handle positioning
        if params.position == "start":
            insert_code = f'''
    # Add image in a new paragraph at the start
    paragraph = doc.paragraphs[0] if len(doc.paragraphs) > 0 else doc.add_paragraph()
    run = paragraph.insert_paragraph_before().add_run()
    picture = run.add_picture("{params.image_path}"{width_param}{height_param})
'''
        else:
            insert_code = f'''
    # Add image at end
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    picture = run.add_picture("{params.image_path}"{width_param}{height_param})
'''

        script = f'''
import json
from docx import Document
from docx.shared import Inches

try:
    doc = Document("{params.file_path}")

{insert_code}

    # Get final dimensions
    width_inches = picture.width / 914400
    height_inches = picture.height / 914400

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "image_path": "{params.image_path}",
        "position": "{params.position}",
        "width": width_inches,
        "height": height_inches,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Image added at {params.position} of document ({result['width']:.2f}\" x {result['height']:.2f}\")")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_edit_content_text(self, params: EditContentParams) -> ToolOutput:
        """Modify existing paragraph text"""
        text_escaped = params.new_text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')

        if params.append:
            update_code = f'para.add_run("\\n{text_escaped}")'
        else:
            update_code = f'para.text = "{text_escaped}"'

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]
        old_text = para.text

        {update_code}

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "old_text": old_text[:100],
            "new_text": para.text[:100],
            "append": {params.append},
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            action = "appended to" if params.append else "updated"
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Paragraph {params.paragraph_index} {action} successfully\nPreview: {result['new_text']}...")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_delete_content_text(self, params: DeleteContentParams) -> ToolOutput:
        """Remove paragraphs by index"""
        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]
        deleted_text = para.text

        # Delete paragraph element
        p = para._element
        p.getparent().remove(p)

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "deleted_text": deleted_text[:100],
            "remaining_paragraphs": len(doc.paragraphs),
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Paragraph {params.paragraph_index} deleted successfully\nDeleted text: {result['deleted_text']}...")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_modify_image(self, params: ModifyImageParams) -> ToolOutput:
        """Update image properties (size)"""
        width_code = f"picture.width = Inches({params.width_inches})" if params.width_inches else ""
        height_code = f"picture.height = Inches({params.height_inches})" if params.height_inches else ""

        script = f'''
import json
from docx import Document
from docx.shared import Inches

try:
    doc = Document("{params.file_path}")

    # Try using inline_shapes if available
    if hasattr(doc, 'inline_shapes'):
        if {params.image_index} < 0 or {params.image_index} >= len(doc.inline_shapes):
            print(json.dumps({{
                "success": False,
                "error": f"Image index {params.image_index} out of range (document has {{len(doc.inline_shapes)}} images)"
            }}))
        else:
            picture = doc.inline_shapes[{params.image_index}]

            {width_code}
            {height_code}

            new_width = picture.width / 914400
            new_height = picture.height / 914400

            doc.save("{params.file_path}")

            result = {{
                "success": True,
                "file_path": "{params.file_path}",
                "image_index": {params.image_index},
                "new_width": new_width,
                "new_height": new_height,
            }}
            print(json.dumps(result))
    else:
        print(json.dumps({{
            "success": False,
            "error": "Image modification requires python-docx with inline_shapes support"
        }}))

except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Image {params.image_index} modified successfully ({result['new_width']:.2f}\" x {result['new_height']:.2f}\")")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_apply_formatting(self, params: ApplyFormattingParams) -> ToolOutput:
        """Apply text formatting to paragraph runs"""
        # Build formatting code
        format_lines = []
        if params.bold is not None:
            format_lines.append(f"run.font.bold = {str(params.bold)}")
        if params.italic is not None:
            format_lines.append(f"run.font.italic = {str(params.italic)}")
        if params.underline is not None:
            format_lines.append(f"run.font.underline = {str(params.underline)}")
        if params.font_name:
            font_escaped = params.font_name.replace('\\', '\\\\').replace('"', '\\"')
            format_lines.append(f'run.font.name = "{font_escaped}"')
        if params.font_size:
            format_lines.append(f"run.font.size = Pt({params.font_size})")
        if params.color:
            # Convert hex color (e.g., "FF0000") to RGB
            format_lines.append(f"run.font.color.rgb = RGBColor(int('{params.color}'[0:2], 16), int('{params.color}'[2:4], 16), int('{params.color}'[4:6], 16))")

        format_code = "\n                ".join(format_lines) if format_lines else "pass"

        script = f'''
import json
from docx import Document
from docx.shared import Pt, RGBColor

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]

        # Apply formatting to all runs in the paragraph
        for run in para.runs:
            if run.text.strip():
                {format_code}

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "formatting_applied": {{
                "bold": {params.bold},
                "italic": {params.italic},
                "underline": {params.underline},
                "font_name": "{params.font_name}" if "{params.font_name}" else None,
                "font_size": {params.font_size},
                "color": "{params.color}" if "{params.color}" else None,
            }},
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Formatting applied to paragraph {params.paragraph_index} successfully")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_search_document(self, params: SearchDocumentParams) -> ToolOutput:
        """Search for text within document paragraphs and tables"""
        search_text_escaped = params.search_text.replace('\\', '\\\\').replace('"', '\\"')
        max_results_repr = repr(params.max_results)

        script = f'''
import json
from docx import Document
import re

try:
    doc = Document("{params.file_path}")

    # Prepare search pattern
    flags = 0 if {params.case_sensitive} else re.IGNORECASE
    pattern = re.compile(re.escape("{search_text_escaped}"), flags=flags)

    results = []
    max_results = {max_results_repr}

    # Search in paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text:
            continue

        for match in pattern.finditer(text):
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            context = text[start:end]

            results.append({{
                "location_type": "paragraph",
                "paragraph_index": para_idx,
                "start_pos": match.start(),
                "end_pos": match.end(),
                "context": context,
                "style": para.style.name if para.style else "Normal"
            }})

            if max_results and len(results) >= max_results:
                break

        if max_results and len(results) >= max_results:
            break

    # Search in tables if enabled and not at max_results
    if {params.search_in_tables} and (not max_results or len(results) < max_results):
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text
                    if not text:
                        continue

                    for match in pattern.finditer(text):
                        start = max(0, match.start() - 50)
                        end = min(len(text), match.end() + 50)
                        context = text[start:end]

                        results.append({{
                            "location_type": "table",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "col_index": col_idx,
                            "start_pos": match.start(),
                            "end_pos": match.end(),
                            "context": context
                        }})

                        if max_results and len(results) >= max_results:
                            break

                    if max_results and len(results) >= max_results:
                        break

                if max_results and len(results) >= max_results:
                    break

            if max_results and len(results) >= max_results:
                break

    result = {{
        "success": True,
        "search_text": "{search_text_escaped}",
        "total_results": len(results),
        "results": results
    }}
    print(json.dumps(result))

except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        output, exit_code = await self._run_python_script(script)

        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            count = result["total_results"]
            display_text = f"✅ Found {count} occurrence(s) of '{params.search_text}'"

            if count > 0:
                display_text += "\n\nMatches:"
                for i, match in enumerate(result["results"][:5], 1):
                    if match["location_type"] == "paragraph":
                        display_text += f"\n{i}. Paragraph {match['paragraph_index']} ({match['style']}): ...{match['context']}..."
                    else:
                        display_text += f"\n{i}. Table {match['table_index']} [Row {match['row_index']}, Col {match['col_index']}]: ...{match['context']}..."

                if count > 5:
                    display_text += f"\n\n... and {count - 5} more"

            return ToolOutput(
                blocks=[TextBlock(text=display_text)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )
