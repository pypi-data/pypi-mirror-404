from docx import Document
from pathlib import Path
from typing import Any, Dict
import logging

logger = logging.getLogger(__name__)


async def execute(args: Dict[str, Any], context: Dict[str, Any]) -> str:
    file_path: str = args.get("file_path", "")

    path = Path(file_path)

    if not path.exists():
        return f"错误：文件不存在 {file_path}"

    if not path.is_file():
        return f"错误：{file_path} 不是文件"

    try:
        doc = Document(str(path))

        info: list[str] = []
        info.append(f"文件大小：{path.stat().st_size} 字节")

        core_props = doc.core_properties
        if core_props:
            info.append("\n文档属性：")
            if core_props.title:
                info.append(f"  标题: {core_props.title}")
            if core_props.author:
                info.append(f"  作者: {core_props.author}")
            if core_props.last_modified_by:
                info.append(f"  最后修改者: {core_props.last_modified_by}")
            if core_props.created:
                info.append(f"  创建时间: {core_props.created}")
            if core_props.modified:
                info.append(f"  修改时间: {core_props.modified}")
            if core_props.comments:
                info.append(f"  备注: {core_props.comments}")

        paragraphs = list(doc.paragraphs)
        tables = doc.tables

        info.append(f"\n段落数：{len(paragraphs)}")
        info.append(f"表格数：{len(tables)}")

        text_content = ""
        for para in paragraphs:
            text = para.text.strip()
            if text:
                text_content += text + "\n"

        if not text_content.strip():
            text_content = "(文档未检测到文本内容)"

        info.append(f"\n文本内容预览（前 5000 字符）：\n{text_content[:5000]}")
        if len(text_content) > 5000:
            info.append(f"\n... (共 {len(text_content)} 字符)")

        if tables:
            info.append("\n表格内容：")
            for i, table in enumerate(tables[:3], 1):
                info.append(f"\n表格 {i}:")
                for row in table.rows[:5]:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cells)
                    info.append(f"  {row_text}")
                if len(table.rows) > 5:
                    info.append(f"  ... (共 {len(table.rows)} 行)")
            if len(tables) > 3:
                info.append(f"\n... (还有 {len(tables) - 3} 个表格)")

        return "\n".join(info)

    except Exception as e:
        logger.exception(f"解析 Word 文档失败: {e}")
        return f"解析 Word 文档失败: {e}"
