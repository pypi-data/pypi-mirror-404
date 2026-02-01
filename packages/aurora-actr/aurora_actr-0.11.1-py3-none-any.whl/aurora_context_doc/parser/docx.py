"""DOCX document parser using python-docx."""

import hashlib
from datetime import datetime, timezone
from pathlib import Path

from aurora_core.chunks import DocChunk

from aurora_context_doc.parser.base import DocumentParser


class DOCXParser(DocumentParser):
    """Parse DOCX documents with heading-based hierarchy.

    Extraction strategy:
    - Detect Heading 1/2/3/4/5 paragraph styles
    - Build hierarchy from heading levels
    - Extract paragraphs and tables as separate chunks
    """

    def __init__(self):
        """Initialize DOCX parser (lazy import python-docx)."""
        self._docx = None

    def _get_docx(self):
        """Lazy import of python-docx.

        Returns:
            docx module

        Raises:
            ImportError: If python-docx is not installed

        """
        if self._docx is None:
            try:
                import docx

                self._docx = docx
            except ImportError as e:
                raise ImportError(
                    "python-docx is required for DOCX parsing. "
                    "Install with: pip install aurora-context-doc[docx]"
                ) from e
        return self._docx

    def supported_extensions(self) -> list[str]:
        """Return supported extensions.

        Returns:
            List of extensions: ["docx", "DOCX"]

        """
        return ["docx", "DOCX"]

    def parse(self, file_path: Path) -> list[DocChunk]:
        """Extract sections from DOCX with heading hierarchy.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of DocChunk objects with hierarchy

        Raises:
            FileNotFoundError: If DOCX file not found
            ValueError: If DOCX is invalid or corrupted
            RuntimeError: If parsing fails

        """
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        docx = self._get_docx()

        try:
            doc = docx.Document(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX: {file_path}") from e

        chunks = []
        timestamp = datetime.now(timezone.utc)

        # Track current section hierarchy
        current_sections = {}  # level -> (chunk_id, section_path)
        chunk_counter = 0

        for paragraph in doc.paragraphs:
            # Check if paragraph is a heading
            level = self._get_heading_level(paragraph)

            if level > 0:
                # Create section chunk for heading
                chunk_id = self._generate_chunk_id(file_path, f"section-{chunk_counter}")
                chunk_counter += 1

                # Determine parent and section path
                parent_id = None
                section_path = []

                if level > 1:
                    # Look for parent at level-1
                    for parent_level in range(level - 1, 0, -1):
                        if parent_level in current_sections:
                            parent_id, parent_path = current_sections[parent_level]
                            section_path = parent_path.copy()
                            break

                # Add current title to path
                title = paragraph.text.strip()
                section_path.append(title)

                # Update current sections
                current_sections[level] = (chunk_id, section_path)

                # Clear deeper levels
                for deeper_level in range(level + 1, 10):
                    current_sections.pop(deeper_level, None)

                # Create section chunk
                chunk = DocChunk(
                    chunk_id=chunk_id,
                    file_path=str(file_path.absolute()),
                    page_start=0,  # DOCX doesn't have explicit pages
                    page_end=0,
                    element_type="section",
                    name=title,
                    content=paragraph.text,
                    parent_chunk_id=parent_id,
                    section_path=section_path,
                    section_level=level,
                    document_type="docx",
                    created_at=timestamp,
                    updated_at=timestamp,
                )

                chunks.append(chunk)

            elif paragraph.text.strip():
                # Regular paragraph - attach to current section
                chunk_id = self._generate_chunk_id(file_path, f"para-{chunk_counter}")
                chunk_counter += 1

                # Find parent section (deepest level)
                parent_id = None
                section_path = []

                for level in range(10, 0, -1):
                    if level in current_sections:
                        parent_id, parent_path = current_sections[level]
                        section_path = parent_path.copy()
                        break

                chunk = DocChunk(
                    chunk_id=chunk_id,
                    file_path=str(file_path.absolute()),
                    page_start=0,
                    page_end=0,
                    element_type="paragraph",
                    name=paragraph.text[:50],  # First 50 chars as name
                    content=paragraph.text,
                    parent_chunk_id=parent_id,
                    section_path=section_path,
                    section_level=0,
                    document_type="docx",
                    created_at=timestamp,
                    updated_at=timestamp,
                )

                chunks.append(chunk)

        # Extract tables
        for i, table in enumerate(doc.tables):
            chunk_id = self._generate_chunk_id(file_path, f"table-{i}")

            # Convert table to text
            table_text = self._table_to_text(table)

            # Find current section for parent
            parent_id = None
            section_path = []

            for level in range(10, 0, -1):
                if level in current_sections:
                    parent_id, parent_path = current_sections[level]
                    section_path = parent_path.copy()
                    break

            chunk = DocChunk(
                chunk_id=chunk_id,
                file_path=str(file_path.absolute()),
                page_start=0,
                page_end=0,
                element_type="table",
                name=f"Table {i + 1}",
                content=table_text,
                parent_chunk_id=parent_id,
                section_path=section_path,
                section_level=0,
                document_type="docx",
                created_at=timestamp,
                updated_at=timestamp,
            )

            chunks.append(chunk)

        return chunks

    def _get_heading_level(self, paragraph) -> int:
        """Determine heading level from paragraph style.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            Heading level (1-5), or 0 if not a heading

        """
        style = paragraph.style.name.lower() if paragraph.style else ""

        if "heading 1" in style:
            return 1
        elif "heading 2" in style:
            return 2
        elif "heading 3" in style:
            return 3
        elif "heading 4" in style:
            return 4
        elif "heading 5" in style:
            return 5

        return 0

    def _table_to_text(self, table) -> str:
        """Convert table to plain text representation.

        Args:
            table: python-docx Table object

        Returns:
            Text representation of table

        """
        lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

        return "\n".join(lines)

    def _generate_chunk_id(self, file_path: Path, suffix: str) -> str:
        """Generate unique chunk ID for a document section.

        Args:
            file_path: Path to document
            suffix: Unique suffix (e.g., "section-1", "table-2")

        Returns:
            Unique chunk ID string

        """
        content = f"{file_path.absolute()}:{suffix}"
        hash_digest = hashlib.sha256(content.encode()).hexdigest()[:16]
        return f"doc:docx:{hash_digest}"


__all__ = ["DOCXParser"]
