"""
File Manager module for processing various file types for conversation context.

This module provides functionality for:
- Extracting text from documents (PDF, DOCX, TXT, MD, CSV)
- Processing images as base64-encoded data
- Token counting for file content
"""

import base64
import logging
import mimetypes
from pathlib import Path
from typing import Dict, Optional, Tuple


class FileManager:
    """Manages file processing for conversation context."""

    # Supported file extensions
    SUPPORTED_TEXT_FILES = {'.txt', '.md', '.csv', '.json', '.yaml', '.yml', '.xml', '.log'}
    SUPPORTED_CODE_FILES = {
        '.py', '.js', '.ts', '.jsx', '.tsx',  # Python, JavaScript, TypeScript
        '.java', '.kt', '.scala',              # JVM languages
        '.c', '.cpp', '.cc', '.h', '.hpp',     # C/C++
        '.cs',                                  # C#
        '.go', '.rs', '.rb',                   # Go, Rust, Ruby
        '.php', '.swift', '.m',                # PHP, Swift, Objective-C
        '.sh', '.bash', '.zsh', '.ps1',        # Shell scripts
        '.sql', '.r',                          # SQL, R
        '.html', '.css', '.scss', '.sass',     # Web
        '.vue', '.svelte',                     # Web frameworks
        '.toml', '.ini', '.cfg', '.conf',      # Config files
        '.dockerfile', '.tf', '.hcl',          # Infrastructure
    }
    SUPPORTED_DOCUMENT_FILES = {'.pdf', '.docx', '.doc'}
    SUPPORTED_IMAGE_FILES = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}

    def __init__(self, bedrock_service=None):
        """
        Initialise the file manager.

        Args:
            bedrock_service: Optional BedrockService instance for token counting
        """
        self.bedrock_service = bedrock_service

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """
        Check if a file type is supported.

        Args:
            file_path: Path to the file

        Returns:
            True if the file type is supported
        """
        ext = Path(file_path).suffix.lower()
        return ext in (cls.SUPPORTED_TEXT_FILES | cls.SUPPORTED_CODE_FILES | cls.SUPPORTED_DOCUMENT_FILES | cls.SUPPORTED_IMAGE_FILES)

    def process_file(self, file_path: str) -> Dict:
        """
        Process a file and extract its content.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary containing:
            - filename: Original filename
            - file_type: File extension
            - file_size: Size in bytes
            - content_text: Extracted text content (for documents)
            - content_base64: Base64 encoded content (for images)
            - mime_type: MIME type (for images)
            - token_count: Token count of extracted content
            - error: Error message if processing failed
        """
        path = Path(file_path)

        if not path.exists():
            return {'error': f"File not found: {file_path}"}

        if not self.is_supported(file_path):
            return {'error': f"Unsupported file type: {path.suffix}"}

        result = {
            'filename': path.name,
            'file_type': path.suffix.lower(),
            'file_size': path.stat().st_size,
            'content_text': None,
            'content_base64': None,
            'mime_type': None,
            'token_count': 0
        }

        ext = path.suffix.lower()

        try:
            # Process text and code files
            if ext in (self.SUPPORTED_TEXT_FILES | self.SUPPORTED_CODE_FILES):
                result['content_text'] = self._extract_text_file(file_path)
                if result['content_text']:
                    result['token_count'] = self._count_tokens(result['content_text'])

            # Process documents
            elif ext in self.SUPPORTED_DOCUMENT_FILES:
                if ext == '.pdf':
                    result['content_text'] = self._extract_pdf(file_path)
                elif ext in {'.docx', '.doc'}:
                    result['content_text'] = self._extract_docx(file_path)

                if result['content_text']:
                    result['token_count'] = self._count_tokens(result['content_text'])

            # Process images
            elif ext in self.SUPPORTED_IMAGE_FILES:
                result['content_base64'], result['mime_type'] = self._encode_image(file_path)
                # Images don't have text token count, but we could estimate based on size
                # For now, we'll use a simple heuristic: ~1 token per 750 bytes
                result['token_count'] = max(1, result['file_size'] // 750)

            logging.info(f"Processed file: {path.name} ({result['file_type']}, {result['token_count']} tokens)")

        except Exception as e:
            logging.error(f"Error processing file {file_path}: {e}")
            result['error'] = str(e)

        return result

    def _extract_text_file(self, file_path: str) -> str:
        """
        Extract text from a plain text file.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fall back to latin-1 if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logging.error(f"Failed to read text file {file_path}: {e}")
                raise

    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content
        """
        try:
            import pypdf

            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            return '\n\n'.join(text_parts)

        except ImportError:
            raise ImportError("pypdf library is required for PDF processing. Install with: pip install pypdf")
        except Exception as e:
            logging.error(f"Failed to extract PDF {file_path}: {e}")
            raise

    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content
        """
        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return '\n\n'.join(text_parts)

        except ImportError:
            raise ImportError("python-docx library is required for DOCX processing. Install with: pip install python-docx")
        except Exception as e:
            logging.error(f"Failed to extract DOCX {file_path}: {e}")
            raise

    def _encode_image(self, file_path: str) -> Tuple[str, str]:
        """
        Encode an image as base64.

        Args:
            file_path: Path to the image file

        Returns:
            Tuple of (base64_string, mime_type)
        """
        try:
            # Read and encode image
            with open(file_path, 'rb') as f:
                image_data = f.read()

            base64_string = base64.b64encode(image_data).decode('utf-8')

            # Determine MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                # Default to common image types
                ext = Path(file_path).suffix.lower()
                mime_map = {
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.png': 'image/png',
                    '.gif': 'image/gif',
                    '.bmp': 'image/bmp',
                    '.webp': 'image/webp'
                }
                mime_type = mime_map.get(ext, 'application/octet-stream')

            return base64_string, mime_type

        except Exception as e:
            logging.error(f"Failed to encode image {file_path}: {e}")
            raise

    def _count_tokens(self, text: str) -> int:
        """
        Count tokens in text.

        Args:
            text: Text to count tokens for

        Returns:
            Token count
        """
        if self.bedrock_service:
            return self.bedrock_service.count_tokens(text)
        else:
            # Rough estimate: ~4 characters per token
            return max(1, len(text) // 4)

    @classmethod
    def get_supported_extensions(cls) -> str:
        """
        Get a formatted string of supported file extensions.

        Returns:
            Comma-separated list of supported extensions
        """
        all_extensions = sorted(
            cls.SUPPORTED_TEXT_FILES | cls.SUPPORTED_DOCUMENT_FILES | cls.SUPPORTED_IMAGE_FILES
        )
        return ', '.join(all_extensions)

    @classmethod
    def scan_directory(cls, directory_path: str, recursive: bool = False) -> list:
        """
        Scan a directory for supported files.

        Args:
            directory_path: Path to the directory to scan
            recursive: If True, scan subdirectories as well

        Returns:
            List of absolute file paths for supported files
        """
        from pathlib import Path

        dir_path = Path(directory_path)

        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory_path}")

        iterator = dir_path.rglob('*') if recursive else dir_path.iterdir()
        supported_files = [
            str(file_path.absolute())
            for file_path in iterator
            if file_path.is_file() and cls.is_supported(str(file_path))
        ]

        return sorted(supported_files)
