import dspy
import os
import logging
try:
    import litellm
except ImportError:
    litellm = None
from wenbi.utils import segment


def configure_lm(model_string, verbose=False, **kwargs):
    """Configure the Language Model with verbose logging support"""
    logger = logging.getLogger(__name__)

    if not model_string:
        model_string = "ollama/qwen3"

    if verbose:
        logger.debug(f"Configuring LLM: {model_string}")

    parts = model_string.strip().split("/")
    provider = parts[0].lower() if parts else ""

    if verbose:
        logger.debug(f"LLM provider: {provider}")

    config = kwargs
    if provider == "ollama":
        config.update(
            {
                "base_url": "http://localhost:11434",
                "model": model_string,
            }
        )
        if verbose:
            logger.debug(f"Ollama configuration: {config}")
        lm = dspy.LM(**config)
    elif provider == "openai":
        config.update(
            {
                "api_base": "https://api.openai.com/v1",
                "model": model_string,
            }
        )
        if verbose:
            logger.debug(f"OpenAI configuration: {config}")
        lm = dspy.LM(**config)
    elif provider == "gemini":
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv(
            "GOOGLE_API_KEY_JSON")
        if not api_key:
            raise ValueError(
                "GOOGLE_API_KEY or GOOGLE_API_KEY_JSON environment variable not set."
            )

        # Extract the actual model name (e.g., "gemini-2.5-flash" from "gemini/gemini-2.5-flash")
        model_name = (
            model_string.split(
                "/", 1)[1] if "/" in model_string else model_string
        )

        # Use the correct format for LiteLLM Gemini integration
        config.update(
            {
                "model": f"gemini/{model_name}",
                "api_key": api_key,
            }
        )
        if verbose:
            logger.debug(f"Gemini configuration: model={config['model']}")
        lm = dspy.LM(**config)
    else:
        config.update({"model": model_string})
        if verbose:
            logger.debug(f"Generic LM configuration: {config}")
        lm = dspy.LM(**config)

    dspy.configure(lm=lm)

    if verbose:
        logger.debug("LLM configuration completed successfully")

    return lm


def translate(
    input_file,
    output_dir="",
    translate_language="Chinese",
    llm="ollama/qwen3",
    chunk_length=20,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    cite_timestamps=False,
    verbose=False,
):
    """
    Translate text content using LLM with verbose logging support.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Translation Process ===")
        logger.debug(f"Input file: {input_file}")
        logger.debug(f"Target language: {translate_language}")
        logger.debug(f"LLM model: {llm}")
        logger.debug(f"Chunk length: {chunk_length}")
        logger.debug(f"Max tokens: {max_tokens}")
        logger.debug(f"Include timestamps: {cite_timestamps}")

    # Configure LLM
    lm = configure_lm(
        llm,
        verbose=verbose,
        max_tokens=max_tokens,
        timeout=timeout,
        temperature=temperature,
    )

    class TranslateSignature(dspy.Signature):
        """Translate the given text to the target language while preserving the original meaning and style."""

        text_to_translate = dspy.InputField(
            desc="Text content to be translated")
        target_language = dspy.InputField(
            desc="Target language for translation")
        translated_text = dspy.OutputField(
            desc="Translated text in the target language"
        )

    translate_module = dspy.Predict(TranslateSignature)

    if verbose:
        logger.debug("LLM translation module initialized")

    # Read and segment the input text
    segmented_text = segment(input_file, chunk_length,
                             cite_timestamps, verbose=verbose)

    # Split into chunks for processing
    chunks = segmented_text.split("\n\n")
    if verbose:
        logger.debug(f"Text divided into {len(chunks)} chunks for processing")

    translated_chunks = []

    for i, chunk in enumerate(chunks, 1):
        if not chunk.strip():
            translated_chunks.append(chunk)
            continue

        if verbose:
            logger.debug(
                f"Translating chunk {
                    i}/{len(chunks)} ({len(chunk)} characters)"
            )

        try:
            # Skip timestamp headers when cite_timestamps is True
            is_timestamp_header = chunk.strip().startswith(
                "### **"
            ) and chunk.strip().endswith("**")

            if cite_timestamps and is_timestamp_header:
                if verbose:
                    logger.debug(
                        f"Preserving timestamp header: {chunk.strip()[:50]}..."
                    )
                translated_chunks.append(chunk)
            else:
                result = translate_module(
                    text_to_translate=chunk, target_language=translate_language
                )
                translated_chunks.append(result.translated_text)

                if verbose:
                    logger.debug(f"Chunk {i} translated successfully")

        except Exception as e:
            error_msg = f"Error translating chunk {i}: {e}"
            if verbose:
                logger.debug(error_msg)
            print(error_msg)
            translated_chunks.append(f"[Translation Error: {chunk}]")

    final_translation = "\n\n".join(translated_chunks)

    if verbose:
        logger.debug(
            f"Translation completed. Final text length: {
                len(final_translation)
            } characters"
        )
        logger.debug("=== Translation Process Completed ===")

    return final_translation


def rewrite(
    input_file,
    output_dir="",
    rewrite_language="Chinese",
    llm="ollama/qwen3",
    chunk_length=20,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    cite_timestamps=False,
    verbose=False,
):
    """
    Rewrite oral language to written form using LLM with verbose logging support.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Rewrite Process ===")
        logger.debug(f"Input file: {input_file}")
        logger.debug(f"Target language: {rewrite_language}")
        logger.debug(f"LLM model: {llm}")
        logger.debug(f"Chunk length: {chunk_length}")
        logger.debug(f"Max tokens: {max_tokens}")
        logger.debug(f"Include timestamps: {cite_timestamps}")

    # Configure LLM
    lm = configure_lm(
        llm,
        verbose=verbose,
        max_tokens=max_tokens,
        timeout=timeout,
        temperature=temperature,
    )

    class RewriteSignature(dspy.Signature):
        """
        Rewrite this text in {rewrite_lang} from oral to written.  Follow these rules strictly:
        1. Correct any basic grammar, punctuation, or usage errors.
        2. Improve clarity while preserving the original meaning and scholarly tone (trying your best not to change the structure of sentence)
        3. IMPORTANT: Maintaining the original meaning and length (97% of original)
        """

        oral_text = dspy.InputField(desc="Oral or spoken text to be rewritten")
        target_language = dspy.InputField(
            desc="Target language for the rewriting")
        written_text = dspy.OutputField(
            desc="Polished written version of the text")

    rewrite_module = dspy.Predict(RewriteSignature)

    if verbose:
        logger.debug("LLM rewrite module initialized")

    # Read and segment the input text
    segmented_text = segment(input_file, chunk_length,
                             cite_timestamps, verbose=verbose)

    # Split into chunks for processing
    chunks = segmented_text.split("\n\n")
    if verbose:
        logger.debug(f"Text divided into {len(chunks)} chunks for processing")

    rewritten_chunks = []

    for i, chunk in enumerate(chunks, 1):
        if not chunk.strip():
            rewritten_chunks.append(chunk)
            continue

        if verbose:
            logger.debug(f"Rewriting chunk {
                         i}/{len(chunks)} ({len(chunk)} characters)")

        try:
            # Skip timestamp headers when cite_timestamps is True
            is_timestamp_header = chunk.strip().startswith(
                "### **"
            ) and chunk.strip().endswith("**")

            if cite_timestamps and is_timestamp_header:
                if verbose:
                    logger.debug(
                        f"Preserving timestamp header: {chunk.strip()[:50]}..."
                    )
                rewritten_chunks.append(chunk)
            else:
                result = rewrite_module(
                    oral_text=chunk, target_language=rewrite_language
                )
                rewritten_chunks.append(result.written_text)

                if verbose:
                    logger.debug(f"Chunk {i} rewritten successfully")

        except Exception as e:
            error_msg = f"Error rewriting chunk {i}: {e}"
            if verbose:
                logger.debug(error_msg)
            print(error_msg)
            rewritten_chunks.append(f"[Rewrite Error: {chunk}]")

    final_rewrite = "\n\n".join(rewritten_chunks)

    if verbose:
        logger.debug(
            f"Rewrite completed. Final text length: {
                len(final_rewrite)} characters"
        )
        logger.debug("=== Rewrite Process Completed ===")

    # Write output file if output_dir is provided
    output_file = None
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}_rewritten.md")
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(final_rewrite)

        if verbose:
            logger.debug(f"Rewrite output saved to: {output_file}")

    return final_rewrite, output_file


def academic(
    input_file,
    output_dir="",
    llm="ollama/qwen3",
    academic_lang="English",
    chunk_length=20,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    cite_timestamps=False,
    verbose=False,
):
    """
    Convert text to academic writing style using LLM with verbose logging support.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Academic Writing Process ===")
        logger.debug(f"Input file: {input_file}")
        logger.debug(f"Academic language: {academic_lang}")
        logger.debug(f"LLM model: {llm}")
        logger.debug(f"Chunk length: {chunk_length}")
        logger.debug(f"Max tokens: {max_tokens}")
        logger.debug(f"Include timestamps: {cite_timestamps}")

    # Configure LLM
    lm = configure_lm(
        llm,
        verbose=verbose,
        max_tokens=max_tokens,
        timeout=timeout,
        temperature=temperature,
    )

    class AcademicSignature(dspy.Signature):
        """
        Rewrite this text in formal academic style in {academic_lang}. Follow these rules strictly:
        1. Correct any basic grammar, punctuation, or usage errors.
        2. Improve clarity while preserving the original meaning and scholarly tone (trying your best not to change the structure of sentence)
        3. Maintaining the original meaning and length (97% of original)
        4. IMPORTANT: Preserve ALL footnote references (e.g., [^1], [^2]) exactly as they appear
        """

        input_text = dspy.InputField(
            desc="Original text to be transformed into academic style"
        )
        target_language = dspy.InputField(
            desc="Target language for academic writing")
        academic_text = dspy.OutputField(
            desc="Text rewritten in formal academic style")

    academic_module = dspy.Predict(AcademicSignature)

    if verbose:
        logger.debug("LLM academic writing module initialized")

    # Handle different input file types
    if input_file.lower().endswith(".docx"):
        if verbose:
            logger.debug("Processing DOCX file")
        content = process_docx(input_file, verbose=verbose)
        segmented_text = content
    else:
        # Read and segment the input text
        segmented_text = segment(
            input_file, chunk_length, cite_timestamps, verbose=verbose
        )

    # Split into chunks for processing
    chunks = segmented_text.split("\n\n")
    if verbose:
        logger.debug(f"Text divided into {len(chunks)} chunks for processing")

    academic_chunks = []

    for i, chunk in enumerate(chunks, 1):
        if not chunk.strip():
            academic_chunks.append(chunk)
            continue

        if verbose:
            logger.debug(
                f"Converting chunk {i}/{len(chunks)} to academic style ({
                    len(chunk)
                } characters)"
            )

        try:
            # Skip timestamp headers when cite_timestamps is True
            is_timestamp_header = chunk.strip().startswith(
                "### **"
            ) and chunk.strip().endswith("**")

            if cite_timestamps and is_timestamp_header:
                if verbose:
                    logger.debug(
                        f"Preserving timestamp header: {chunk.strip()[:50]}..."
                    )
                academic_chunks.append(chunk)
            else:
                result = academic_module(
                    input_text=chunk, target_language=academic_lang
                )
                academic_chunks.append(result.academic_text)

                if verbose:
                    logger.debug(
                        f"Chunk {i} converted to academic style successfully")

        except Exception as e:
            error_msg = f"Error converting chunk {i} to academic style: {e}"
            if verbose:
                logger.debug(error_msg)
            print(error_msg)
            academic_chunks.append(f"[Academic Conversion Error: {chunk}]")

    final_academic = "\n\n".join(academic_chunks)

    if verbose:
        logger.debug(
            f"Academic conversion completed. Final text length: {
                len(final_academic)
            } characters"
        )
        logger.debug("=== Academic Writing Process Completed ===")

    return final_academic


def process_docx(input_file, verbose=False):
    """
    Process DOCX files with verbose logging support.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug(f"Processing DOCX file: {input_file}")

    try:
        from docx import Document

        doc = Document(input_file)
        content = []

        paragraph_count = len(doc.paragraphs)
        if verbose:
            logger.debug(f"Found {paragraph_count} paragraphs in DOCX")

        for i, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
                if verbose and i % 50 == 0:  # Log progress every 50 paragraphs
                    logger.debug(f"Processed paragraph {i}/{paragraph_count}")

        result = "\n\n".join(content)

        if verbose:
            logger.debug(
                f"DOCX processing completed. Extracted {len(content)} paragraphs, {
                    len(result)
                } characters total"
            )

        return result

    except ImportError:
        error_msg = "python-docx library not installed. Please install it with: pip install python-docx"
        if verbose:
            logger.debug(error_msg)
        raise ImportError(error_msg)
    except Exception as e:
        error_msg = f"Error processing DOCX file: {e}"
        if verbose:
            logger.debug(error_msg)
        raise Exception(error_msg)


def read_markdown_file(file_path, verbose=False):
    """
    Read markdown file content with verbose logging support.
    
    Returns: (markdown_content, file_path)
    """
    logger = logging.getLogger(__name__)
    
    if verbose:
        logger.debug(f"Reading markdown file: {file_path}")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if verbose:
            logger.debug(f"Markdown file read successfully. Content length: {len(content)} characters")
        
        return content, file_path
    
    except Exception as e:
        error_msg = f"Error reading markdown file: {e}"
        if verbose:
            logger.debug(error_msg)
        raise Exception(error_msg)


def convert_slides_to_markdown(
    slides_file,
    output_dir="",
    image_export_mode="embedded",
    verbose=False,
):
    """
    Convert PPTX/PDF to markdown using marker-pdf Python API with verbose logging support.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Slides to Markdown Conversion ===")
        logger.debug(f"Slides file: {slides_file}")
        logger.debug(f"Image export mode: {image_export_mode}")

    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict

        if verbose:
            logger.debug("Loading marker-pdf models...")
        
        # Create model dict for conversion
        artifact_dict = create_model_dict()

        if verbose:
            logger.debug("Initializing PDF converter...")

        # Convert slides to markdown
        converter = PdfConverter(
            artifact_dict=artifact_dict,
        )
        
        markdown_output = converter(slides_file)
        markdown_text = markdown_output.markdown

        if verbose:
            logger.debug(
                f"Conversion completed. Markdown length: {len(markdown_text)} characters"
            )

        # Handle image export mode
        if image_export_mode == "none":
            # Remove image references from markdown
            import re
            markdown_text = re.sub(r'!\[.*?\]\(.*?\)', '', markdown_text)
            if verbose:
                logger.debug("Stripped image references from markdown")

        # Save to file if output_dir provided
        output_file = None
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            base_name = os.path.splitext(os.path.basename(slides_file))[0]
            output_file = os.path.join(output_dir, f"{base_name}_slides.md")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(markdown_text)

            if verbose:
                logger.debug(f"Slides markdown saved to: {output_file}")

        return markdown_text, output_file

    except ImportError as e:
        error_msg = f"marker-pdf library not installed. Please install it with: pip install marker-pdf"
        if verbose:
            logger.debug(error_msg)
        raise ImportError(error_msg)
    except Exception as e:
        error_msg = f"Error converting slides to markdown: {e}"
        if verbose:
            logger.debug(error_msg)
        raise Exception(error_msg)


def convert_single_slide_image(
    image_path,
    langs=['Chinese', 'English'],
    output_dir=None,
    verbose=False,
):
    """
    OCR a single image using marker-pdf functionality
    Returns dictionary with OCR text and metadata
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Single Image OCR ===")
        logger.debug(f"Image file: {image_path}")

    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from PIL import Image
        import tempfile

        if verbose:
            logger.debug("Loading marker-pdf models...")
        
        # Create model dict for conversion
        artifact_dict = create_model_dict()

        if verbose:
            logger.debug("Processing single image...")

        # Create a temporary PDF from the image to use marker-pdf
        temp_dir = tempfile.mkdtemp()
        temp_pdf = os.path.join(temp_dir, "temp.pdf")
        
        # Convert image to PDF
        img = Image.open(image_path)
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        
        img.save(temp_pdf, "PDF", resolution=150.0)
        
        # Convert PDF to markdown
        converter = PdfConverter(
            artifact_dict=artifact_dict,
        )
        
        markdown_output = converter(temp_pdf)
        
        if verbose:
            logger.debug(f"OCR completed. Text length: {len(markdown_output.markdown)} characters")

        # Cleanup temp file
        import shutil
        shutil.rmtree(temp_dir)

        # Extract confidence and metadata (marker doesn't provide this for images directly)
        result = {
            'text': markdown_output.markdown.strip(),
            'confidence': 0.8,  # Default confidence for image OCR
            'page_id': 1,
            'metadata': {
                'source': image_path,
                'languages': langs
            }
        }

        return result

    except ImportError as e:
        error_msg = f"marker-pdf library not installed. Please install it with: pip install marker-pdf"
        if verbose:
            logger.debug(error_msg)
        raise ImportError(error_msg)
    except Exception as e:
        error_msg = f"Error OCR processing image: {e}"
        if verbose:
            logger.debug(error_msg)
        raise Exception(error_msg)


def align_slides_with_speech(
    slide_content,
    speech_content,
    llm="ollama/qwen3",
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    verbose=False,
):
    """
    Find speech content that aligns with slide content using LLM with verbose logging support.
    Uses both exact phrase matching and semantic similarity.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Slide-Speech Alignment ===")
        logger.debug(f"LLM model: {llm}")
        logger.debug(f"Slide content length: {len(slide_content)} characters")
        logger.debug(f"Speech content length: {len(speech_content)} characters")

    # Configure LLM
    lm = configure_lm(
        llm,
        verbose=verbose,
        max_tokens=max_tokens,
        timeout=timeout,
        temperature=temperature,
    )

    class AlignSlideWithSpeech(dspy.Signature):
        """
        1. Find the given slide content aligns with speech content.
        2. Use both exact phrase matching and semantic similarity.
        3. Then put the slide content before the alignment of speech content.
        4. keep the slide content and speech content 99% unchange. 
        """

        slide_content = dspy.InputField(
            desc="Content from a slide in the presentation"
        )
        speech_content = dspy.InputField(
            desc="Full transcribed speech content"
        )
        aligned_speech_section = dspy.OutputField(
            desc="The matching speech section that aligns with the slide. Return 'NO_MATCH' if no alignment found."
        )
        confidence = dspy.OutputField(
            desc="Confidence level of alignment: high, medium, low, or none"
        )

    align_module = dspy.Predict(AlignSlideWithSpeech)

    if verbose:
        logger.debug("LLM alignment module initialized")

    try:
        result = align_module(slide_content=slide_content, speech_content=speech_content)

        if verbose:
            logger.debug(f"Alignment result - Confidence: {result.confidence}")
            logger.debug(f"Aligned section length: {len(result.aligned_speech_section)} characters")

        return {
            "aligned_section": result.aligned_speech_section,
            "confidence": result.confidence.lower(),
        }

    except Exception as e:
        error_msg = f"Error during alignment: {e}"
        if verbose:
            logger.debug(error_msg)
        logger.debug("=== Slide-Speech Alignment Completed ===")
        return {
            "aligned_section": "NO_MATCH",
            "confidence": "none",
        }


def combine_speech_and_slides(
    speech_markdown,
    slides_markdown,
    llm="ollama/qwen3",
    output_dir="",
    cite_timestamps=False,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    verbose=False,
):
    """
    Combine speech and slides markdown by finding and inserting slide content
    before matching speech sections. Uses LLM for alignment matching.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Speech and Slides Combination ===")
        logger.debug(f"Speech content length: {len(speech_markdown)} characters")
        logger.debug(f"Slides content length: {len(slides_markdown)} characters")
        logger.debug(f"Cite timestamps: {cite_timestamps}")

    # Split slides into individual slides (separated by heading markers)
    slides = _extract_slides(slides_markdown, verbose)

    if verbose:
        logger.debug(f"Extracted {len(slides)} slides from presentation")

    # Split speech into paragraphs - robust content preservation
    # Try multiple splitting strategies to ensure maximum content preservation
    strategies = [
        lambda x: [p.strip() for p in x.split("\n\n") if p.strip()],  # Double newlines
        lambda x: [p.strip() for p in x.split("\n") if p.strip()],    # Single newlines
        lambda x: [p.strip() for p in x.replace('\n\n', '\n').split('\n') if p.strip()],  # Normalize then split
    ]
    
    best_paragraphs = []
    best_content_ratio = 0
    
    for strategy in strategies:
        test_paragraphs = strategy(speech_markdown)
        content_ratio = sum(len(p) for p in test_paragraphs) / len(speech_markdown) if speech_markdown else 0
        
        if content_ratio > best_content_ratio and len(test_paragraphs) > 0:
            best_content_ratio = content_ratio
            best_paragraphs = test_paragraphs
    
    speech_paragraphs = best_paragraphs
    
    # Fallback: if all strategies fail, use the entire content as one paragraph
    if not speech_paragraphs and speech_markdown.strip():
        speech_paragraphs = [speech_markdown.strip()]
        if verbose:
            logger.debug("Used fallback: entire speech as single paragraph")

    if verbose:
        logger.debug(f"Speech split into {len(speech_paragraphs)} paragraphs")
        logger.debug(f"Original speech content length: {len(speech_markdown)} characters")
        logger.debug(f"Total speech paragraphs content length: {sum(len(p) for p in speech_paragraphs)} characters")
        # Check for content preservation
        preserved_ratio = sum(len(p) for p in speech_paragraphs) / len(speech_markdown) if speech_markdown else 0
        logger.debug(f"Content preservation ratio: {preserved_ratio:.2%}")

    # Track which speech paragraphs have been used
    used_indices = set()
    aligned_results = []

    # Process each slide in order
    for slide_idx, slide_content in enumerate(slides, 1):
        if verbose:
            logger.debug(f"Processing slide {slide_idx}/{len(slides)}")

        # Try to find alignment in speech
        alignment = align_slides_with_speech(
            slide_content,
            speech_markdown,
            llm=llm,
            max_tokens=max_tokens,
            timeout=timeout,
            temperature=temperature,
            verbose=verbose,
        )

        if alignment["confidence"] in ["high", "medium"]:
            # Find the matching paragraph index
            matched_idx = _find_matching_paragraph(
                alignment["aligned_section"], speech_paragraphs, verbose
            )

            if matched_idx is not None:
                aligned_results.append(
                    {
                        "slide_num": slide_idx,
                        "slide_content": slide_content,
                        "speech_idx": matched_idx,
                        "confidence": alignment["confidence"],
                    }
                )
                used_indices.add(matched_idx)
                if verbose:
                    logger.debug(
                        f"Slide {slide_idx} aligned with speech paragraph {matched_idx} (confidence: {alignment['confidence']})"
                    )
            else:
                if verbose:
                    logger.debug(f"Slide {slide_idx} - Could not find matching paragraph")
                aligned_results.append(
                    {
                        "slide_num": slide_idx,
                        "slide_content": slide_content,
                        "speech_idx": None,
                        "confidence": "none",
                    }
                )
        else:
            if verbose:
                logger.debug(f"Slide {slide_idx} - No alignment found (confidence: {alignment['confidence']})")
            aligned_results.append(
                {
                    "slide_num": slide_idx,
                    "slide_content": slide_content,
                    "speech_idx": None,
                    "confidence": "none",
                }
            )

    # Build combined markdown by inserting slides before their matching speech sections
    combined_content = _build_combined_markdown(
        speech_paragraphs, aligned_results, cite_timestamps, verbose
    )
    
    # Final content preservation check
    if verbose:
        original_speech_length = len(speech_markdown)
        final_combined_length = len(combined_content)
        preservation_ratio = final_combined_length / original_speech_length if original_speech_length else 0
        logger.debug(f"Content preservation check: {preservation_ratio:.2%}")
        
        # If preservation is too low, warn and potentially add missing content
        if preservation_ratio < 0.80:
            logger.warning(f"Low content preservation detected: {preservation_ratio:.2%}")
            # Check if any speech content is completely missing
            speech_content_in_combined = " ".join(speech_paragraphs)
            if len(speech_content_in_combined) < original_speech_length * 0.80:
                logger.warning("Significant speech content may be missing from combined output!")

    if verbose:
        logger.debug(f"Combined markdown length: {len(combined_content)} characters")
        # Calculate final preservation ratio
        original_speech_length = len(speech_markdown)
        final_combined_length = len(combined_content)
        preservation_ratio = final_combined_length / original_speech_length if original_speech_length else 0
        logger.debug(f"Final content preservation ratio: {preservation_ratio:.2%}")
        if preservation_ratio < 0.90:
            logger.warning(f"Content preservation ratio is low: {preservation_ratio:.2%} - some speech content may be missing!")
        logger.debug("=== Speech and Slides Combination Completed ===")

    return combined_content


def combine_speech_and_slides_enhanced(
    speech_markdown,
    slides_markdown,
    llm="ollama/qwen3",
    output_dir="",
    cite_timestamps=False,
    max_tokens=50000,
    timeout=3600,
    temperature=0.1,
    verbose=False,
):
    """
    Enhanced version of combine_speech_and_slides using multi-layered similarity analysis
    instead of LLM-based alignment. Implements keyword extraction, semantic similarity,
    temporal constraints, and content preservation guarantees.
    """
    logger = logging.getLogger(__name__)

    if verbose:
        logger.debug("=== Starting Enhanced Speech and Slides Combination ===")
        logger.debug(f"Speech content length: {len(speech_markdown)} characters")
        logger.debug(f"Slides content length: {len(slides_markdown)} characters")

    # Import enhanced functions
    from wenbi.enhanced_combination import (
        calculate_combined_similarity,
        create_similarity_matrix,
        find_optimal_alignment,
        distribute_unaligned_slides,
        build_enhanced_combined_markdown,
        _extract_slides as enhanced_extract_slides
    )

    # Step 1: Extract slides from markdown
    slides = enhanced_extract_slides(slides_markdown, verbose)

    if verbose:
        logger.debug(f"Extracted {len(slides)} slides from presentation")

    # Step 2: Split speech into paragraphs (reuse existing logic)
    strategies = [
        lambda x: [p.strip() for p in x.split("\n\n") if p.strip()],  # Double newlines
        lambda x: [p.strip() for p in x.split("\n") if p.strip()],    # Single newlines
        lambda x: [p.strip() for p in x.replace('\n\n', '\n').split('\n') if p.strip()],  # Normalize then split
    ]
    
    best_paragraphs = []
    best_content_ratio = 0
    
    for strategy in strategies:
        test_paragraphs = strategy(speech_markdown)
        content_ratio = sum(len(p) for p in test_paragraphs) / len(speech_markdown) if speech_markdown else 0
        
        if content_ratio > best_content_ratio and len(test_paragraphs) > 0:
            best_content_ratio = content_ratio
            best_paragraphs = test_paragraphs
    
    speech_paragraphs = best_paragraphs
    
    # Fallback: if all strategies fail, use entire content as one paragraph
    if not speech_paragraphs and speech_markdown.strip():
        speech_paragraphs = [speech_markdown.strip()]
        if verbose:
            logger.debug("Used fallback: entire speech as single paragraph")

    if verbose:
        logger.debug(f"Speech split into {len(speech_paragraphs)} paragraphs")
        preserved_ratio = sum(len(p) for p in speech_paragraphs) / len(speech_markdown) if speech_markdown else 0
        logger.debug(f"Content preservation ratio: {preserved_ratio:.2%}")

    # Step 3: Create similarity matrix for all slide-speech pairs
    similarity_matrix = create_similarity_matrix(slides, speech_paragraphs, verbose)

    # Step 4: Find optimal alignment with hybrid temporal constraints
    aligned_slides = find_optimal_alignment(similarity_matrix, len(slides), len(speech_paragraphs), verbose)

    # Step 5: Handle unaligned slides with even distribution
    aligned_slides = distribute_unaligned_slides(slides, speech_paragraphs, aligned_slides, verbose)

    # Step 6: Build enhanced combined markdown with content preservation
    combined_content = build_enhanced_combined_markdown(
        speech_paragraphs, slides, aligned_slides, cite_timestamps, verbose
    )

    # Step 7: Final content preservation check
    if verbose:
        original_speech_length = len(speech_markdown)
        final_combined_length = len(combined_content)
        preservation_ratio = final_combined_length / original_speech_length if original_speech_length else 0
        logger.debug(f"Final content preservation ratio: {preservation_ratio:.2%}")
        
        # Count aligned vs unaligned slides
        aligned_count = len([slide for slide, speeches in aligned_slides.items() if speeches])
        unaligned_count = len(slides) - aligned_count
        logger.debug(f"Alignment summary: {aligned_count} aligned, {unaligned_count} distributed slides")
        
        if preservation_ratio < 0.90:
            logger.warning(f"Content preservation ratio is low: {preservation_ratio:.2%} - some speech content may be missing!")
        
        logger.debug("=== Enhanced Speech and Slides Combination Completed ===")

    return combined_content


def _extract_slides(slides_markdown, verbose=False):
    """Extract individual slides from markdown (by splitting on major headings)"""
    logger = logging.getLogger(__name__)

    # Split by ## or # headings (common slide markers)
    import re

    slides = []
    current_slide = []

    lines = slides_markdown.split("\n")
    for line in lines:
        # Check if this is a slide heading (## or #)
        if re.match(r"^#+\s", line) and current_slide:
            # Save current slide and start new one
            slides.append("\n".join(current_slide).strip())
            current_slide = [line]
        else:
            current_slide.append(line)

    # Don't forget last slide
    if current_slide:
        slides.append("\n".join(current_slide).strip())

    if verbose:
        logger.debug(f"Extracted {len(slides)} slides from markdown")

    return [s for s in slides if s.strip()]  # Filter empty slides


def _find_matching_paragraph(matched_section, speech_paragraphs, verbose=False):
    """Find which paragraph index contains the matched section"""
    logger = logging.getLogger(__name__)

    if matched_section == "NO_MATCH":
        return None

    # Try exact match first
    for idx, para in enumerate(speech_paragraphs):
        if matched_section.lower() in para.lower():
            if verbose:
                logger.debug(f"Found exact match at paragraph {idx}")
            return idx

    # Fuzzy match: find most similar paragraph
    best_match_idx = None
    best_similarity = 0

    for idx, para in enumerate(speech_paragraphs):
        # Count overlapping words
        matched_words = set(matched_section.lower().split())
        para_words = set(para.lower().split())
        overlap = len(matched_words & para_words)

        if overlap > best_similarity:
            best_similarity = overlap
            best_match_idx = idx

    if best_match_idx is not None and best_similarity > 2:
        if verbose:
            logger.debug(
                f"Found fuzzy match at paragraph {best_match_idx} with {best_similarity} overlapping words"
            )
        return best_match_idx

    if verbose:
        logger.debug("No matching paragraph found")
    return None


def _build_combined_markdown(speech_paragraphs, aligned_results, cite_timestamps=False, verbose=False):
    """Build combined markdown by inserting slides before matching speech sections"""
    logger = logging.getLogger(__name__)

    # Create a mapping of speech index -> list of slides to insert before it
    slides_before_speech = {}
    unaligned_slides = []

    for result in aligned_results:
        if result["speech_idx"] is not None:
            if result["speech_idx"] not in slides_before_speech:
                slides_before_speech[result["speech_idx"]] = []
            slides_before_speech[result["speech_idx"]].append(result["slide_content"])
        else:
            unaligned_slides.append(result)

    # Build combined content
    combined = []

    for idx, para in enumerate(speech_paragraphs):
        # Insert slides before this paragraph if they align with it
        if idx in slides_before_speech:
            for slide_content in slides_before_speech[idx]:
                combined.append(slide_content)
                combined.append("")  # Blank line separator

        combined.append(para)
        combined.append("")  # Blank line separator

    # Add unaligned slides at the end (in order)
    if unaligned_slides:
        if verbose:
            logger.debug(f"Adding {len(unaligned_slides)} unaligned slides at the end")
        combined.append("---\n\n## Unaligned Slides\n")
        for result in unaligned_slides:
            combined.append(result["slide_content"])
            combined.append("")

    result = "\n".join(combined).strip()

    if verbose:
        logger.debug(f"Combined content built with {len(speech_paragraphs)} speech paragraphs")
        logger.debug(f"Final combined content length: {len(result)} characters")
        # Calculate final preservation ratio (need to pass original speech length)
        # Note: This will be calculated in the calling function

    return result
