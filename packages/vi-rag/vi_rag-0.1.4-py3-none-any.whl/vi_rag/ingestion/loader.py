from ..core.document import DocumentNode
from abc import ABC, abstractmethod
import re
import underthesea
import os
import hashlib
import json
import logging
from typing import Optional, Dict, List, Tuple
from pypdf import PdfReader

# Setup logger
logger = logging.getLogger(__name__)


class BaseLoader(ABC):
    def __init__(self, file_path : str):
        self.file_path = file_path
    @abstractmethod
    def load(self, file_path: str) -> DocumentNode:
        """
        Load a document from the given file_path and return a DocumentNode.
        """
        pass

    @abstractmethod
    def validate(self, file_path: str) -> bool:
        """
        Check if the file at file_path is compatible with this loader.
        """
        pass

    @abstractmethod
    def get_metadata(self, file_path: str) -> dict:
        """
        Extract and return metadata from the document at file_path.
        """
        pass


class PDFLoader(BaseLoader):
    """
    Loader for PDF documents with duplicate detection using MD5 hashing.
    """
    # Class-level storage for loaded documents (MD5 hash -> document info)
    _loaded_documents: Dict[str, Dict] = {}
    _loaded_documents_file = "loaded_documents_cache.json"
    
    # In-memory cache for DocumentNode objects (MD5 hash -> DocumentNode)
    _cached_documents: Dict[str, 'DocumentNode'] = {}
    
    def __init__(self, file_path: str, use_pymupdf: bool = True, cache_dir: str = "."):
        """
        Initialize PDFLoader.
        
        Args:
            file_path: Path to the PDF file
            use_pymupdf: If True, use PyMuPDF. If False, use pypdf.
            cache_dir: Directory to store the loaded documents cache
        """
        super().__init__(file_path)
        self.use_pymupdf = use_pymupdf
        self.cache_file = os.path.join(cache_dir, self._loaded_documents_file)
        self._load_cache()
    
    def _load_cache(self):
        """Load the cache of previously loaded documents from file."""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    PDFLoader._loaded_documents = json.load(f)
            except Exception:
                PDFLoader._loaded_documents = {}
    
    def _save_cache(self):
        """Save the cache of loaded documents to file."""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(PDFLoader._loaded_documents, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"Could not save cache: {e}")
    
    def _calculate_md5(self, text: str) -> str:
        """Calculate MD5 hash of text content."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def check_document_loaded(self) -> Optional['DocumentNode']:
        """
        Check if document has already been loaded.
        Uses self.file_path from initialization.
        
        Returns:
            DocumentNode if already loaded, None otherwise
        """
        if not self.validate(self.file_path):
            return None
        
        # Extract text content and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Check in-memory cache first
        if content_md5 in PDFLoader._cached_documents:
            logger.info(f"Document found in cache: {os.path.basename(self.file_path)}")
            logger.info(f"  MD5: {content_md5}")
            return PDFLoader._cached_documents[content_md5]
        
        return None

    def _safe_extract_page_text(self, page) -> str:
        try:
            return page.extract_text(extraction_mode="layout") or ""
        except Exception:
            try:
                return page.extract_text() or ""
            except Exception:
                return ""

    
    def _extract_text(self, file_path):
        text_chunks = []

        reader = PdfReader(file_path)
        for page in reader.pages:
            text = self._safe_extract_page_text(page)
            if text:
                text_chunks.append(text)

        return "\n".join(text_chunks)


    
    def validate(self, file_path: str) -> bool:
        """
        Check if the file is a valid PDF file.
        """
        if not os.path.exists(file_path):
            return False
        if not file_path.lower().endswith('.pdf'):
            return False
        
        try:
            if self.use_pymupdf:
                try:
                    import fitz
                    doc = fitz.open(file_path)
                    doc.close()
                    return True
                except ImportError:
                    # Fallback to pypdf
                    self.use_pymupdf = False
                    return self.validate(file_path)
            else:
                from pypdf import PdfReader
                PdfReader(file_path)
                return True
        except Exception:
            return False
    
    def get_metadata(self, file_path: str = None) -> dict:
        """
        Extract simple metadata from the PDF file.
        Focus on essential information including parent_id.
        
        Args:
            file_path: Optional path to PDF file. If not provided, uses self.file_path
        """
        target_path = file_path or self.file_path
        metadata = {
            'file_path': target_path,
            'file_name': os.path.basename(target_path),
            'file_type': 'pdf',
            'parent_id': None  # Important: will be set when building hierarchy
        }
        
        try:
            if self.use_pymupdf:
                import fitz
                doc = fitz.open(target_path)
                metadata['page_count'] = doc.page_count
                
                # Add title if available
                pdf_metadata = doc.metadata
                if pdf_metadata and pdf_metadata.get('title'):
                    metadata['title'] = pdf_metadata.get('title')
                
                doc.close()
            else:
                from pypdf import PdfReader
                reader = PdfReader(target_path)
                metadata['page_count'] = len(reader.pages)
                
                # Add title if available
                if reader.metadata and reader.metadata.get('/Title'):
                    metadata['title'] = reader.metadata.get('/Title')
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def load(self) -> DocumentNode:
        """
        Load a PDF document and return a DocumentNode.
        Uses self.file_path from initialization.
        Always returns a DocumentNode (from cache or freshly loaded).
        
        Returns:
            DocumentNode
        """
        if not self.validate(self.file_path):   
            raise ValueError(f"Invalid PDF file: {self.file_path}")
        
        # Check in-memory cache first
        cached_doc = self.check_document_loaded()
        if cached_doc:
            logger.info("Returning cached document")
            return cached_doc
        
        # Extract text and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Get metadata
        metadata = self.get_metadata()
        metadata['content_md5'] = content_md5
        
        # Create root document node with unique ID
        root_node = DocumentNode(
            level=0,
            title=metadata.get('title') or os.path.basename(self.file_path),
            text="",
            parent_id=None,  # Root node has no parent
            metadata=metadata
        )
        
        # Process each page and create child nodes
        if self.use_pymupdf:
            import fitz
            doc = fitz.open(self.file_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text().strip()
                
                # Create unique page node
                page_node = DocumentNode(
                    level=1,
                    title=f"Page {page_num + 1}",
                    text=page_text,
                    parent_id=root_node.id,  # Link to parent
                    metadata={
                        'page_number': page_num + 1,
                        'parent_id': root_node.id  # Important: track parent
                    }
                )
                root_node.add_child(page_node)
            doc.close()
        else:
            reader = PdfReader(self.file_path)

            for page_num, page in enumerate(reader.pages):
                page_text = self._safe_extract_page_text(page).strip()

                page_node = DocumentNode(
                    level=1,
                    title=f"Page {page_num + 1}",
                    text=page_text,
                    parent_id=root_node.id,
                    metadata={
                        'page_number': page_num + 1,
                        'parent_id': root_node.id
                    }
                )
                root_node.add_child(page_node)

        
        # Concatenate all page texts
        root_node.text = "\n\n".join([child.text for child in root_node.children if child.text])
        
        # Register this document as loaded
        from datetime import datetime
        PDFLoader._loaded_documents[content_md5] = {
            'file_name': os.path.basename(self.file_path),
            'original_path': self.file_path,
            'loaded_at': datetime.now().isoformat(),
            'root_node_id': root_node.id,
            'page_count': len(root_node.children)
        }
        self._save_cache()
        
        # Store DocumentNode in memory cache
        PDFLoader._cached_documents[content_md5] = root_node
        
        logger.info(f"Document loaded successfully: {os.path.basename(self.file_path)}")
        logger.debug(f"  MD5: {content_md5}")
        logger.debug(f"  Root Node ID: {root_node.id}")
        logger.debug(f"  Pages: {len(root_node.children)}")
        
        return root_node


class TXTLoader(BaseLoader):
    """
    Loader for TXT documents with duplicate detection using MD5 hashing.
    """
    # Class-level storage for loaded documents (MD5 hash -> document info)
    _loaded_documents: Dict[str, Dict] = {}
    _loaded_documents_file = "loaded_txt_documents_cache.json"
    
    # In-memory cache for DocumentNode objects (MD5 hash -> DocumentNode)
    _cached_documents: Dict[str, 'DocumentNode'] = {}
    
    def __init__(self, file_path: str, encoding: str = 'utf-8', cache_dir: str = "."):
        """
        Initialize TXTLoader.
        
        Args:
            file_path: Path to the TXT file
            encoding: Text encoding (default: utf-8)
            cache_dir: Directory to store the loaded documents cache
        """
        super().__init__(file_path)
        self.encoding = encoding
        self.cache_file = os.path.join(cache_dir, self._loaded_documents_file)
        self._load_cache()
    
    def _load_cache(self):
        """Load the cache of previously loaded documents from file."""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    TXTLoader._loaded_documents = json.load(f)
            except Exception:
                TXTLoader._loaded_documents = {}
    
    def _save_cache(self):
        """Save the cache of loaded documents to file."""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(TXTLoader._loaded_documents, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"Could not save cache: {e}")
    
    def _calculate_md5(self, text: str) -> str:
        """Calculate MD5 hash of text content."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def check_document_loaded(self) -> Optional['DocumentNode']:
        """
        Check if document has already been loaded.
        Uses self.file_path from initialization.
        
        Returns:
            DocumentNode if already loaded, None otherwise
        """
        if not self.validate(self.file_path):
            return None
        
        # Extract text content and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Check in-memory cache first
        if content_md5 in TXTLoader._cached_documents:
            logger.info(f"Document found in cache: {os.path.basename(self.file_path)}")
            logger.info(f"  MD5: {content_md5}")
            return TXTLoader._cached_documents[content_md5]
        
        return None
    
    def _extract_text(self, file_path: str) -> str:
        """Extract all text from TXT file."""
        try:
            with open(file_path, 'r', encoding=self.encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        return f.read()
                except:
                    continue
            raise ValueError(f"Could not decode file with any supported encoding")
    
    def validate(self, file_path: str) -> bool:
        """
        Check if the file is a valid TXT file.
        """
        if not os.path.exists(file_path):
            return False
        if not file_path.lower().endswith('.txt'):
            return False
        
        try:
            # Try to read the file
            self._extract_text(file_path)
            return True
        except Exception:
            return False
    
    def get_metadata(self, file_path: str) -> dict:
        """
        Extract simple metadata from the TXT file.
        Focus on essential information including parent_id.
        """
        metadata = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_type': 'txt',
            'parent_id': None,  # Important: will be set when building hierarchy
            'encoding': self.encoding
        }
        
        try:
            # Get file size
            metadata['file_size'] = os.path.getsize(file_path)
            
            # Count lines
            text = self._extract_text(file_path)
            metadata['line_count'] = len(text.split('\n'))
            metadata['char_count'] = len(text)
            metadata['word_count'] = len(text.split())
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def load(self) -> DocumentNode:
        """
        Load a TXT document and return a DocumentNode.
        Uses self.file_path from initialization.
        Always returns a DocumentNode (from cache or freshly loaded).
        
        Returns:
            DocumentNode
        """
        if not self.validate(self.file_path):
            raise ValueError(f"Invalid TXT file: {self.file_path}")
        
        # Check in-memory cache first
        cached_doc = self.check_document_loaded()
        if cached_doc:
            logger.info("Returning cached document")
            return cached_doc
        
        # Extract text and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Get metadata
        metadata = self.get_metadata(self.file_path)
        metadata['content_md5'] = content_md5
        
        # Create root document node with unique ID
        root_node = DocumentNode(
            level=0,
            title=os.path.basename(self.file_path),
            text=text_content,
            parent_id=None,  # Root node has no parent
            metadata=metadata
        )
        
        # Register this document as loaded
        from datetime import datetime
        TXTLoader._loaded_documents[content_md5] = {
            'file_name': os.path.basename(self.file_path),
            'original_path': self.file_path,
            'loaded_at': datetime.now().isoformat(),
            'root_node_id': root_node.id,
            'line_count': metadata.get('line_count', 0)
        }
        self._save_cache()
        
        # Store DocumentNode in memory cache
        TXTLoader._cached_documents[content_md5] = root_node
        
        logger.info(f"Document loaded successfully: {os.path.basename(self.file_path)}")
        logger.debug(f"  MD5: {content_md5}")
        logger.debug(f"  Root Node ID: {root_node.id}")
        logger.debug(f"  Lines: {metadata.get('line_count', 0)}")
        
        return root_node


class DOCXLoader(BaseLoader):
    """
    Loader for DOCX documents with duplicate detection using MD5 hashing.
    """
    # Class-level storage for loaded documents (MD5 hash -> document info)
    _loaded_documents: Dict[str, Dict] = {}
    _loaded_documents_file = "loaded_docx_documents_cache.json"
    
    # In-memory cache for DocumentNode objects (MD5 hash -> DocumentNode)
    _cached_documents: Dict[str, 'DocumentNode'] = {}
    
    def __init__(self, file_path: str, cache_dir: str = "."):
        """
        Initialize DOCXLoader.
        
        Args:
            file_path: Path to the DOCX file
            cache_dir: Directory to store the loaded documents cache
        """
        super().__init__(file_path)
        self.cache_file = os.path.join(cache_dir, self._loaded_documents_file)
        self._load_cache()
    
    def _load_cache(self):
        """Load the cache of previously loaded documents from file."""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    DOCXLoader._loaded_documents = json.load(f)
            except Exception:
                DOCXLoader._loaded_documents = {}
    
    def _save_cache(self):
        """Save the cache of loaded documents to file."""
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(DOCXLoader._loaded_documents, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.warning(f"Could not save cache: {e}")
    
    def _calculate_md5(self, text: str) -> str:
        """Calculate MD5 hash of text content."""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def check_document_loaded(self) -> Optional['DocumentNode']:
        """
        Check if document has already been loaded.
        Uses self.file_path from initialization.
        
        Returns:
            DocumentNode if already loaded, None otherwise
        """
        if not self.validate(self.file_path):
            return None
        
        # Extract text content and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Check in-memory cache first
        if content_md5 in DOCXLoader._cached_documents:
            logger.info(f"Document found in cache: {os.path.basename(self.file_path)}")
            logger.info(f"  MD5: {content_md5}")
            return DOCXLoader._cached_documents[content_md5]
        
        return None
    
    def _extract_text(self, file_path: str) -> str:
        """Extract all text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install it with: pip install python-docx")
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n\n".join(text_parts)
    
    def validate(self, file_path: str) -> bool:
        """
        Check if the file is a valid DOCX file.
        """
        if not os.path.exists(file_path):
            return False
        if not file_path.lower().endswith('.docx'):
            return False
        
        try:
            from docx import Document
            Document(file_path)
            return True
        except ImportError:
            logger.warning("python-docx not installed")
            return False
        except Exception:
            return False
    
    def get_metadata(self, file_path: str) -> dict:
        """
        Extract simple metadata from the DOCX file.
        Focus on essential information including parent_id.
        """
        metadata = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_type': 'docx',
            'parent_id': None  # Important: will be set when building hierarchy
        }
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Get core properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            
            # Count elements
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            metadata['section_count'] = len(doc.sections)
            
        except Exception as e:
            metadata['error'] = str(e)
        
        return metadata
    
    def load(self) -> DocumentNode:
        """
        Load a DOCX document and return a DocumentNode.
        Uses self.file_path from initialization.
        Always returns a DocumentNode (from cache or freshly loaded).
        
        Returns:
            DocumentNode
        """
        if not self.validate(self.file_path):
            raise ValueError(f"Invalid DOCX file: {self.file_path}")
        
        # Check in-memory cache first
        cached_doc = self.check_document_loaded()
        if cached_doc:
            logger.info("Returning cached document")
            return cached_doc
        
        # Extract text and calculate MD5
        text_content = self._extract_text(self.file_path)
        content_md5 = self._calculate_md5(text_content)
        
        # Get metadata
        metadata = self.get_metadata(self.file_path)
        metadata['content_md5'] = content_md5
        
        # Create root document node with unique ID
        from docx import Document
        doc = Document(self.file_path)
        
        root_node = DocumentNode(
            level=0,
            title=metadata.get('title') or os.path.basename(self.file_path),
            text="",
            parent_id=None,  # Root node has no parent
            metadata=metadata
        )
        
        # Process each paragraph as a child node
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_node = DocumentNode(
                    level=1,
                    title=f"Paragraph {idx + 1}",
                    text=paragraph.text.strip(),
                    parent_id=root_node.id,  # Link to parent
                    metadata={
                        'paragraph_number': idx + 1,
                        'parent_id': root_node.id,  # Important: track parent
                        'style': paragraph.style.name if paragraph.style else None
                    }
                )
                root_node.add_child(para_node)
        
        # Add tables as separate child nodes
        for idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            
            if table_text:
                table_node = DocumentNode(
                    level=1,
                    title=f"Table {idx + 1}",
                    text="\n".join(table_text),
                    parent_id=root_node.id,  # Link to parent
                    metadata={
                        'table_number': idx + 1,
                        'parent_id': root_node.id,  # Important: track parent
                        'row_count': len(table.rows),
                        'column_count': len(table.columns)
                    }
                )
                root_node.add_child(table_node)
        
        # Concatenate all child texts
        root_node.text = "\n\n".join([child.text for child in root_node.children if child.text])
        
        # Register this document as loaded
        from datetime import datetime
        DOCXLoader._loaded_documents[content_md5] = {
            'file_name': os.path.basename(self.file_path),
            'original_path': self.file_path,
            'loaded_at': datetime.now().isoformat(),
            'root_node_id': root_node.id,
            'paragraph_count': len([c for c in root_node.children if 'paragraph_number' in c.metadata]),
            'table_count': len([c for c in root_node.children if 'table_number' in c.metadata])
        }
        self._save_cache()
        
        # Store DocumentNode in memory cache
        DOCXLoader._cached_documents[content_md5] = root_node
        
        logger.info(f"Document loaded successfully: {os.path.basename(self.file_path)}")
        logger.debug(f"  MD5: {content_md5}")
        logger.debug(f"  Root Node ID: {root_node.id}")
        logger.debug(f"  Paragraphs: {metadata.get('paragraph_count', 0)}")
        logger.debug(f"  Tables: {metadata.get('table_count', 0)}")
        
        return root_node


class DocumentLoader:
    """
    Universal document loader that automatically detects file type and uses appropriate loader.
    Supports PDF, TXT, and DOCX files with MD5 duplicate detection.
    
    Usage:
        loader = DocumentLoader("document.pdf")
        doc = loader.load()
    """
    
    # Map of file extensions to loader classes
    LOADER_MAP = {
        '.pdf': PDFLoader,
        '.txt': TXTLoader,
        '.docx': DOCXLoader,
    }
    
    def __init__(
        self, 
        file_path: str,
        cache_dir: str = ".",
        auto_chunk: bool = False,
        parent_size: int = 2000,
        child_size: int = 400,
        overlap: int = 50,
        **kwargs
    ):
        """
        Initialize DocumentLoader with automatic file type detection.
        
        Args:
            file_path: Path to the document file
            cache_dir: Directory to store cache files
            auto_chunk: If True, automatically chunk documents (for load_and_chunk)
            parent_size: Size of parent chunks (default: 2000)
            child_size: Size of child chunks (default: 400)
            overlap: Overlap between chunks (default: 50)
            **kwargs: Additional arguments passed to specific loaders
                     - use_pymupdf: bool (for PDFLoader)
                     - encoding: str (for TXTLoader)
        """
        self.file_path = file_path
        self.cache_dir = cache_dir
        self.auto_chunk = auto_chunk
        self.parent_size = parent_size
        self.child_size = child_size
        self.overlap = overlap
        self.kwargs = kwargs
        
        # Detect file type and initialize appropriate loader
        self.file_extension = self._get_file_extension(file_path)
        self.loader = self._create_loader()
    
    def _get_file_extension(self, file_path: str) -> str:
        """Get file extension in lowercase."""
        _, ext = os.path.splitext(file_path)
        return ext.lower()
    
    def _create_loader(self) -> BaseLoader:
        """Create appropriate loader based on file extension."""
        if self.file_extension not in self.LOADER_MAP:
            supported = ', '.join(self.LOADER_MAP.keys())
            raise ValueError(
                f"Unsupported file type: {self.file_extension}. "
                f"Supported types: {supported}"
            )
        
        loader_class = self.LOADER_MAP[self.file_extension]
        
        # Prepare arguments for specific loaders
        loader_kwargs = {'file_path': self.file_path, 'cache_dir': self.cache_dir}
        
        # Add specific arguments based on loader type
        if loader_class == PDFLoader:
            if 'use_pymupdf' in self.kwargs:
                loader_kwargs['use_pymupdf'] = self.kwargs['use_pymupdf']
        elif loader_class == TXTLoader:
            if 'encoding' in self.kwargs:
                loader_kwargs['encoding'] = self.kwargs['encoding']
        
        return loader_class(**loader_kwargs)
    
    def load(self) -> DocumentNode:
        """
        Load the document using the appropriate loader.
        Always returns a DocumentNode (from cache or freshly loaded).
        
        Returns:
            DocumentNode
        """
        return self.loader.load()
    
    def validate(self) -> bool:
        """
        Validate the document file.
        
        Returns:
            True if file is valid, False otherwise
        """
        return self.loader.validate(self.loader.file_path)
    
    def get_metadata(self) -> dict:
        """
        Get metadata from the document.
        
        Returns:
            Dictionary containing document metadata
        """
        # Need to get metadata from specific loader methods
        return self.loader.get_metadata(self.loader.file_path)
    
    def check_document_loaded(self) -> Optional['DocumentNode']:
        """
        Check if document has already been loaded.
        
        Returns:
            DocumentNode if already loaded, None otherwise
        """
        return self.loader.check_document_loaded()
    
    def load_and_chunk(self) -> Tuple['DocumentNode', List[Dict], List[Dict]]:
        """
        Load document and create hierarchical chunks automatically.
        
        This method ALWAYS performs chunking when called, regardless of the auto_chunk setting.
        If you want to load a document without chunking, use the load() method instead.
        
        Returns:
            Tuple of (document, parents, children)
            - document: The loaded DocumentNode
            - parents: List of parent-level chunks
            - children: List of child-level chunks
            
        Raises:
            ValueError: If document loading or chunking fails
        """
        # Load document (from cache or fresh)
        try:
            document = self.load()
            logger.info(f"Document loaded for chunking: {os.path.basename(self.file_path)}")
            logger.debug(f"  Document ID: {document.id}")
            logger.debug(f"  Text length: {len(document.text)}")
            logger.debug(f"  Child nodes: {len(document.children)}")
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise ValueError(f"Failed to load document: {e}") from e
        
        # Import and create chunker
        try:
            from .chunker import HierarchicalChunker
            
            # Create chunks - always chunk when this method is explicitly called
            logger.info(f"Chunking with settings: parent_size={self.parent_size}, child_size={self.child_size}, overlap={self.overlap}")
            chunker = HierarchicalChunker(
                parent_size=self.parent_size,
                child_size=self.child_size,
                overlap=self.overlap
            )
            parents, children = chunker.build_chunks(document)
            
            # Validate results
            if parents is None or children is None:
                raise ValueError("Chunking returned None")
            
            logger.info(f"Chunking complete: {len(parents)} parent chunks, {len(children)} child chunks")
            
            return document, parents, children
            
        except ImportError as e:
            logger.error(f"Failed to import chunker: {e}")
            raise ValueError("HierarchicalChunker not available") from e
        except Exception as e:
            logger.error(f"Error during chunking: {e}")
            raise ValueError(f"Failed to chunk document: {e}") from e
    
    def get_loader_type(self) -> str:
        """
        Get the type of loader being used.
        
        Returns:
            Name of the loader class
        """
        return self.loader.__class__.__name__
    
    @staticmethod
    def get_supported_extensions() -> list:
        """
        Get list of supported file extensions.
        
        Returns:
            List of supported file extensions
        """
        return list(DocumentLoader.LOADER_MAP.keys())
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file type is supported, False otherwise
        """
        _, ext = os.path.splitext(file_path)
        return ext.lower() in DocumentLoader.LOADER_MAP
    
    @classmethod
    def load_multiple(cls, file_paths: list, cache_dir: str = ".", 
                     **kwargs) -> Dict[str, Optional[DocumentNode]]:
        """
        Load multiple documents at once.
        
        Args:
            file_paths: List of file paths to load
            cache_dir: Directory to store cache files
            **kwargs: Additional arguments passed to specific loaders
            
        Returns:
            Dictionary mapping file paths to DocumentNodes (or None if failed)
            
        Note:
            Documents are automatically cached. If a document has been loaded before,
            the cached version will be returned.
        """
        results = {}
        
        for file_path in file_paths:
            try:
                if not cls.is_supported(file_path):
                    logger.warning(f"Skipping unsupported file: {file_path}")
                    results[file_path] = None
                    continue
                
                loader = cls(file_path, cache_dir=cache_dir, **kwargs)
                doc_node = loader.load()
                results[file_path] = doc_node
                
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                results[file_path] = None
        
        return results
    
    def __repr__(self) -> str:
        """String representation of DocumentLoader."""
        return f"DocumentLoader(file='{os.path.basename(self.file_path)}', type='{self.file_extension}', loader='{self.get_loader_type()}')"
        