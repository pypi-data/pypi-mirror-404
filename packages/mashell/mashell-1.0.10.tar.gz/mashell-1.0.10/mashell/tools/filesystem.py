"""Native filesystem tools - read/list/search without shell dependency."""

import fnmatch
import os
import re
from pathlib import Path
from typing import Any

from mashell.tools.base import BaseTool, ToolResult


def is_binary(data: bytes, sample_size: int = 8192) -> bool:
    """Check if data appears to be binary."""
    if not data:
        return False  # Empty file is not binary
    # Check for null bytes (strong indicator of binary)
    if b"\x00" in data[:sample_size]:
        return True
    # Check ratio of non-text bytes
    text_chars = bytearray({7, 8, 9, 10, 12, 13, 27} | set(range(0x20, 0x100)) - {0x7F})
    check_size = min(len(data), sample_size)
    non_text = sum(1 for byte in data[:check_size] if byte not in text_chars)
    return non_text / check_size > 0.30


def smart_truncate(content: str, max_lines: int = 500, context_lines: int = 50) -> str:
    """Truncate content intelligently, keeping head and tail."""
    lines = content.splitlines()
    if len(lines) <= max_lines:
        return content

    head = lines[:context_lines]
    tail = lines[-context_lines:]
    omitted = len(lines) - 2 * context_lines

    return "\n".join(head) + f"\n\n[... {omitted} lines omitted ...]\n\n" + "\n".join(tail)


def add_line_numbers(content: str, start_line: int = 1) -> str:
    """Add line numbers to content."""
    lines = content.splitlines()
    width = len(str(start_line + len(lines) - 1))
    numbered = [f"{i + start_line:>{width}} │ {line}" for i, line in enumerate(lines)]
    return "\n".join(numbered)


class ReadFileTool(BaseTool):
    """Read file contents with smart handling."""

    name = "read_file"
    description = """Read the contents of a file.

Returns file content with line numbers for easy reference.
Handles binary files, large files, and encoding issues gracefully.

## Parameters
- path: File path (absolute or relative to working directory)
- start_line: Optional, start reading from this line (1-indexed)
- end_line: Optional, stop reading at this line (inclusive)

## Examples
- Read entire file: read_file("src/main.py")
- Read lines 10-50: read_file("config.yaml", start_line=10, end_line=50)
- Read first 100 lines: read_file("large.log", end_line=100)"""

    parameters: dict[str, Any] = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to the file to read"},
            "start_line": {
                "type": "integer",
                "description": "Start reading from this line (1-indexed, optional)",
            },
            "end_line": {
                "type": "integer",
                "description": "Stop reading at this line (inclusive, optional)",
            },
        },
        "required": ["path"],
    }

    requires_permission = False  # Reading is safe
    permission_level = "auto"

    def _read_pdf(self, file_path: Path) -> str | None:
        """Try to read PDF file content."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i} ---\n{page_text}")
            return "\n\n".join(text_parts) if text_parts else None
        except ImportError:
            return None  # pdfplumber not installed
        except Exception:
            return None

    def _read_docx(self, file_path: Path) -> str | None:
        """Try to read Word document content."""
        try:
            from docx import Document

            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs) if paragraphs else None
        except ImportError:
            return None  # python-docx not installed
        except Exception:
            return None

    async def execute(
        self,
        path: str,
        start_line: int | None = None,
        end_line: int | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        """Read file contents."""
        try:
            file_path = Path(path).expanduser()

            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            if not file_path.is_file():
                return ToolResult(success=False, output="", error=f"Not a file: {path}")

            # Check for special file types
            suffix = file_path.suffix.lower()

            # Handle PDF
            if suffix == ".pdf":
                content = self._read_pdf(file_path)
                if content is None:
                    return ToolResult(
                        success=False,
                        output="",
                        error="Cannot read PDF (install pdfplumber: pip install pdfplumber)",
                    )
                output = f"[{path}] (PDF)\n\n{content}"
                return ToolResult(success=True, output=smart_truncate(output, max_lines=1000))

            # Handle Word documents
            if suffix in (".docx", ".doc"):
                if suffix == ".doc":
                    return ToolResult(
                        success=False, output="", error="Old .doc format not supported, only .docx"
                    )
                content = self._read_docx(file_path)
                if content is None:
                    return ToolResult(
                        success=False,
                        output="",
                        error="Cannot read DOCX (install python-docx: pip install python-docx)",
                    )
                output = f"[{path}] (Word Document)\n\n{content}"
                return ToolResult(success=True, output=smart_truncate(output, max_lines=1000))

            # Read raw bytes first to check for binary
            raw_data = file_path.read_bytes()

            if is_binary(raw_data):
                size = len(raw_data)
                return ToolResult(success=True, output=f"[Binary file, {size:,} bytes]")

            # Try to decode
            try:
                content = raw_data.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    content = raw_data.decode("latin-1")
                except UnicodeDecodeError:
                    return ToolResult(
                        success=False, output="", error=f"Unable to decode file: {path}"
                    )

            # Handle line range
            lines = content.splitlines()
            total_lines = len(lines)

            if start_line is not None or end_line is not None:
                start_idx = (start_line - 1) if start_line else 0
                end_idx = end_line if end_line else total_lines
                start_idx = max(0, start_idx)
                end_idx = min(total_lines, end_idx)

                lines = lines[start_idx:end_idx]
                content = "\n".join(lines)
                line_offset = start_idx + 1

                # Add line numbers
                output = add_line_numbers(content, line_offset)
                output = f"[{path}] Lines {start_idx + 1}-{end_idx} of {total_lines}\n\n{output}"
            else:
                # Smart truncate for large files
                if total_lines > 500:
                    content = smart_truncate(content)

                output = add_line_numbers(content)
                output = f"[{path}] {total_lines} lines\n\n{output}"

            return ToolResult(success=True, output=output)

        except PermissionError:
            return ToolResult(success=False, output="", error=f"Permission denied: {path}")
        except Exception as e:
            return ToolResult(success=False, output="", error=f"Error reading file: {e}")


class ListDirTool(BaseTool):
    """List directory contents."""

    name = "list_dir"
    description = """List contents of a directory.

Returns a structured view of files and subdirectories with:
- File sizes (human-readable)
- File types (file/directory)
- Last modified time

## Parameters
- path: Directory path (default: current directory)
- pattern: Optional glob pattern to filter (e.g., "*.py")
- recursive: If true, list recursively (default: false)
- max_depth: Maximum depth for recursive listing (default: 3)

## Examples
- List current dir: list_dir(".")
- List with pattern: list_dir("src", pattern="*.py")
- List recursively: list_dir(".", recursive=true, max_depth=2)"""

    parameters: dict[str, Any] = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Directory path to list", "default": "."},
            "pattern": {
                "type": "string",
                "description": "Glob pattern to filter files (e.g., '*.py')",
            },
            "recursive": {"type": "boolean", "description": "List recursively", "default": False},
            "max_depth": {
                "type": "integer",
                "description": "Maximum depth for recursive listing",
                "default": 3,
            },
        },
        "required": [],
    }

    requires_permission = False
    permission_level = "auto"

    def _format_size(self, size: int) -> str:
        """Format file size in human-readable form."""
        size_f = float(size)
        for unit in ["B", "KB", "MB", "GB"]:
            if size_f < 1024:
                return f"{size_f:>7.1f} {unit}" if unit != "B" else f"{int(size_f):>7} {unit}"
            size_f /= 1024
        return f"{size_f:>7.1f} TB"

    def _list_entries(
        self,
        dir_path: Path,
        pattern: str | None,
        recursive: bool,
        max_depth: int,
        current_depth: int = 0,
        prefix: str = "",
    ) -> list[str]:
        """List directory entries."""
        entries = []

        try:
            items = sorted(dir_path.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
        except PermissionError:
            return [f"{prefix}[Permission denied]"]

        for item in items:
            # Skip hidden files unless pattern explicitly includes them
            if item.name.startswith(".") and not (pattern and pattern.startswith(".")):
                continue

            if item.is_dir():
                # Only show/recurse into directories if recursive mode is on
                if recursive and current_depth < max_depth:
                    # Check if pattern matches directory name
                    if pattern and fnmatch.fnmatch(item.name, pattern):
                        entries.append(f"{prefix}{item.name}/")
                    # Always recurse (to find matching files inside)
                    sub_entries = self._list_entries(
                        item, pattern, recursive, max_depth, current_depth + 1, prefix + "  "
                    )
                    # Only add directory header if it has matching contents
                    if sub_entries and not (pattern and fnmatch.fnmatch(item.name, pattern)):
                        entries.append(f"{prefix}{item.name}/")
                    entries.extend(sub_entries)
                elif not pattern:
                    # No pattern, non-recursive: just show directory name
                    entries.append(f"{prefix}{item.name}/")
            else:
                # Apply pattern filter to files
                if pattern and not fnmatch.fnmatch(item.name, pattern):
                    continue
                try:
                    stat = item.stat()
                    size = self._format_size(stat.st_size)
                    entries.append(f"{prefix}{item.name}  ({size})")
                except (OSError, PermissionError):
                    entries.append(f"{prefix}{item.name}  (unknown size)")

        return entries

    async def execute(
        self,
        path: str = ".",
        pattern: str | None = None,
        recursive: bool = False,
        max_depth: int = 3,
        **kwargs: Any,
    ) -> ToolResult:
        """List directory contents."""
        try:
            dir_path = Path(path).expanduser()

            if not dir_path.exists():
                return ToolResult(success=False, output="", error=f"Directory not found: {path}")

            if not dir_path.is_dir():
                return ToolResult(success=False, output="", error=f"Not a directory: {path}")

            entries = self._list_entries(dir_path, pattern, recursive, max_depth)

            if not entries:
                output = f"[{path}] Empty directory"
                if pattern:
                    output += f" (or no files matching '{pattern}')"
            else:
                header = f"[{path}]"
                if pattern:
                    header += f" (filtered: {pattern})"
                output = header + "\n\n" + "\n".join(entries)

            return ToolResult(success=True, output=output)

        except PermissionError:
            return ToolResult(success=False, output="", error=f"Permission denied: {path}")
        except Exception as e:
            return ToolResult(success=False, output="", error=f"Error listing directory: {e}")


class SearchFilesTool(BaseTool):
    """Search for patterns in files."""

    name = "search_files"
    description = """Search for a pattern in files (like grep).

Returns matching lines with file paths and line numbers.

## Parameters
- pattern: Text or regex pattern to search for
- path: Directory or file to search in (default: current directory)
- file_pattern: Glob pattern for files to search (e.g., "*.py")
- is_regex: Treat pattern as regex (default: false)
- ignore_case: Case-insensitive search (default: true)
- max_results: Maximum number of results (default: 100)

## Examples
- Search text: search_files("TODO", path="src")
- Search Python files: search_files("import os", file_pattern="*.py")
- Regex search: search_files("def\\s+\\w+\\(", is_regex=true)"""

    parameters: dict[str, Any] = {
        "type": "object",
        "properties": {
            "pattern": {"type": "string", "description": "Text or regex pattern to search for"},
            "path": {
                "type": "string",
                "description": "Directory or file to search in",
                "default": ".",
            },
            "file_pattern": {
                "type": "string",
                "description": "Glob pattern for files to search (e.g., '*.py')",
            },
            "is_regex": {
                "type": "boolean",
                "description": "Treat pattern as regex",
                "default": False,
            },
            "ignore_case": {
                "type": "boolean",
                "description": "Case-insensitive search",
                "default": True,
            },
            "max_results": {
                "type": "integer",
                "description": "Maximum number of results",
                "default": 100,
            },
        },
        "required": ["pattern"],
    }

    requires_permission = False
    permission_level = "auto"

    def _search_file(
        self,
        file_path: Path,
        pattern: re.Pattern,
        max_results: int,
        results: list[str],
    ) -> int:
        """Search a single file, return number of matches found."""
        try:
            # Skip binary files
            raw_data = file_path.read_bytes()[:8192]
            if is_binary(raw_data):
                return 0

            content = file_path.read_text(encoding="utf-8", errors="replace")
            lines = content.splitlines()

            matches_in_file = 0
            for i, line in enumerate(lines, 1):
                if len(results) >= max_results:
                    return matches_in_file

                if pattern.search(line):
                    # Truncate very long lines
                    display_line = line[:200] + "..." if len(line) > 200 else line
                    results.append(f"{file_path}:{i}: {display_line}")
                    matches_in_file += 1

            return matches_in_file

        except (PermissionError, OSError, UnicodeDecodeError):
            return 0

    async def execute(
        self,
        pattern: str,
        path: str = ".",
        file_pattern: str | None = None,
        is_regex: bool = False,
        ignore_case: bool = True,
        max_results: int = 100,
        **kwargs: Any,
    ) -> ToolResult:
        """Search for pattern in files."""
        try:
            search_path = Path(path).expanduser()

            if not search_path.exists():
                return ToolResult(success=False, output="", error=f"Path not found: {path}")

            # Compile pattern
            flags = re.IGNORECASE if ignore_case else 0
            try:
                if is_regex:
                    regex = re.compile(pattern, flags)
                else:
                    regex = re.compile(re.escape(pattern), flags)
            except re.error as e:
                return ToolResult(success=False, output="", error=f"Invalid regex pattern: {e}")

            results: list[str] = []

            if search_path.is_file():
                self._search_file(search_path, regex, max_results, results)
            else:
                # Walk directory
                for root, dirs, files in os.walk(search_path):
                    # Skip hidden directories
                    dirs[:] = [d for d in dirs if not d.startswith(".")]

                    for filename in files:
                        if filename.startswith("."):
                            continue

                        if file_pattern and not fnmatch.fnmatch(filename, file_pattern):
                            continue

                        file_path = Path(root) / filename
                        self._search_file(file_path, regex, max_results, results)

                        if len(results) >= max_results:
                            break

                    if len(results) >= max_results:
                        break

            if not results:
                output = f"No matches found for '{pattern}'"
                if file_pattern:
                    output += f" in {file_pattern} files"
            else:
                header = f"Found {len(results)} matches"
                if len(results) >= max_results:
                    header += f" (truncated at {max_results})"
                output = header + "\n\n" + "\n".join(results)

            return ToolResult(success=True, output=output)

        except Exception as e:
            return ToolResult(success=False, output="", error=f"Error searching: {e}")


class WriteFileTool(BaseTool):
    """Write content to a file."""

    name = "write_file"
    description = """Write content to a file.

Creates the file if it doesn't exist, or overwrites if it does.
Creates parent directories as needed.

## Parameters
- path: File path to write to
- content: Content to write

## Examples
- Write new file: write_file("config.json", '{"key": "value"}')
- Overwrite file: write_file("README.md", "# New content")"""

    parameters: dict[str, Any] = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to the file to write"},
            "content": {"type": "string", "description": "Content to write to the file"},
        },
        "required": ["path", "content"],
    }

    requires_permission = True  # Writing requires confirmation
    permission_level = "always_ask"

    async def execute(
        self,
        path: str,
        content: str,
        **kwargs: Any,
    ) -> ToolResult:
        """Write content to file."""
        try:
            file_path = Path(path).expanduser()

            # Create parent directories if needed
            file_path.parent.mkdir(parents=True, exist_ok=True)

            # Write the file
            file_path.write_text(content, encoding="utf-8")

            lines = len(content.splitlines())
            size = len(content.encode("utf-8"))

            return ToolResult(
                success=True, output=f"Wrote {lines} lines ({size:,} bytes) to {path}"
            )

        except PermissionError:
            return ToolResult(success=False, output="", error=f"Permission denied: {path}")
        except Exception as e:
            return ToolResult(success=False, output="", error=f"Error writing file: {e}")


class EditDocxTool(BaseTool):
    """Edit Word documents (.docx)."""

    name = "edit_docx"
    description = """Edit a Word document (.docx file).

Supports multiple edit operations:
- find_replace: Find and replace text throughout the document
- insert_after: Insert a new paragraph after a paragraph containing specific text
- update_paragraph: Update the entire text of a paragraph containing specific text
- save_as: Save to a new file (keeps original unchanged)

## Parameters
- path: Path to the .docx file
- operations: List of edit operations to perform
- save_as: Optional new path to save to (if not provided, overwrites original)

## Operations format
Each operation is an object with:
- type: "find_replace" | "insert_after" | "update_paragraph"
- find: Text to search for (partial match)
- replace/text: New text content

## Examples
edit_docx("resume.docx", [
    {"type": "find_replace", "find": "old title", "replace": "new title"},
    {"type": "insert_after", "find": "Work Experience", "text": "New bullet point"},
    {"type": "update_paragraph", "find": "Microsoft", "text": "Microsoft AI, Principal EM"}
], save_as="resume_updated.docx")"""

    parameters: dict[str, Any] = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to the .docx file to edit"},
            "operations": {
                "type": "array",
                "description": "List of edit operations",
                "items": {
                    "type": "object",
                    "properties": {
                        "type": {
                            "type": "string",
                            "enum": ["find_replace", "insert_after", "update_paragraph"],
                            "description": "Type of edit operation",
                        },
                        "find": {
                            "type": "string",
                            "description": "Text to search for (partial match)",
                        },
                        "replace": {
                            "type": "string",
                            "description": "Replacement text (for find_replace)",
                        },
                        "text": {
                            "type": "string",
                            "description": "New text content (for insert_after, update_paragraph)",
                        },
                    },
                    "required": ["type", "find"],
                },
            },
            "save_as": {
                "type": "string",
                "description": "Optional: Save to a new file path instead of overwriting",
            },
        },
        "required": ["path", "operations"],
    }

    requires_permission = True
    permission_level = "always_ask"

    async def execute(
        self,
        path: str,
        operations: list[dict[str, str]],
        save_as: str | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        """Edit a Word document."""
        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                success=False,
                output="",
                error="python-docx not installed. Run: pip install python-docx",
            )

        try:
            file_path = Path(path).expanduser()

            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            if file_path.suffix.lower() != ".docx":
                return ToolResult(success=False, output="", error=f"Not a .docx file: {path}")

            doc = Document(str(file_path))
            changes = []

            for op in operations:
                op_type = op.get("type")
                find_text = op.get("find", "")

                if op_type == "find_replace":
                    replace_text = op.get("replace", "")
                    count = 0
                    for para in doc.paragraphs:
                        if find_text in para.text:
                            # Preserve runs structure for simple replacements
                            for run in para.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                                    count += 1
                    if count > 0:
                        changes.append(
                            f"Replaced '{find_text}' → '{replace_text}' ({count} occurrences)"
                        )
                    else:
                        changes.append(f"⚠️ '{find_text}' not found for replacement")

                elif op_type == "update_paragraph":
                    new_text = op.get("text", "")
                    found = False
                    for para in doc.paragraphs:
                        if find_text in para.text:
                            old_text = para.text
                            para.text = new_text
                            changes.append(
                                f"Updated paragraph: '{old_text[:40]}...' → '{new_text[:40]}...'"
                            )
                            found = True
                            break
                    if not found:
                        changes.append(f"⚠️ Paragraph containing '{find_text}' not found")

                elif op_type == "insert_after":
                    new_text = op.get("text", "")
                    found = False
                    for i, para in enumerate(doc.paragraphs):
                        if find_text in para.text:
                            # Insert after this paragraph
                            new_para = doc.paragraphs[i]._element
                            new_p = doc.add_paragraph(new_text)._element
                            new_para.addnext(new_p)
                            changes.append(
                                f"Inserted after '{find_text[:30]}...': '{new_text[:40]}...'"
                            )
                            found = True
                            break
                    if not found:
                        changes.append(f"⚠️ Paragraph '{find_text}' not found for insert")

                else:
                    changes.append(f"⚠️ Unknown operation type: {op_type}")

            # Save
            output_path = Path(save_as).expanduser() if save_as else file_path
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            summary = "\n".join(f"  • {c}" for c in changes)
            return ToolResult(
                success=True, output=f"Edited Word document:\n{summary}\n\nSaved to: {output_path}"
            )

        except Exception as e:
            return ToolResult(success=False, output="", error=f"Error editing document: {e}")
