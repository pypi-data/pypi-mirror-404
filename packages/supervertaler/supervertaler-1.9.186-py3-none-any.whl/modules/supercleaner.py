"""
Supercleaner Module for Supervertaler
======================================

Cleans up DOCX documents before translation by removing formatting issues,
excessive tags, and OCR artifacts. Combines functionality similar to:
- TransTools Document Cleaner (tag/formatting cleanup)
- TransTools Unbreaker (incorrect line break removal)

Author: Michael Beijer / Supervertaler
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import re
from typing import List, Dict, Any
import logging


class DocumentCleaner:
    """
    Clean DOCX documents by removing formatting issues and excessive tags.
    Also includes Unbreaker functionality to fix incorrect line/paragraph breaks.
    """

    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.operations_performed = []
        self.incorrect_breaks_found = []

    def clean_document(self, input_path: str, output_path: str, operations: Dict[str, bool]) -> Dict[str, Any]:
        """
        Clean a DOCX document based on selected operations

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save cleaned DOCX file
            operations: Dictionary of operation names and whether to perform them

        Returns:
            Dictionary with statistics about operations performed
        """
        try:
            doc = Document(input_path)
            stats = {
                'paragraphs_processed': 0,
                'runs_processed': 0,
                'changes_made': 0,
                'operations': []
            }

            # Process all paragraphs
            for paragraph in doc.paragraphs:
                stats['paragraphs_processed'] += 1

                # Process all runs in paragraph
                for run in paragraph.runs:
                    stats['runs_processed'] += 1

                    # Perform selected operations
                    if operations.get('remove_text_shading', False):
                        if self._remove_text_shading(run):
                            stats['changes_made'] += 1

                    if operations.get('remove_highlighting', False):
                        if self._remove_highlighting(run):
                            stats['changes_made'] += 1

                    if operations.get('font_color_to_automatic', False):
                        if self._set_font_color_automatic(run):
                            stats['changes_made'] += 1

                    if operations.get('remove_character_styles', False):
                        if self._remove_character_styles(run):
                            stats['changes_made'] += 1

                # Paragraph-level operations
                if operations.get('normalize_font_color', False):
                    if self._normalize_paragraph_font_color(paragraph):
                        stats['changes_made'] += 1

                if operations.get('normalize_font_size', False):
                    if self._normalize_paragraph_font_size(paragraph):
                        stats['changes_made'] += 1

                if operations.get('normalize_font', False):
                    if self._normalize_paragraph_font(paragraph):
                        stats['changes_made'] += 1

                if operations.get('set_default_spacing', False):
                    if self._set_default_spacing(paragraph):
                        stats['changes_made'] += 1

            # Text content operations
            if operations.get('remove_manual_hyphens', False):
                count = self._remove_manual_hyphens(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Removed {count} manual hyphens")

            if operations.get('replace_special_symbols', False):
                count = self._replace_special_symbols(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Replaced {count} special symbols")

            if operations.get('simplify_quotes_and_dashes', False):
                count = self._simplify_quotes_and_dashes(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Simplified {count} quotes/dashes to ASCII")

            # Unbreaker operations - fix incorrect line/paragraph breaks
            if operations.get('fix_line_breaks', False):
                count = self._fix_incorrect_line_breaks(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Fixed {count} incorrect line breaks")

            if operations.get('join_broken_sentences', False):
                count = self._join_broken_sentences(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Joined {count} broken sentences")

            # Remove excessive spaces
            if operations.get('remove_excessive_spaces', False):
                count = self._remove_excessive_spaces(doc)
                stats['changes_made'] += count
                if count > 0:
                    stats['operations'].append(f"Cleaned up {count} runs with excessive spaces")

            if operations.get('accept_tracked_changes', False):
                # Note: python-docx doesn't fully support tracked changes
                # This would require a more complex implementation
                stats['operations'].append("Tracked changes acceptance not yet implemented")

            # Save cleaned document
            doc.save(output_path)

            return stats

        except Exception as e:
            self.logger.error(f"Error cleaning document: {e}")
            raise

    def _remove_text_shading(self, run) -> bool:
        """Remove background shading from text run"""
        try:
            if run.font.highlight_color is not None or hasattr(run._element, 'shd'):
                # Remove shading from the run element
                shd = run._element.get_or_add_rPr().find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    shd.getparent().remove(shd)
                    return True
        except Exception:
            pass
        return False

    def _remove_highlighting(self, run) -> bool:
        """Remove text highlighting"""
        try:
            if run.font.highlight_color is not None:
                run.font.highlight_color = None
                return True
        except Exception:
            pass
        return False

    def _set_font_color_automatic(self, run) -> bool:
        """Change font color from explicit colors to automatic"""
        try:
            if run.font.color is not None and run.font.color.rgb is not None:
                # Set to automatic (None)
                run.font.color.rgb = None
                return True
        except Exception:
            pass
        return False

    def _remove_character_styles(self, run) -> bool:
        """Remove character styles, keeping only direct formatting"""
        try:
            if run.style is not None and run.style.name != 'Default Paragraph Font':
                run.style = None
                return True
        except Exception:
            pass
        return False

    def _normalize_paragraph_font_color(self, paragraph) -> bool:
        """Normalize font color across all runs in paragraph to the most common color"""
        try:
            if not paragraph.runs:
                return False

            # Find most common color
            colors = {}
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                    colors[color] = colors.get(color, 0) + 1

            if not colors:
                return False

            # Get most common color
            most_common = max(colors, key=colors.get)

            # Apply to all runs
            changed = False
            for run in paragraph.runs:
                if run.font.color is None or run.font.color.rgb != most_common:
                    run.font.color.rgb = most_common
                    changed = True

            return changed
        except Exception:
            pass
        return False

    def _normalize_paragraph_font_size(self, paragraph) -> bool:
        """Normalize font size across all runs in paragraph to the most common size"""
        try:
            if not paragraph.runs:
                return False

            # Find most common size
            sizes = {}
            for run in paragraph.runs:
                if run.font.size:
                    size = run.font.size
                    sizes[size] = sizes.get(size, 0) + 1

            if not sizes:
                return False

            # Get most common size
            most_common = max(sizes, key=sizes.get)

            # Apply to all runs
            changed = False
            for run in paragraph.runs:
                if run.font.size != most_common:
                    run.font.size = most_common
                    changed = True

            return changed
        except Exception:
            pass
        return False

    def _normalize_paragraph_font(self, paragraph) -> bool:
        """Normalize font name across all runs in paragraph to the most common font"""
        try:
            if not paragraph.runs:
                return False

            # Find most common font
            fonts = {}
            for run in paragraph.runs:
                if run.font.name:
                    font = run.font.name
                    fonts[font] = fonts.get(font, 0) + 1

            if not fonts:
                return False

            # Get most common font
            most_common = max(fonts, key=fonts.get)

            # Apply to all runs
            changed = False
            for run in paragraph.runs:
                if run.font.name != most_common:
                    run.font.name = most_common
                    changed = True

            return changed
        except Exception:
            pass
        return False

    def _set_default_spacing(self, paragraph) -> bool:
        """Set default paragraph spacing"""
        try:
            # Set line spacing to single (1.0)
            if paragraph.paragraph_format.line_spacing != 1.0:
                paragraph.paragraph_format.line_spacing = 1.0
                return True
        except Exception:
            pass
        return False

    def _remove_manual_hyphens(self, doc) -> int:
        """Remove manual/soft hyphens from document"""
        count = 0
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if '\u00AD' in run.text or '\u002D' in run.text:  # Soft hyphen and regular hyphen
                        original = run.text
                        # Remove soft hyphens
                        run.text = run.text.replace('\u00AD', '')
                        # Remove hyphens at end of lines (manual hyphenation)
                        run.text = re.sub(r'-\s+', '', run.text)
                        if run.text != original:
                            count += 1
        except Exception:
            pass
        return count

    def _replace_special_symbols(self, doc) -> int:
        """Replace problematic special symbols (mainly non-breaking spaces and ellipsis)"""
        count = 0
        replacements = {
            '\u2026': '...', # Ellipsis
            '\u00A0': ' ',  # Non-breaking space (important for TM matching)
        }

        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original = run.text
                    for special, regular in replacements.items():
                        run.text = run.text.replace(special, regular)
                    if run.text != original:
                        count += 1
        except Exception:
            pass
        return count

    def _simplify_quotes_and_dashes(self, doc) -> int:
        """Convert typographic quotes and dashes to simple ASCII equivalents (OPTIONAL)"""
        count = 0
        replacements = {
            '\u2018': "'",  # Left single quotation mark → straight apostrophe
            '\u2019': "'",  # Right single quotation mark → straight apostrophe
            '\u201C': '"',  # Left double quotation mark → straight quote
            '\u201D': '"',  # Right double quotation mark → straight quote
            '\u2013': '-',  # En dash → hyphen
            '\u2014': '-',  # Em dash → hyphen (NOT double hyphen)
        }

        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original = run.text
                    for special, regular in replacements.items():
                        run.text = run.text.replace(special, regular)
                    if run.text != original:
                        count += 1
        except Exception:
            pass
        return count

    # ============================================================================
    # UNBREAKER FUNCTIONALITY - Fix incorrect line/paragraph breaks
    # ============================================================================

    def _fix_incorrect_line_breaks(self, doc) -> int:
        """
        Fix incorrect line breaks (manual line breaks within sentences).
        Detects line breaks that occur mid-sentence and removes them.
        """
        count = 0
        try:
            for paragraph in doc.paragraphs:
                # Check for line breaks within the paragraph text
                original_text = paragraph.text

                # Line break character in Word is '\v' or '\x0B'
                if '\v' in original_text or '\x0B' in original_text:
                    # Check if these are likely incorrect (mid-sentence)
                    if self._is_likely_incorrect_break(original_text):
                        # Remove line breaks and replace with space
                        new_text = original_text.replace('\v', ' ').replace('\x0B', ' ')
                        # Clean up multiple spaces
                        new_text = re.sub(r'\s+', ' ', new_text)

                        # Update paragraph text
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                            # Clear other runs
                            for i in range(len(paragraph.runs) - 1, 0, -1):
                                paragraph.runs[i].text = ''
                            count += 1
        except Exception as e:
            self.logger.error(f"Error fixing line breaks: {e}")
        return count

    def _join_broken_sentences(self, doc) -> int:
        """
        Join sentences that were incorrectly split across paragraphs.
        Detects paragraphs that don't end with sentence-ending punctuation
        and joins them with the next paragraph.
        
        DISABLED BY DEFAULT - This operation is too aggressive and causes
        words to stick together. Needs more sophisticated logic to detect
        true broken sentences vs intentional paragraph breaks.
        """
        count = 0
        # TEMPORARILY DISABLED due to word spacing bugs
        # The current logic joins too many paragraphs incorrectly
        return count
        
        # Original code kept for reference but not executed:
        # try:
        #     paragraphs = list(doc.paragraphs)
        #     i = 0
        #
        #     while i < len(paragraphs) - 1:
        #         current_para = paragraphs[i]
        #         next_para = paragraphs[i + 1]
        #
        #         current_text = current_para.text.strip()
        #         next_text = next_para.text.strip()
        #
        #         # Skip empty paragraphs
        #         if not current_text or not next_text:
        #             i += 1
        #             continue
        #
        #         # Check if current paragraph ends mid-sentence
        #         if self._is_broken_sentence(current_text):
        #             # Join paragraphs WITH PROPER SPACING
        #             joined_text = current_text + ' ' + next_text
        #
        #             # Update current paragraph
        #             if current_para.runs:
        #                 current_para.runs[0].text = joined_text
        #                 # Clear other runs
        #                 for j in range(len(current_para.runs) - 1, 0, -1):
        #                     current_para.runs[j].text = ''
        #
        #             # Clear next paragraph
        #             if next_para.runs:
        #                 for run in next_para.runs:
        #                     run.text = ''
        #
        #             count += 1
        #
        #         i += 1
        #
        # except Exception as e:
        #     self.logger.error(f"Error joining broken sentences: {e}")
        # return count

    def _is_likely_incorrect_break(self, text: str) -> bool:
        """Check if a line break is likely incorrect (mid-sentence)"""
        # Line breaks before lowercase letters are often incorrect
        if re.search(r'\v[a-z]', text) or re.search(r'\x0B[a-z]', text):
            return True
        # Line breaks not followed by capital letters or numbers
        if re.search(r'\v[^A-Z0-9\s]', text) or re.search(r'\x0B[^A-Z0-9\s]', text):
            return True
        return False

    def _is_broken_sentence(self, text: str) -> bool:
        """Check if text appears to be a broken sentence (doesn't end properly)"""
        # Ends with sentence-ending punctuation
        sentence_enders = ('.', '!', '?', ':', ';')

        # Skip if ends with sentence-ending punctuation
        if text.endswith(sentence_enders):
            return False

        # Likely broken if ends with lowercase letter
        if text and text[-1].islower():
            return True

        # Likely broken if ends with comma
        if text.endswith(','):
            return True

        # Likely broken if very short (less than 50 chars)
        if len(text) < 50:
            return True

        return False

    # ============================================================================
    # REMOVE EXCESSIVE SPACES FUNCTIONALITY
    # ============================================================================

    def _remove_excessive_spaces(self, doc) -> int:
        """
        Remove excessive spaces between words and around punctuation.
        
        CRITICAL: We work on full paragraph text, not individual runs,
        because runs are formatting boundaries and may split words.
        Removing trailing spaces from runs causes words to stick together!
        """
        count = 0
        try:
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                
                # Only process if there's text
                if not original_text or not original_text.strip():
                    continue

                # Work on the full paragraph text
                text = original_text
                
                # Replace multiple spaces (2+) with single space
                text = re.sub(r'  +', ' ', text)

                # Remove spaces before punctuation (but be careful with abbreviations)
                text = re.sub(r' +([,;:!?)])', r'\1', text)

                # Remove spaces after opening punctuation
                text = re.sub(r'([(]) +', r'\1', text)
                
                # Remove leading/trailing spaces from paragraph
                text = text.strip()

                # Only update if changed
                if text != original_text:
                    # Reconstruct paragraph with cleaned text
                    # Keep the first run and put all text there, clear others
                    if paragraph.runs:
                        paragraph.runs[0].text = text
                        # Clear remaining runs
                        for i in range(len(paragraph.runs) - 1, 0, -1):
                            paragraph.runs[i].text = ''
                        count += 1

        except Exception as e:
            self.logger.error(f"Error removing excessive spaces: {e}")
        return count


def clean_document_simple(input_path: str, output_path: str = None,
                         quick_clean: bool = True) -> Dict[str, Any]:
    """
    Convenience function for quick document cleaning with default settings

    Args:
        input_path: Path to input DOCX file
        output_path: Path to save cleaned file (if None, overwrites input)
        quick_clean: If True, applies common cleaning operations

    Returns:
        Statistics dictionary
    """
    if output_path is None:
        output_path = input_path

    cleaner = DocumentCleaner()

    # Default quick clean operations (most useful for OCR/PDF cleanup)
    operations = {
        # Document Cleaner operations
        'remove_text_shading': quick_clean,
        'remove_highlighting': quick_clean,
        'font_color_to_automatic': quick_clean,
        'normalize_font_color': quick_clean,
        'normalize_font_size': quick_clean,
        'normalize_font': quick_clean,
        'set_default_spacing': quick_clean,
        'remove_manual_hyphens': quick_clean,
        'replace_special_symbols': quick_clean,  # Only non-breaking spaces and ellipsis
        'simplify_quotes_and_dashes': False,  # OPTIONAL - converts curly quotes/em-dashes to ASCII
        'remove_character_styles': False,  # More aggressive, optional

        # Unbreaker operations
        'fix_line_breaks': quick_clean,
        'join_broken_sentences': False,  # DISABLED - too aggressive, causes word spacing issues

        # Remove excessive spaces
        'remove_excessive_spaces': quick_clean,

        # Not yet implemented
        'accept_tracked_changes': False,
    }

    return cleaner.clean_document(input_path, output_path, operations)


if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_cleaner.py input.docx [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '_cleaned.docx')

    print(f"Cleaning document: {input_file}")
    stats = clean_document_simple(input_file, output_file)

    print(f"\nCleaning complete!")
    print(f"  Paragraphs processed: {stats['paragraphs_processed']}")
    print(f"  Runs processed: {stats['runs_processed']}")
    print(f"  Changes made: {stats['changes_made']}")
    print(f"  Output saved to: {output_file}")
