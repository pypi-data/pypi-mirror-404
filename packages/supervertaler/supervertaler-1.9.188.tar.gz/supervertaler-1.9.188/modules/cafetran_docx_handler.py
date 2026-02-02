"""
CafeTran Bilingual DOCX Handler

This module handles the import and export of CafeTran bilingual DOCX files.
CafeTran uses a simple table-based format with pipe symbols (|) to mark formatted text.

Format Structure:
- Table with columns: ID | Source | Target | Notes | *
- Pipe symbols (|) surround formatted text in the source column
- Examples:
  - |Atalanta| = underlined text
  - Biagio Pagano| = bold text (pipe at end)
  - |text| = formatted text (underlined)
  
The pipe symbols are preserved during translation and applied to the target text.
"""

import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_UNDERLINE
import re


class FormattedSegment:
    """
    Represents a text segment with formatting information using pipe symbols.
    """
    def __init__(self, segment_id, source_with_pipes, target_with_pipes="", notes=""):
        self.segment_id = segment_id
        self.source_with_pipes = source_with_pipes  # Source text with | markers
        self.target_with_pipes = target_with_pipes  # Target text with | markers
        self.notes = notes
        
    @property
    def plain_text(self):
        """Get source text with pipe symbols removed for translation."""
        return self.source_with_pipes.replace('|', '')
    
    def __repr__(self):
        return f"FormattedSegment(id={self.segment_id}, source='{self.source_with_pipes[:50]}...', target='{self.target_with_pipes[:50]}...')"


class CafeTranDOCXHandler:
    """
    Handler for CafeTran bilingual DOCX files.
    
    This class provides methods to:
    - Load and parse CafeTran bilingual DOCX files
    - Extract source segments with formatting markers (pipe symbols)
    - Update target segments with translations
    - Save modified files while preserving formatting
    """
    
    def __init__(self):
        self.doc = None
        self.table = None
        self.segments = []
        self.file_path = None
        self.header_row = None
        
    def load(self, file_path):
        """
        Load a CafeTran bilingual DOCX file.
        
        Args:
            file_path: Path to the CafeTran bilingual DOCX file
            
        Returns:
            bool: True if loaded successfully, False otherwise
        """
        try:
            self.file_path = file_path
            self.doc = Document(file_path)
            
            # CafeTran bilingual files should have exactly one table
            if len(self.doc.tables) == 0:
                print(f"ERROR: No table found in {file_path}")
                return False
                
            self.table = self.doc.tables[0]
            
            # Verify the header row (first row should be: ID, filename, filename, Notes, *)
            if len(self.table.rows) < 2:
                print(f"ERROR: Table has insufficient rows")
                return False
                
            self.header_row = [cell.text.strip() for cell in self.table.rows[0].cells]
            
            # Check if this looks like a CafeTran bilingual DOCX
            if self.header_row[0] != 'ID':
                print(f"ERROR: First column header should be 'ID', got '{self.header_row[0]}'")
                return False
                
            print(f"Successfully loaded CafeTran bilingual DOCX: {file_path}")
            print(f"Header: {self.header_row}")
            print(f"Total rows (including header): {len(self.table.rows)}")
            
            return True
            
        except Exception as e:
            print(f"ERROR loading CafeTran DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def extract_source_segments(self):
        """
        Extract all source segments from the CafeTran bilingual DOCX.
        
        Returns:
            list: List of FormattedSegment objects with pipe symbols preserved
        """
        self.segments = []
        
        if not self.table:
            print("ERROR: No table loaded")
            return []
        
        # Skip header row (index 0), process data rows
        for i, row in enumerate(self.table.rows[1:], start=1):
            try:
                cells = row.cells
                
                # Extract data from columns
                segment_id = cells[0].text.strip()
                source = cells[1].text.strip()
                target = cells[2].text.strip() if len(cells) > 2 else ""
                notes = cells[3].text.strip() if len(cells) > 3 else ""
                
                # Create FormattedSegment with pipe symbols preserved
                segment = FormattedSegment(
                    segment_id=segment_id,
                    source_with_pipes=source,
                    target_with_pipes=target,
                    notes=notes
                )
                
                self.segments.append(segment)
                
            except Exception as e:
                print(f"WARNING: Error processing row {i}: {e}")
                continue
        
        print(f"Extracted {len(self.segments)} segments from CafeTran DOCX")
        return self.segments
    
    def update_target_segments(self, translations):
        """
        Update target segments with translations.
        
        This method takes plain translations and applies the pipe symbol formatting
        from the source segments to create properly formatted target segments.
        
        Args:
            translations: List of translated strings (without pipe symbols)
            
        Returns:
            bool: True if update successful, False otherwise
        """
        if not self.segments:
            print("ERROR: No segments loaded. Call extract_source_segments() first.")
            return False
            
        if len(translations) != len(self.segments):
            print(f"ERROR: Translation count ({len(translations)}) doesn't match segment count ({len(self.segments)})")
            return False
        
        print(f"Updating {len(translations)} target segments with pipe formatting...")
        
        for i, (segment, translation) in enumerate(zip(self.segments, translations)):
            try:
                # Apply pipe symbol formatting from source to translation
                formatted_translation = self._apply_pipe_formatting(
                    source_with_pipes=segment.source_with_pipes,
                    translation_plain=translation
                )
                
                segment.target_with_pipes = formatted_translation
                
            except Exception as e:
                print(f"WARNING: Error updating segment {i} (ID: {segment.segment_id}): {e}")
                # Fallback: use plain translation without formatting
                segment.target_with_pipes = translation
                
        print("Target segments updated successfully")
        return True
    
    def _apply_pipe_formatting(self, source_with_pipes, translation_plain):
        """
        Apply pipe symbol formatting from source to translation.
        
        This method is a placeholder that will be called by the main application.
        The actual formatting will be done by asking the AI to intelligently place
        pipe symbols in the translation based on their positions in the source.
        
        For now, this just returns the plain translation. The AI-based formatting
        will be handled in the main application during the translation process.
        
        Args:
            source_with_pipes: Source text with pipe symbols
            translation_plain: Translated text without pipe symbols
            
        Returns:
            str: Translation (will have pipes added by AI in the main app)
        """
        # Return plain translation for now
        # The main app will handle AI-based pipe placement during translation
        return translation_plain
    
    def save(self, output_path=None):
        """
        Save the CafeTran bilingual DOCX with updated target segments.
        
        Args:
            output_path: Optional path for output file. If None, overwrites original.
            
        Returns:
            bool: True if saved successfully, False otherwise
        """
        if not self.doc or not self.table:
            print("ERROR: No document loaded")
            return False
            
        if not self.segments:
            print("ERROR: No segments to save")
            return False
        
        try:
            # Update the table cells with translated content
            for i, segment in enumerate(self.segments):
                row_idx = i + 1  # +1 because row 0 is header
                
                if row_idx >= len(self.table.rows):
                    print(f"WARNING: Row index {row_idx} out of range, skipping segment {segment.segment_id}")
                    continue
                
                # Update target cell (column 2) with formatted pipe symbols
                target_cell = self.table.rows[row_idx].cells[2]
                
                # Clear existing content
                target_cell.text = ''
                
                # Add content with formatted pipe symbols (bold + red)
                self._add_text_with_formatted_pipes(target_cell, segment.target_with_pipes)
            
            # Save the document
            save_path = output_path if output_path else self.file_path
            self.doc.save(save_path)
            
            print(f"Successfully saved CafeTran bilingual DOCX to: {save_path}")
            return True
            
        except Exception as e:
            print(f"ERROR saving CafeTran DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_text_with_formatted_pipes(self, cell, text_with_pipes):
        """
        Add text to a cell with pipe symbols formatted as bold and red.
        
        Args:
            cell: The table cell to add text to
            text_with_pipes: Text containing pipe symbols
        """
        # Split text by pipe symbols
        parts = text_with_pipes.split('|')
        
        # Add a paragraph to the cell
        if len(cell.paragraphs) == 0:
            paragraph = cell.add_paragraph()
        else:
            paragraph = cell.paragraphs[0]
            paragraph.clear()
        
        # Add text parts with pipes formatted
        for i, part in enumerate(parts):
            if i > 0:
                # Add the pipe symbol with bold + red formatting
                pipe_run = paragraph.add_run('|')
                pipe_run.bold = True
                pipe_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            
            if part:  # Only add non-empty parts
                # Add the regular text
                text_run = paragraph.add_run(part)
    
    @staticmethod
    def is_cafetran_bilingual_docx(file_path):
        """
        Check if a DOCX file is a CafeTran bilingual DOCX.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            bool: True if file appears to be CafeTran bilingual DOCX, False otherwise
        """
        try:
            doc = Document(file_path)
            
            if len(doc.tables) == 0:
                return False
            
            table = doc.tables[0]
            
            if len(table.rows) < 2:
                return False
            
            # Check header row
            header = [cell.text.strip() for cell in table.rows[0].cells]
            
            # CafeTran bilingual DOCX should have:
            # - First column: "ID"
            # - At least 4-5 columns
            if len(header) >= 4 and header[0] == 'ID':
                return True
            
            return False
            
        except Exception as e:
            print(f"Error checking if file is CafeTran bilingual DOCX: {e}")
            return False


# Test function for standalone execution
def test_handler():
    """Test the CafeTran DOCX handler with a sample file."""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python cafetran_docx_handler.py <cafetran_bilingual.docx>")
        return
    
    file_path = sys.argv[1]
    
    print(f"\n{'='*60}")
    print(f"Testing CafeTran DOCX Handler")
    print(f"{'='*60}\n")
    
    # Test 1: Check if file is CafeTran bilingual DOCX
    print("Test 1: Checking file format...")
    is_cafetran = CafeTranDOCXHandler.is_cafetran_bilingual_docx(file_path)
    print(f"Is CafeTran bilingual DOCX: {is_cafetran}\n")
    
    if not is_cafetran:
        print("File does not appear to be a CafeTran bilingual DOCX")
        return
    
    # Test 2: Load file
    print("Test 2: Loading file...")
    handler = CafeTranDOCXHandler()
    if not handler.load(file_path):
        print("Failed to load file")
        return
    print("File loaded successfully\n")
    
    # Test 3: Extract segments
    print("Test 3: Extracting segments...")
    segments = handler.extract_source_segments()
    print(f"Extracted {len(segments)} segments\n")
    
    # Display first 5 segments
    print("First 5 segments:")
    for i, seg in enumerate(segments[:5]):
        print(f"\n  Segment {i+1} (ID: {seg.segment_id}):")
        print(f"    Source with pipes: {seg.source_with_pipes}")
        print(f"    Plain text: {seg.plain_text}")
        print(f"    Target: {seg.target_with_pipes if seg.target_with_pipes else '(empty)'}")
    
    print(f"\n{'='*60}")
    print("Testing complete!")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    test_handler()
