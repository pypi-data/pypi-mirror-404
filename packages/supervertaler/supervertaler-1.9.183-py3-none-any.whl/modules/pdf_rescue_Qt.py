"""
PDF Rescue Module - Qt Edition
Embeddable version of the AI-powered OCR tool for extracting text from poorly formatted PDFs
Supports multiple AI providers: OpenAI GPT-4 Vision, Anthropic Claude Vision, Google Gemini Vision

This module can be embedded in the main Supervertaler Qt application as a tab.
Can also be run independently as a standalone application.
"""

import os
import base64
from pathlib import Path
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QListWidget, QTextEdit,
    QComboBox, QCheckBox, QProgressBar, QFileDialog, QMessageBox, QSplitter,
    QGroupBox, QFrame, QDialog, QDialogButtonBox, QPlainTextEdit, QApplication,
    QStyleOptionButton
)
from PyQt6.QtCore import Qt, QTimer, QPointF, QRect
from PyQt6.QtGui import QFont, QTextOption, QPainter, QPen, QColor, QStandardItemModel, QStandardItem
from docx import Document
from docx.shared import Pt
import fitz  # PyMuPDF
import re


class CheckmarkCheckBox(QCheckBox):
    """Custom checkbox with green background and white checkmark when checked - matches Supervertaler Qt style"""
    
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.setStyleSheet("""
            QCheckBox {
                font-size: 9pt;
                spacing: 6px;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
                border: 2px solid #999;
                border-radius: 3px;
                background-color: white;
            }
            QCheckBox::indicator:checked {
                background-color: #4CAF50;
                border-color: #4CAF50;
            }
            QCheckBox::indicator:hover {
                border-color: #666;
            }
            QCheckBox::indicator:checked:hover {
                background-color: #45a049;
                border-color: #45a049;
            }
        """)
    
    def paintEvent(self, a0):
        """Override paint event to draw white checkmark when checked"""
        super().paintEvent(a0)
        
        if self.isChecked():
            # Get the indicator rectangle using QStyle
            opt = QStyleOptionButton()
            self.initStyleOption(opt)
            indicator_rect = self.style().subElementRect(
                self.style().SubElement.SE_CheckBoxIndicator,
                opt,
                self
            )
            
            if indicator_rect.isValid():
                # Draw white checkmark
                painter = QPainter(self)
                painter.setRenderHint(QPainter.RenderHint.Antialiasing)
                # Slightly thinner pen for better fit on smaller displays
                pen_width = max(2.0, min(indicator_rect.width(), indicator_rect.height()) * 0.12)
                painter.setPen(QPen(QColor(255, 255, 255), pen_width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
                painter.setBrush(QColor(255, 255, 255))
                
                # Draw checkmark (‚úì shape) - coordinates relative to indicator
                # Add padding to prevent clipping on smaller displays
                x = indicator_rect.x()
                y = indicator_rect.y()
                w = indicator_rect.width()
                h = indicator_rect.height()
                
                # Add padding (15% on all sides) to ensure checkmark doesn't get cut off on smaller displays
                padding = min(w, h) * 0.15
                x += padding
                y += padding
                w -= padding * 2
                h -= padding * 2
                
                # Checkmark path: bottom-left to middle, then middle to top-right
                # Using proportions that create a nice checkmark shape with proper padding
                check_x1 = x + w * 0.10  # Left point (more padding from left)
                check_y1 = y + h * 0.50
                check_x2 = x + w * 0.35  # Middle-bottom point
                check_y2 = y + h * 0.70  # Bottom point (with padding from bottom)
                check_x3 = x + w * 0.90  # Right point (more padding from right)
                check_y3 = y + h * 0.25  # Top point (with padding from top)
                
                # Draw two lines forming the checkmark
                painter.drawLine(QPointF(check_x2, check_y2), QPointF(check_x3, check_y3))
                painter.drawLine(QPointF(check_x1, check_y1), QPointF(check_x2, check_y2))
                
                painter.end()


class PDFRescueQt:
    """
    PDF Rescue feature - extract text from images using AI OCR
    Can be embedded in any PyQt6 application as a tab or panel
    """
    
    def __init__(self, parent_app, standalone=False):
        """
        Initialize PDF Rescue module
        
        Args:
            parent_app: Reference to the main application (needs .load_api_keys() method or .api_keys attribute)
            standalone: If True, running as standalone app. If False, embedded in Supervertaler
        """
        self.parent_app = parent_app
        self.standalone = standalone
        self.clients = {}  # Dictionary to store clients for different providers
        self.image_files = []
        self.extracted_texts = {}
        
        # Load API keys for all providers
        self.api_keys = {}
        if hasattr(parent_app, 'load_api_keys'):
            # Supervertaler_Qt style
            self.api_keys = parent_app.load_api_keys()
        elif hasattr(parent_app, 'api_keys'):
            # Direct api_keys dict
            self.api_keys = parent_app.api_keys
        
        # Initialize clients for available providers
        self._initialize_clients()
    
    def _initialize_clients(self):
        """Initialize API clients for all available providers"""
        # OpenAI
        openai_key = self.api_keys.get('openai') or self.api_keys.get('openai_api_key')
        if openai_key:
            try:
                from openai import OpenAI
                self.clients['openai'] = OpenAI(api_key=openai_key)
                self.log_message("‚úì OpenAI client initialized")
            except Exception as e:
                self.log_message(f"‚ö† Failed to initialize OpenAI: {e}")
        
        # Anthropic Claude
        claude_key = self.api_keys.get('claude') or self.api_keys.get('anthropic')
        if claude_key:
            try:
                import anthropic
                self.clients['claude'] = anthropic.Anthropic(api_key=claude_key)
                self.log_message("‚úì Claude client initialized")
            except ImportError:
                self.log_message("‚ö† Claude requested but 'anthropic' library not installed. Run: pip install anthropic")
            except Exception as e:
                self.log_message(f"‚ö† Failed to initialize Claude: {e}")
        
        # Google Gemini
        gemini_key = self.api_keys.get('gemini') or self.api_keys.get('google')
        if gemini_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=gemini_key)
                self.clients['gemini'] = genai
                self.log_message("‚úì Gemini client initialized")
            except ImportError:
                self.log_message("‚ö† Gemini requested but 'google-generativeai' library not installed. Run: pip install google-generativeai")
            except Exception as e:
                self.log_message(f"‚ö† Failed to initialize Gemini: {e}")
    
    def log_message(self, message: str):
        """Log a message to the parent app's log if available"""
        if hasattr(self.parent_app, 'log'):
            self.parent_app.log(f"[PDF Rescue] {message}")
        else:
            print(f"[PDF Rescue] {message}")
    
    def create_tab(self, parent):
        """
        Create the PDF Rescue tab UI
        
        Args:
            parent: The parent widget (QWidget)
        """
        # Save current state before recreating UI
        saved_files = self.image_files.copy() if hasattr(self, 'image_files') else []
        saved_texts = self.extracted_texts.copy() if hasattr(self, 'extracted_texts') else []
        
        # Main layout
        main_layout = QVBoxLayout(parent)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(5)  # Reduced from 10 to 5 for tighter spacing
        
        # Header (matches Universal Lookup / AutoFingers style)
        header = QLabel("üîç PDF Rescue - AI-Powered OCR")
        header.setStyleSheet("font-size: 16pt; font-weight: bold; color: #1976D2;")
        main_layout.addWidget(header, 0)  # 0 = no stretch, stays compact
        
        # Description box (matches Universal Lookup / AutoFingers style)
        description = QLabel(
            "Extract text from image-based PDFs using AI vision OCR. Designed for scanned documents, screenshots, "
            "and PDFs without accessible text. Not recommended for PDFs with selectable text - use professional tools like Adobe Acrobat instead."
        )
        description.setWordWrap(True)
        description.setStyleSheet("color: #666; padding: 5px; background-color: #E3F2FD; border-radius: 3px;")
        main_layout.addWidget(description, 0)  # 0 = no stretch, stays compact
        
        # Split view: Files on left, Preview on right
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # LEFT: File list
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(3, 0, 3, 3)  # Zero top margin to eliminate gap
        left_layout.setSpacing(3)
        
        files_label = QLabel("Images to Process")
        files_label.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        left_layout.addWidget(files_label)
        
        # File list
        self.file_listbox = QListWidget()
        self.file_listbox.setFont(QFont("Consolas", 9))
        self.file_listbox.itemSelectionChanged.connect(self._on_file_select)
        left_layout.addWidget(self.file_listbox)
        
        # Buttons
        btn_layout = QHBoxLayout()
        
        pdf_btn = QPushButton("üìÑ Import PDF")
        pdf_btn.clicked.connect(self._import_from_pdf)
        pdf_btn.setStyleSheet("background-color: #9C27B0; color: white; font-weight: bold; padding: 4px 8px;")
        pdf_btn.setToolTip("Extract all pages from a PDF file and convert them to images for OCR processing")
        btn_layout.addWidget(pdf_btn)
        
        add_files_btn = QPushButton("‚ûï Add Image Files")
        add_files_btn.clicked.connect(self._add_files)
        add_files_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 4px 8px;")
        add_files_btn.setToolTip("Supported formats: .jpg, .jpeg, .png, .bmp, .gif, .tiff")
        btn_layout.addWidget(add_files_btn)
        
        add_folder_btn = QPushButton("üìÇ Folder")
        add_folder_btn.clicked.connect(self._add_folder)
        add_folder_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 4px 8px;")
        add_folder_btn.setToolTip("Add all image files from a selected folder")
        btn_layout.addWidget(add_folder_btn)
        
        clear_btn = QPushButton("Clear")
        clear_btn.clicked.connect(self._clear_list)
        clear_btn.setStyleSheet("background-color: #9E9E9E; color: white; padding: 4px 8px;")
        clear_btn.setToolTip("Remove all files from the list")
        btn_layout.addWidget(clear_btn)
        
        left_layout.addLayout(btn_layout)
        
        # Processing options (moved into left panel to eliminate wasted space)
        options_group = QGroupBox("Processing Options")
        options_layout = QVBoxLayout(options_group)
        options_layout.setContentsMargins(8, 8, 8, 3)  # Minimal bottom margin
        options_layout.setSpacing(3)  # Reduced spacing
        
        # === MODEL SELECTOR ===
        # Model selection and formatting option
        model_layout = QHBoxLayout()
        
        model_label = QLabel("AI Model:")
        model_label.setFont(QFont("Segoe UI", 9))
        model_layout.addWidget(model_label)
        
        self.model_combo = QComboBox()
        # Organize models by provider with separators
        self.model_combo.addItem("--- OpenAI ---")
        self.model_combo.addItems(["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-4", "gpt-5"])
        self.model_combo.addItem("--- Claude (Anthropic) ---")
        self.model_combo.addItems(["claude-3-5-sonnet-20241022", "claude-3-5-haiku-20241022", "claude-3-opus-20240229"])
        self.model_combo.addItem("--- Gemini (Google) ---")
        self.model_combo.addItems(["gemini-2.0-flash-exp", "gemini-1.5-pro-002", "gemini-1.5-flash-002"])
        
        # Set default and style separator items
        self.model_combo.setCurrentText("gpt-4o")
        
        # Make separator items non-selectable by disabling them
        combo_model = self.model_combo.model()
        if isinstance(combo_model, QStandardItemModel):
            for i in range(self.model_combo.count()):
                if self.model_combo.itemText(i).startswith("---"):
                    item = combo_model.item(i)
                    if item:
                        item.setEnabled(False)
                        # Make separators visually distinct (gray, centered)
                        item.setFlags(Qt.ItemFlag.NoItemFlags)
        
        self.model_combo.setEditable(False)
        self.model_combo.setToolTip("Select AI model for vision OCR processing (OpenAI, Claude, or Gemini)")
        model_layout.addWidget(self.model_combo)
        
        model_layout.addSpacing(20)
        
        self.preserve_formatting_check = CheckmarkCheckBox("Preserve formatting (bold/italic/underline)")
        self.preserve_formatting_check.setChecked(True)
        self.preserve_formatting_check.setFont(QFont("Segoe UI", 9))
        self.preserve_formatting_check.setToolTip("When enabled, the AI will use markdown to preserve text formatting (bold, italic, underline)")
        model_layout.addWidget(self.preserve_formatting_check)
        
        model_layout.addStretch()
        options_layout.addLayout(model_layout)
        
        # Model descriptions (prominent display in left panel)
        model_desc_label = QLabel("Model Capabilities:")
        model_desc_label.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        options_layout.addWidget(model_desc_label)
        
        self.model_descriptions_text = QPlainTextEdit()
        self.model_descriptions_text.setFont(QFont("Segoe UI", 9))
        self.model_descriptions_text.setReadOnly(True)
        # Size naturally to content - no fixed max height to eliminate wasted space below
        model_descriptions = """‚Ä¢ gpt-4o (Recommended): Fast and accurate; best balance of speed, quality, and cost; excellent for most documents including tables
‚Ä¢ gpt-4o-mini: Fast and economical; good for simple documents; may struggle with complex layouts or tables
‚Ä¢ gpt-4-turbo: Large context window (128k tokens); good for very long documents; slightly slower but handles extensive content well
‚Ä¢ gpt-4: Classic, reliable baseline; consistent quality; good for standard documents, though slower than gpt-4o
‚Ä¢ gpt-5 (Advanced Reasoning): Reasoning model; may improve table extraction and complex layouts; slower and more expensive; best for: complex tables, technical documents, structured data extraction"""
        self.model_descriptions_text.setPlainText(model_descriptions)
        options_layout.addWidget(self.model_descriptions_text)
        
        left_layout.addWidget(options_group)
        
        # Add stretch to push everything up and eliminate wasted space below
        left_layout.addStretch()
        
        splitter.addWidget(left_widget)
        
        # RIGHT: Extraction Instructions and Text preview with vertical splitter
        right_splitter = QSplitter(Qt.Orientation.Vertical)
        
        # Top: Extraction Instructions
        instructions_widget = QWidget()
        instructions_widget_layout = QVBoxLayout(instructions_widget)
        instructions_widget_layout.setContentsMargins(0, 0, 0, 0)
        instructions_widget_layout.setSpacing(0)
        
        instructions_group = QGroupBox("Extraction Instructions")
        instructions_group_layout = QVBoxLayout(instructions_group)
        instructions_group_layout.setContentsMargins(8, 8, 8, 8)
        
        instructions_header = QHBoxLayout()
        instructions_header.setContentsMargins(0, 0, 0, 5)
        instructions_header.addStretch()
        
        show_prompt_btn = QPushButton("üëÅÔ∏è Show Prompt")
        show_prompt_btn.clicked.connect(self._show_full_prompt)
        show_prompt_btn.setStyleSheet("background-color: #9C27B0; color: white; padding: 2px 8px;")
        show_prompt_btn.setToolTip("Preview the exact prompt that will be sent to the AI model")
        instructions_header.addWidget(show_prompt_btn)
        
        instructions_group_layout.addLayout(instructions_header)
        
        self.instructions_text = QPlainTextEdit()
        self.instructions_text.setFont(QFont("Segoe UI", 9))
        self.instructions_text.setMinimumHeight(60)  # Reduced minimum height
        default_instructions = """Extract all text from this image. The image is a screenshot from a poorly formatted PDF.
Please:
- Extract all visible text accurately
- Fix any obvious OCR errors or formatting issues
- Remove extraneous line breaks within paragraphs
- Preserve intentional paragraph breaks
- Maintain the logical flow and structure of the content
- For redacted/blacked-out text: insert a descriptive placeholder in square brackets in the document's language (e.g., [naam] for Dutch names, [name] for English names, [bedrag] for amounts, etc.)
- For stamps, signatures, or images: insert a descriptive placeholder in square brackets in the document's language (e.g., [handtekening], [stempel], [signature], [stamp], etc.)
- For any non-text elements that would normally appear: describe them briefly in square brackets
- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__
- Output clean, readable text only (no commentary)"""
        self.instructions_text.setPlainText(default_instructions)
        instructions_group_layout.addWidget(self.instructions_text)
        
        instructions_widget_layout.addWidget(instructions_group)
        right_splitter.addWidget(instructions_widget)
        
        # Bottom: Extracted Text Preview
        preview_widget = QWidget()
        preview_widget_layout = QVBoxLayout(preview_widget)
        preview_widget_layout.setContentsMargins(0, 0, 0, 0)
        preview_widget_layout.setSpacing(0)
        
        preview_label = QLabel("Extracted Text Preview")
        preview_label.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        preview_label.setContentsMargins(0, 5, 0, 5)
        preview_widget_layout.addWidget(preview_label)
        
        self.preview_text = QTextEdit()
        self.preview_text.setFont(QFont("Segoe UI", 9))
        self.preview_text.setReadOnly(False)  # Allow editing
        self.preview_text.setWordWrapMode(QTextOption.WrapMode.WordWrap)
        preview_widget_layout.addWidget(self.preview_text)
        
        right_splitter.addWidget(preview_widget)
        
        # Set initial splitter sizes for right panel (30% instructions, 70% preview)
        right_splitter.setSizes([150, 350])
        
        splitter.addWidget(right_splitter)
        
        # Set splitter sizes (1:2 ratio)
        splitter.setSizes([300, 600])
        
        main_layout.addWidget(splitter, 1)  # 1 = stretch factor, expands to fill space
        
        # Action buttons
        action_layout = QHBoxLayout()
        action_layout.setSpacing(5)
        
        process_selected_btn = QPushButton("üîç Process Selected")
        process_selected_btn.clicked.connect(self._process_selected)
        process_selected_btn.setStyleSheet("background-color: #FF9800; color: white; font-weight: bold; padding: 6px 15px;")
        process_selected_btn.setToolTip("Process the currently selected image with AI OCR to extract text")
        action_layout.addWidget(process_selected_btn)
        
        process_all_btn = QPushButton("‚ö° Process ALL")
        process_all_btn.clicked.connect(self._process_all)
        process_all_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 6px 15px;")
        process_all_btn.setToolTip("Process all images in the list with AI OCR. This will use API credits and may take several minutes.")
        action_layout.addWidget(process_all_btn)
        
        save_docx_btn = QPushButton("üíæ Export Markdown && Word")
        save_docx_btn.clicked.connect(self._export_markdown_and_word)
        save_docx_btn.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold; padding: 6px 15px;")
        save_docx_btn.setToolTip("Export extracted text as Markdown (.md), Word document (.docx), and session report")
        action_layout.addWidget(save_docx_btn)
        
        copy_all_btn = QPushButton("üìã Copy All")
        copy_all_btn.clicked.connect(self._copy_all_text)
        copy_all_btn.setStyleSheet("background-color: #607D8B; color: white; font-weight: bold; padding: 6px 15px;")
        copy_all_btn.setToolTip("Copy all extracted text from all processed images to clipboard")
        action_layout.addWidget(copy_all_btn)
        
        session_report_btn = QPushButton("üìä Session Report")
        session_report_btn.clicked.connect(self._save_session_report)
        session_report_btn.setStyleSheet("background-color: #795548; color: white; font-weight: bold; padding: 6px 15px;")
        session_report_btn.setToolTip("Generate and save a detailed session report in Markdown format with statistics and extracted text")
        action_layout.addWidget(session_report_btn)
        
        action_layout.addStretch()
        main_layout.addLayout(action_layout, 0)  # 0 = no stretch, stays compact
        
        # Status
        self.status_label = QLabel("Ready - Add images to begin")
        self.status_label.setFont(QFont("Segoe UI", 9))
        self.status_label.setStyleSheet("color: #666;")
        main_layout.addWidget(self.status_label, 0)  # 0 = no stretch, stays compact
        
        # Progress bar
        self.progress = QProgressBar()
        self.progress.setMinimum(0)
        self.progress.setMaximum(100)
        self.progress.setValue(0)
        main_layout.addWidget(self.progress, 0)  # 0 = no stretch, stays compact
        
        # Restore state after UI creation
        self.image_files = saved_files
        self.extracted_texts = saved_texts
        if self.image_files:
            self._update_listbox()
    
    # === File Management Methods ===
    
    def _import_from_pdf(self):
        """Import images directly from a PDF file (simple OCR-only version)"""
        pdf_file, _ = QFileDialog.getOpenFileName(
            parent=None,
            caption="Select PDF File",
            filter="PDF files (*.pdf);;All files (*.*)"
        )
        
        if not pdf_file:
            return
        
        try:
            # Open PDF
            doc = fitz.open(pdf_file)
            total_pages = len(doc)
            
            if total_pages == 0:
                QMessageBox.warning(None, "Empty PDF", "The selected PDF has no pages.")
                return
            
            # Create folder for extracted images next to the PDF
            pdf_path = Path(pdf_file)
            pdf_name = pdf_path.stem
            images_folder = pdf_path.parent / f"{pdf_name}_images"
            
            # Create folder if it doesn't exist
            images_folder.mkdir(exist_ok=True)
            temp_dir = str(images_folder)
            
            # Log start
            self.log_message(f"Starting PDF import: {pdf_path.name}")
            self.log_message(f"Total pages: {total_pages}")
            
            # Extract each page as an image
            extracted_count = 0
            self.status_label.setText(f"Extracting pages from PDF...")
            QApplication.processEvents()
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # Render page to pixmap (image) at 2x resolution for better quality
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # Save as PNG
                img_filename = f"{pdf_name}_page_{page_num + 1:03d}.png"
                img_path = os.path.join(temp_dir, img_filename)
                pix.save(img_path)
                
                # Add to image list
                if img_path not in self.image_files:
                    self.image_files.append(img_path)
                    extracted_count += 1
                    
                    # Log each page
                    self.log_message(f"  Page {page_num + 1}/{total_pages} extracted: {img_filename}")
                
                # Update progress
                self.status_label.setText(
                    f"Extracting page {page_num + 1}/{total_pages}..."
                )
                self.progress.setValue(page_num + 1)
                self.progress.setMaximum(total_pages)
                QApplication.processEvents()
            
            doc.close()
            
            # Update list
            self._update_listbox()
            
            # Log completion
            self.log_message(f"PDF import complete: {extracted_count} pages extracted to {images_folder}")
            self.status_label.setText(f"‚úì Imported {extracted_count} pages from PDF")
            
            QMessageBox.information(
                None,
                "PDF Import Complete",
                f"Successfully extracted {extracted_count} pages from PDF!\n\n"
                f"Images saved to:\n{images_folder}\n\n"
                f"You can now process them with AI OCR."
            )
            
        except Exception as e:
            QMessageBox.critical(None, "PDF Import Error", f"Failed to import PDF:\n\n{str(e)}")
            self.log_message(f"ERROR importing PDF: {str(e)}")
            self.status_label.setText("PDF import failed")
    
    def _add_files(self):
        """Add individual image files"""
        files, _ = QFileDialog.getOpenFileNames(
            parent=None,
            caption="Select Image Files",
            filter="Image files (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;All files (*.*)"
        )
        
        if files:
            for file in files:
                if file not in self.image_files:
                    self.image_files.append(file)
            self._update_listbox()
            self.status_label.setText(f"Added {len(files)} file(s)")
            self.log_message(f"Added {len(files)} image file(s)")
    
    def _add_folder(self):
        """Add all images from a folder"""
        folder = QFileDialog.getExistingDirectory(
            parent=None,
            caption="Select Folder with Images"
        )
        
        if folder:
            image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}
            files = []
            
            for file in sorted(os.listdir(folder)):
                file_path = os.path.join(folder, file)
                if os.path.isfile(file_path):
                    ext = os.path.splitext(file)[1].lower()
                    if ext in image_extensions and file_path not in self.image_files:
                        files.append(file_path)
            
            self.image_files.extend(files)
            self._update_listbox()
            self.status_label.setText(f"Added {len(files)} file(s) from folder")
            self.log_message(f"Added {len(files)} file(s) from folder: {folder}")
    
    def _clear_list(self):
        """Clear all files"""
        if self.image_files:
            reply = QMessageBox.question(
                None,
                "Clear",
                "Remove all files?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                self.image_files = []
                self.extracted_texts = {}
                self._update_listbox()
                self.preview_text.clear()
                self.status_label.setText("List cleared")
    
    def _update_listbox(self):
        """Update file listbox with checkmarks for processed items"""
        self.file_listbox.clear()
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            status = "‚úì " if file in self.extracted_texts else ""
            self.file_listbox.addItem(f"{status}{i:2d}. {filename}")
    
    def _on_file_select(self):
        """Show extracted text when file is selected"""
        selected_items = self.file_listbox.selectedItems()
        if not selected_items:
            return
        
        idx = self.file_listbox.row(selected_items[0])
        if idx < len(self.image_files):
            file = self.image_files[idx]
            if file in self.extracted_texts:
                self.preview_text.setPlainText(self.extracted_texts[file])
    
    def _show_full_prompt(self):
        """Show the exact prompt that will be sent to the AI"""
        instructions = self.instructions_text.toPlainText().strip()
        
        # Apply formatting modifications like in _extract_text_from_image
        if self.preserve_formatting_check.isChecked():
            if "markdown for text formatting" not in instructions:
                instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
        else:
            instructions = instructions.replace(
                "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
            ).replace(
                "- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
            )
        
        # Create popup dialog
        popup = QDialog()
        popup.setWindowTitle("Full Prompt Preview")
        popup.resize(700, 600)
        
        layout = QVBoxLayout(popup)
        
        # Title
        title = QLabel("Exact Prompt Sent to OpenAI API")
        title.setFont(QFont("Segoe UI", 12, QFont.Weight.Bold))
        layout.addWidget(title)
        
        # Info frame
        info_group = QGroupBox("Configuration")
        info_layout = QVBoxLayout(info_group)
        
        model_label = QLabel(f"Model: {self.model_combo.currentText()}")
        model_label.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        info_layout.addWidget(model_label)
        
        formatting_status = "‚úì Enabled" if self.preserve_formatting_check.isChecked() else "‚úó Disabled"
        formatting_label = QLabel(f"Formatting Preservation: {formatting_status}")
        formatting_label.setFont(QFont("Segoe UI", 9))
        info_layout.addWidget(formatting_label)
        
        tokens_label = QLabel("Max Tokens: 4000")
        tokens_label.setFont(QFont("Segoe UI", 9))
        info_layout.addWidget(tokens_label)
        
        layout.addWidget(info_group)
        
        # Prompt text
        prompt_group = QGroupBox("Full Instructions Text")
        prompt_layout = QVBoxLayout(prompt_group)
        
        prompt_text = QPlainTextEdit()
        prompt_text.setFont(QFont("Consolas", 9))
        prompt_text.setPlainText(instructions)
        prompt_text.setReadOnly(True)
        prompt_layout.addWidget(prompt_text)
        
        layout.addWidget(prompt_group)
        
        # Note
        note = QLabel(
            "Note: The image is sent as base64-encoded data along with these instructions."
        )
        note.setStyleSheet("color: #666;")
        note.setFont(QFont("Segoe UI", 8))
        layout.addWidget(note)
        
        # Close button
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        button_box.rejected.connect(popup.close)
        layout.addWidget(button_box)
        
        popup.exec()
    
    def _save_session_report(self):
        """Generate and save a session report in markdown format"""
        if not self.extracted_texts:
            QMessageBox.warning(
                None,
                "No Data",
                "No OCR processing has been performed yet.\n\n"
                "Process some images first to generate a session report."
            )
            return
        
        # Ask for save location
        output_file, _ = QFileDialog.getSaveFileName(
            parent=None,
            caption="Save Session Report",
            filter="Markdown files (*.md);;Text files (*.txt);;All files (*.*)",
            initialFilter="Markdown files (*.md)"
        )
        
        if not output_file:
            return
        
        try:
            from datetime import datetime
            
            # Generate report content
            report_lines = []
            report_lines.append("# PDF Rescue - Session Report\n")
            report_lines.append("**Generated by [Supervertaler](https://supervertaler.com/) ‚Ä¢ by Michael Beijer**\n\n")
            report_lines.append(f"**Date**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            report_lines.append("---\n\n")
            
            # Configuration section
            report_lines.append("## Configuration\n\n")
            report_lines.append(f"- **Model**: {self.model_combo.currentText()}\n")
            formatting_status = "Enabled ‚úì" if self.preserve_formatting_check.isChecked() else "Disabled ‚úó"
            report_lines.append(f"- **Formatting Preservation**: {formatting_status}\n")
            report_lines.append(f"- **Total Images Processed**: {len(self.extracted_texts)}\n")
            report_lines.append(f"- **Total Images in List**: {len(self.image_files)}\n\n")
            
            # Instructions used
            report_lines.append("## Extraction Instructions\n\n")
            report_lines.append("```\n")
            instructions = self.instructions_text.toPlainText().strip()
            if self.preserve_formatting_check.isChecked():
                if "markdown for text formatting" not in instructions:
                    instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
            report_lines.append(instructions)
            report_lines.append("\n```\n\n")
            
            # Processing summary
            report_lines.append("## Processing Summary\n\n")
            report_lines.append("| # | Image File | Status |\n")
            report_lines.append("|---|------------|--------|\n")
            
            for i, file in enumerate(self.image_files, 1):
                filename = os.path.basename(file)
                status = "‚úì Processed" if file in self.extracted_texts else "‚ßó Pending"
                report_lines.append(f"| {i} | {filename} | {status} |\n")
            
            report_lines.append("\n---\n\n")
            
            # Extracted text for each image
            report_lines.append("## Extracted Text\n\n")
            
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    filename = os.path.basename(file)
                    report_lines.append(f"### Page {i}: {filename}\n\n")
                    report_lines.append("```\n")
                    report_lines.append(self.extracted_texts[file])
                    report_lines.append("\n```\n\n")
                    report_lines.append("---\n\n")
            
            # Statistics
            report_lines.append("## Statistics\n\n")
            texts_list = list(self.extracted_texts.values())
            total_chars = sum(len(text) for text in texts_list)
            total_words = sum(len(text.split()) for text in texts_list)
            report_lines.append(f"- **Total Characters Extracted**: {total_chars:,}\n")
            report_lines.append(f"- **Total Words Extracted**: {total_words:,}\n")
            report_lines.append(f"- **Average Characters per Page**: {total_chars // len(self.extracted_texts) if self.extracted_texts else 0:,}\n")
            report_lines.append(f"- **Average Words per Page**: {total_words // len(self.extracted_texts) if self.extracted_texts else 0:,}\n\n")
            
            # Footer
            report_lines.append("---\n\n")
            report_lines.append("*Report generated by **PDF Rescue** - AI-Powered OCR Tool*\n\n")
            report_lines.append("*Part of [**Supervertaler**](https://supervertaler.com/) ‚Ä¢ by Michael Beijer*\n")
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.writelines(report_lines)
            
            self.log_message(f"Session report saved: {Path(output_file).name}")
            self.status_label.setText(f"‚úì Report saved to {os.path.basename(output_file)}")
            
            reply = QMessageBox.question(
                None,
                "Success",
                f"Session report saved successfully!\n\n"
                f"File: {Path(output_file).name}\n\n"
                "Open the report now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                os.startfile(output_file)
        
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Failed to save report:\n\n{str(e)}")
    
    # === OCR Processing Methods ===
    
    def _encode_image(self, image_path):
        """Encode image to base64"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def _extract_text_from_image(self, image_path):
        """Use AI Vision to extract text from image (supports OpenAI, Claude, Gemini)"""
        try:
            instructions = self.instructions_text.toPlainText().strip()
            
            # Add or remove formatting instruction based on checkbox
            if self.preserve_formatting_check.isChecked():
                if "markdown for text formatting" not in instructions:
                    instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
            else:
                # Remove markdown instruction if present
                instructions = instructions.replace(
                    "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
                ).replace(
                    "- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
                )
            
            # Get selected model
            model = self.model_combo.currentText()
            
            # Skip separator items
            if model.startswith("---"):
                return "[ERROR: Please select a valid model, not a separator]"
            
            # Determine provider from model name
            provider = self._get_provider_from_model(model)
            
            if provider not in self.clients:
                return f"[ERROR: {provider.title()} client not initialized. Check API key in api_keys.txt]"
            
            # Call appropriate provider
            if provider == "openai":
                return self._extract_with_openai(image_path, model, instructions)
            elif provider == "claude":
                return self._extract_with_claude(image_path, model, instructions)
            elif provider == "gemini":
                return self._extract_with_gemini(image_path, model, instructions)
            else:
                return f"[ERROR: Unknown provider: {provider}]"
        
        except Exception as e:
            return f"[ERROR extracting text: {str(e)}]"
    
    def _get_provider_from_model(self, model):
        """Determine provider from model name"""
        if model.startswith("gpt") or model.startswith("o1") or model.startswith("o3"):
            return "openai"
        elif model.startswith("claude"):
            return "claude"
        elif model.startswith("gemini"):
            return "gemini"
        else:
            return "unknown"
    
    def _extract_with_openai(self, image_path, model, instructions):
        """Extract text using OpenAI Vision API"""
        base64_image = self._encode_image(image_path)
        
        api_params = {
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": instructions},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ]
        }
        
        # Use appropriate token parameter based on model
        if model.startswith("gpt-5") or model.startswith("o1"):
            api_params["max_completion_tokens"] = 4000
        else:
            api_params["max_tokens"] = 4000
        
        response = self.clients['openai'].chat.completions.create(**api_params)
        return response.choices[0].message.content
    
    def _extract_with_claude(self, image_path, model, instructions):
        """Extract text using Claude Vision API"""
        import base64
        
        # Read image and encode to base64
        with open(image_path, "rb") as image_file:
            image_data = base64.standard_b64encode(image_file.read()).decode("utf-8")
        
        # Determine media type from file extension
        ext = os.path.splitext(image_path)[1].lower()
        media_type_map = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp'
        }
        media_type = media_type_map.get(ext, 'image/jpeg')
        
        response = self.clients['claude'].messages.create(
            model=model,
            max_tokens=4000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": image_data,
                            },
                        },
                        {
                            "type": "text",
                            "text": instructions
                        }
                    ],
                }
            ],
        )
        
        return response.content[0].text
    
    def _extract_with_gemini(self, image_path, model, instructions):
        """Extract text using Gemini Vision API"""
        from PIL import Image
        
        # Load image
        img = Image.open(image_path)
        
        # Create model instance
        gemini_model = self.clients['gemini'].GenerativeModel(model)
        
        # Generate content with image and prompt
        response = gemini_model.generate_content([instructions, img])
        
        return response.text
    
    def _process_selected(self):
        """Process currently selected image"""
        selected_items = self.file_listbox.selectedItems()
        if not selected_items:
            QMessageBox.warning(None, "No Selection", "Please select an image to process")
            return
        
        idx = self.file_listbox.row(selected_items[0])
        if idx >= len(self.image_files):
            return
        
        file = self.image_files[idx]
        filename = os.path.basename(file)
        
        # Get provider info for status display
        model = self.model_combo.currentText()
        provider = self._get_provider_from_model(model).title()
        
        self.log_message(f"Processing selected image with {provider} ({model}): {filename}")
        self.status_label.setText(f"Processing with {provider}... {filename}")
        QApplication.processEvents()
        
        text = self._extract_text_from_image(file)
        self.extracted_texts[file] = text
        
        self.preview_text.setPlainText(text)
        
        self._update_listbox()
        self.log_message(f"Successfully processed: {filename}")
        self.status_label.setText(f"‚úì Processed {filename} with {provider}")
    
    def _process_all(self):
        """Process all images in the list"""
        if not self.image_files:
            QMessageBox.warning(None, "No Files", "Please add images first")
            return
        
        # Get provider info for confirmation
        model = self.model_combo.currentText()
        provider = self._get_provider_from_model(model).title()
        
        reply = QMessageBox.question(
            None,
            "Process All",
            f"Process all {len(self.image_files)} images with {provider} ({model})?\n\n"
            "This will use API credits and may take several minutes.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        self.log_message(f"Starting batch processing with {provider}: {len(self.image_files)} images")
        self.progress.setMaximum(len(self.image_files))
        self.progress.setValue(0)
        
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            self.status_label.setText(f"[{provider}] Processing {i}/{len(self.image_files)}: {filename}...")
            QApplication.processEvents()
            
            if file not in self.extracted_texts:
                text = self._extract_text_from_image(file)
                self.extracted_texts[file] = text
                self.log_message(f"  [{i}/{len(self.image_files)}] Processed: {filename}")
            else:
                self.log_message(f"  [{i}/{len(self.image_files)}] Skipped (already processed): {filename}")
            
            self.progress.setValue(i)
            self._update_listbox()
        
        self.log_message(f"Batch processing complete: {len(self.image_files)} images processed")
        self.status_label.setText(f"‚úì Processed all {len(self.image_files)} images with {provider}!")
        QMessageBox.information(
            None,
            "Complete",
            f"Successfully processed {len(self.image_files)} images with {provider}!\n\n"
            "Click 'Save DOCX' to export the text."
        )
    
    # === Export Methods ===
    
    def _add_formatted_text(self, doc, text):
        """
        Add text to document with markdown formatting parsed
        Supports: **bold**, *italic*, __underline__
        Also handles multi-column layouts with [START COLUMN X] / [END COLUMN X] markers
        """
        # Check if text has column markers
        if '[START COLUMN' in text or '[COLUMN 1]' in text or '[COLUMN 2]' in text:
            self._add_multi_column_text(doc, text)
            return
        
        # Split text into paragraphs (separated by double newlines or more)
        # This prevents treating every single line as a paragraph
        paragraphs = re.split(r'\n\s*\n', text)
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Replace single newlines within paragraph with spaces (removes extraneous line breaks)
            para_text = para_text.replace('\n', ' ').strip()
            
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(6)  # Reduced from 12 to 6 for tighter spacing
            
            # Parse markdown formatting using regex
            remaining = para_text
            
            while remaining:
                # Check for bold (**text**)
                bold_match = re.match(r'\*\*(.*?)\*\*', remaining)
                if bold_match:
                    run = para.add_run(bold_match.group(1))
                    run.bold = True
                    remaining = remaining[bold_match.end():]
                    continue
                
                # Check for underline (__text__)
                underline_match = re.match(r'__(.*?)__', remaining)
                if underline_match:
                    run = para.add_run(underline_match.group(1))
                    run.underline = True
                    remaining = remaining[underline_match.end():]
                    continue
                
                # Check for italic (*text*)
                italic_match = re.match(r'\*(.*?)\*', remaining)
                if italic_match:
                    run = para.add_run(italic_match.group(1))
                    run.italic = True
                    remaining = remaining[italic_match.end():]
                    continue
                
                # No formatting - add plain text until next marker or end
                next_marker = len(remaining)
                for marker in ['**', '*', '__']:
                    pos = remaining.find(marker)
                    if pos != -1 and pos < next_marker:
                        next_marker = pos
                
                if next_marker == 0:
                    # Edge case: marker at start but no match (e.g., single * or **)
                    para.add_run(remaining[0])
                    remaining = remaining[1:]
                else:
                    plain_text = remaining[:next_marker] if next_marker < len(remaining) else remaining
                    if plain_text:
                        para.add_run(plain_text)
                    remaining = remaining[next_marker:]
    
    def _add_multi_column_text(self, doc, text):
        """
        Handle multi-column text layout using a Word table
        Supports markers like [START COLUMN 1], [END COLUMN 1], etc.
        """
        # Parse columns from text
        columns = {}
        current_column = None
        lines = text.split('\n')
        
        for line in lines:
            # Check for column start marker
            if '[START COLUMN' in line.upper() or '[COLUMN' in line.upper():
                # Extract column number
                import re
                match = re.search(r'\[(?:START )?COLUMN[:\s]+(\d+)\]', line, re.IGNORECASE)
                if match:
                    current_column = int(match.group(1))
                    if current_column not in columns:
                        columns[current_column] = []
                continue
            
            # Check for column end marker
            if '[END COLUMN' in line.upper():
                current_column = None
                continue
            
            # Add line to current column
            if current_column is not None:
                columns[current_column].append(line)
        
        # If we found columns, create a table layout
        if columns:
            num_columns = max(columns.keys()) if columns else 2
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = 'Table Grid'
            
            # Set column widths to be equal
            from docx.shared import Inches
            for col_idx in range(num_columns):
                for cell in table.columns[col_idx].cells:
                    cell.width = Inches(6.5 / num_columns)
            
            # Fill each column
            for col_num in sorted(columns.keys()):
                if col_num <= num_columns:
                    cell = table.cell(0, col_num - 1)
                    column_text = '\n'.join(columns[col_num])
                    
                    # Remove the cell's default paragraph and add formatted text
                    cell.text = ''
                    for para_text in column_text.split('\n'):
                        if para_text.strip():
                            para = cell.add_paragraph(para_text.strip())
                            para.paragraph_format.line_spacing = 1.0
                            para.paragraph_format.space_after = Pt(6)
            
            # Remove table borders for cleaner look
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
                    if tcBorders is None:
                        from docx.oxml import OxmlElement
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    # Set all borders to none
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = tcBorders.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + border_name)
                        if border is None:
                            from docx.oxml import OxmlElement
                            border = OxmlElement(f'w:{border_name}')
                            tcBorders.append(border)
                        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'none')
        else:
            # No columns found, fall back to regular formatting
            self._add_formatted_text(doc, text)
    
    def _parse_markdown_table(self, table_lines):
        """Parse markdown table lines into rows and cells"""
        rows = []
        for line in table_lines:
            line = line.strip()
            if not line or not line.startswith('|'):
                continue
            # Split by pipe, remove first and last empty elements
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Skip separator row (contains only dashes, colons, and spaces)
            is_separator = all(
                cell.replace('-', '').replace(':', '').replace(' ', '').strip() == '' 
                for cell in cells
            )
            if is_separator:
                continue
            rows.append(cells)
        return rows
    
    def _extract_markdown_tables(self, text):
        """Extract markdown tables from text and return (text_without_tables, tables_list)"""
        lines = text.split('\n')
        result_lines = []
        tables = []
        current_table_lines = []
        in_table = False
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            # Check if line looks like a table row (starts and ends with |)
            is_table_row = stripped.startswith('|') and stripped.endswith('|') and '|' in stripped[1:-1]
            
            if is_table_row:
                if not in_table:
                    in_table = True
                    current_table_lines = []
                current_table_lines.append(line)
            else:
                # End of table
                if in_table:
                    # Parse the table
                    parsed_table = self._parse_markdown_table(current_table_lines)
                    if parsed_table and len(parsed_table) > 0:
                        tables.append(parsed_table)
                    # Add placeholder for table
                    result_lines.append(f"[TABLE_{len(tables) - 1}]")
                    in_table = False
                    current_table_lines = []
                result_lines.append(line)
        
        # Handle table at end of text
        if in_table and current_table_lines:
            parsed_table = self._parse_markdown_table(current_table_lines)
            if parsed_table and len(parsed_table) > 0:
                tables.append(parsed_table)
            result_lines.append(f"[TABLE_{len(tables) - 1}]")
        
        text_without_tables = '\n'.join(result_lines)
        return text_without_tables, tables
    
    def _add_markdown_table_to_doc(self, doc, table_rows):
        """Add a Word table from markdown table rows"""
        if not table_rows or len(table_rows) == 0:
            return
        
        # Determine number of columns (use first row, usually header)
        num_cols = len(table_rows[0]) if table_rows else 1
        num_rows = len(table_rows)
        
        # Create Word table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Populate table cells
        for row_idx, row_data in enumerate(table_rows):
            # Ensure row has enough cells
            while len(row_data) < num_cols:
                row_data.append('')
            
            for col_idx in range(num_cols):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ''
                # Remove markdown formatting from table cells (keep plain text for now)
                # Could enhance this to support formatting in cells later
                cell_text = cell_text.replace('**', '').replace('*', '').replace('__', '')
                table.rows[row_idx].cells[col_idx].text = cell_text
        
        # Add spacing after table
        doc.add_paragraph()
    
    def _add_formatted_text_with_tables(self, doc, text):
        """Add text to document with markdown formatting and tables parsed"""
        # Extract tables first
        text_without_tables, tables = self._extract_markdown_tables(text)
        
        # Split text into paragraphs (now with table placeholders)
        parts = text_without_tables.split('\n')
        
        current_paragraph_text = []
        
        for part in parts:
            # Check if this is a table placeholder
            if part.strip().startswith('[TABLE_') and part.strip().endswith(']'):
                # Add any accumulated paragraph text
                if current_paragraph_text:
                    para_text = '\n'.join(current_paragraph_text).strip()
                    if para_text:
                        self._add_formatted_text(doc, para_text)
                    current_paragraph_text = []
                
                # Extract table index
                try:
                    table_idx = int(part.strip()[7:-1])  # Extract number from [TABLE_0]
                    if 0 <= table_idx < len(tables):
                        self._add_markdown_table_to_doc(doc, tables[table_idx])
                except (ValueError, IndexError):
                    pass
            else:
                current_paragraph_text.append(part)
        
        # Add any remaining paragraph text
        if current_paragraph_text:
            para_text = '\n'.join(current_paragraph_text).strip()
            if para_text:
                self._add_formatted_text(doc, para_text)
    
    def _generate_markdown_file(self, output_file):
        """Generate markdown file from extracted texts"""
        with open(output_file, 'w', encoding='utf-8') as f:
            # Header
            f.write("# Extracted Text from Images\n\n")
            f.write("*Generated by PDF Rescue (a Supervertaler module)*\n\n")
            f.write("---\n\n")
            
            # Add extracted text in order
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    # Page header
                    f.write(f"## Page {i}: {os.path.basename(file)}\n\n")
                    
                    # Text content (already in markdown format)
                    text = self.extracted_texts[file]
                    f.write(text)
                    f.write("\n\n")
                    f.write("---\n\n")
    
    def _markdown_to_docx(self, md_file, docx_file):
        """Convert markdown file to DOCX with proper table handling"""
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Parse markdown content line by line
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()  # Keep left whitespace, remove right
            
            # Skip empty lines (add spacing as paragraph)
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue
            
            # Handle headers
            if line.startswith('# '):
                title = doc.add_heading(line[2:].strip(), 0)
                title.runs[0].font.size = Pt(16)
                i += 1
                continue
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:].strip(), level=2)
                heading.runs[0].font.size = Pt(12)
                i += 1
                continue
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue
            
            # Handle horizontal rule
            if line.strip() == '---' or line.strip() == '***':
                doc.add_paragraph()
                i += 1
                continue
            
            # Handle italics line (branding) - simple single-line italic
            stripped = line.strip()
            if stripped.startswith('*') and stripped.endswith('*') and stripped.count('*') == 2 and len(stripped) > 2:
                para = doc.add_paragraph()
                run = para.add_run(stripped[1:-1])
                run.italic = True
                run.font.size = Pt(9)
                i += 1
                continue
            
            # Check if this might be a table
            if line.strip().startswith('|') and line.strip().endswith('|') and '|' in line.strip()[1:-1]:
                # Collect all table lines
                table_lines = [line]
                i += 1
                while i < len(lines):
                    next_line = lines[i].rstrip()
                    # Continue if it's a table row or separator row
                    if next_line.strip().startswith('|') and next_line.strip().endswith('|'):
                        table_lines.append(next_line)
                        i += 1
                    else:
                        break
                
                # Parse and add table
                parsed_table = self._parse_markdown_table(table_lines)
                if parsed_table and len(parsed_table) > 0:
                    self._add_markdown_table_to_doc(doc, parsed_table)
                continue
            
            # Regular paragraph text - collect until empty line or special marker
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].rstrip()
                # Stop at empty line or special markers
                if not next_line.strip():
                    break
                if next_line.strip().startswith('#') or next_line.strip() in ['---', '***']:
                    break
                if next_line.strip().startswith('|') and next_line.strip().endswith('|'):
                    break
                para_lines.append(next_line)
                i += 1
            
            # Process paragraph with tables
            para_text = '\n'.join(para_lines)
            if para_text.strip():
                self._add_formatted_text_with_tables(doc, para_text)
        
        # Save document
        doc.save(docx_file)
    
    def _export_markdown_and_word(self):
        """Export extracted text as markdown, convert to DOCX, and generate session report"""
        if not self.extracted_texts:
            QMessageBox.warning(
                None,
                "No Text",
                "No extracted text to export.\n\n"
                "Process images first."
            )
            return
        
        # Ask for base file name (will generate .md, .docx, and _report.md)
        output_file, _ = QFileDialog.getSaveFileName(
            parent=None,
            caption="Export Markdown & Word Documents",
            filter="Markdown files (*.md);;All files (*.*)",
            initialFilter="Markdown files (*.md)"
        )
        
        if not output_file:
            return
        
        # Generate file names
        base_path = Path(output_file)
        if base_path.suffix != '.md':
            base_path = base_path.with_suffix('.md')
        
        md_file = str(base_path)
        docx_file = str(base_path.with_suffix('.docx'))
        report_file = str(base_path.parent / f"{base_path.stem}_report.md")
        
        self.log_message(f"Exporting documents: {base_path.name}")
        
        try:
            # Step 1: Generate markdown file
            self._generate_markdown_file(md_file)
            self.log_message(f"‚úì Markdown file created: {base_path.name}")
            
            # Step 2: Convert markdown to DOCX
            self._markdown_to_docx(md_file, docx_file)
            self.log_message(f"‚úì Word document created: {base_path.stem}.docx")
            
            # Step 3: Generate session report
            self._save_session_report_to_file(report_file)
            self.log_message(f"‚úì Session report created: {base_path.stem}_report.md")
            
            # Success message
            files_created = f"‚úì {base_path.name}\n‚úì {base_path.stem}.docx\n‚úì {base_path.stem}_report.md"
            self.status_label.setText(f"‚úì Exported 3 files successfully")
            
            reply = QMessageBox.question(
                None,
                "Export Complete",
                f"Successfully exported all documents!\n\n"
                f"Files created:\n{files_created}\n\n"
                "Open the files now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                os.startfile(md_file)
                os.startfile(docx_file)
                os.startfile(report_file)
        
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Failed to export documents:\n\n{str(e)}")
            self.log_message(f"‚úó Export error: {str(e)}")
    
    def _save_session_report_to_file(self, output_file):
        """Generate and save session report to specified file (internal method)"""
        from datetime import datetime
        
        # Generate report content
        report_lines = []
        report_lines.append("# PDF Rescue - Session Report\n")
        report_lines.append("**Generated by [Supervertaler](https://supervertaler.com/) ‚Ä¢ by Michael Beijer**\n\n")
        report_lines.append(f"**Date**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        report_lines.append("---\n\n")
        
        # Configuration section
        report_lines.append("## Configuration\n\n")
        report_lines.append(f"- **Model**: {self.model_combo.currentText()}\n")
        formatting_status = "Enabled ‚úì" if self.preserve_formatting_check.isChecked() else "Disabled ‚úó"
        report_lines.append(f"- **Formatting Preservation**: {formatting_status}\n")
        report_lines.append(f"- **Total Images Processed**: {len(self.extracted_texts)}\n")
        report_lines.append(f"- **Total Images in List**: {len(self.image_files)}\n\n")
        
        # Instructions used
        report_lines.append("## Extraction Instructions\n\n")
        report_lines.append("```\n")
        instructions = self.instructions_text.toPlainText().strip()
        if self.preserve_formatting_check.isChecked():
            if "markdown for text formatting" not in instructions:
                instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
        report_lines.append(instructions)
        report_lines.append("\n```\n\n")
        
        # Processing summary
        report_lines.append("## Processing Summary\n\n")
        report_lines.append("| # | Image File | Status |\n")
        report_lines.append("|---|------------|--------|\n")
        
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            status = "‚úì Processed" if file in self.extracted_texts else "‚ßó Pending"
            report_lines.append(f"| {i} | {filename} | {status} |\n")
        
        report_lines.append("\n---\n\n")
        
        # Extracted text for each image
        report_lines.append("## Extracted Text\n\n")
        
        for i, file in enumerate(self.image_files, 1):
            if file in self.extracted_texts:
                filename = os.path.basename(file)
                report_lines.append(f"### Page {i}: {filename}\n\n")
                report_lines.append("```\n")
                report_lines.append(self.extracted_texts[file])
                report_lines.append("\n```\n\n")
                report_lines.append("---\n\n")
        
        # Statistics
        report_lines.append("## Statistics\n\n")
        texts_list = list(self.extracted_texts.values())
        total_chars = sum(len(text) for text in texts_list)
        total_words = sum(len(text.split()) for text in texts_list)
        report_lines.append(f"- **Total Characters Extracted**: {total_chars:,}\n")
        report_lines.append(f"- **Total Words Extracted**: {total_words:,}\n")
        report_lines.append(f"- **Average Characters per Page**: {total_chars // len(self.extracted_texts) if self.extracted_texts else 0:,}\n")
        report_lines.append(f"- **Average Words per Page**: {total_words // len(self.extracted_texts) if self.extracted_texts else 0:,}\n\n")
        
        # Footer
        report_lines.append("---\n\n")
        report_lines.append("*Report generated by **PDF Rescue** - AI-Powered OCR Tool*\n\n")
        report_lines.append("*Part of [**Supervertaler**](https://supervertaler.com/) ‚Ä¢ by Michael Beijer*\n")
        
        # Write to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.writelines(report_lines)
    
    def _save_to_docx(self):
        """Save all extracted text to a Word document"""
        if not self.extracted_texts:
            QMessageBox.warning(
                None,
                "No Text",
                "No extracted text to save.\n\n"
                "Process images first."
            )
            return
        
        output_file, _ = QFileDialog.getSaveFileName(
            parent=None,
            caption="Save Extracted Text",
            filter="Word Document (*.docx);;All files (*.*)",
            initialFilter="Word Document (*.docx)"
        )
        
        if not output_file:
            return
        
        self.log_message(f"Saving extracted text to DOCX: {Path(output_file).name}")
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Extracted Text from Images', 0)
            title.runs[0].font.size = Pt(16)
            
            # Add branding with hyperlink to Supervertaler
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            def add_hyperlink(paragraph, text, url):
                """Add a hyperlink to a paragraph"""
                # Get the paragraph element
                part = paragraph.part
                r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
                
                # Create the hyperlink element
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                
                # Create a new run element
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                
                # Add hyperlink style
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'Hyperlink')
                rPr.append(rStyle)
                new_run.append(rPr)
                
                # Add the text
                new_run.text = text
                hyperlink.append(new_run)
                
                # Add hyperlink to paragraph
                paragraph._p.append(hyperlink)
                return hyperlink
            
            branding_para = doc.add_paragraph()
            
            # Add text before hyperlink
            run1 = branding_para.add_run('Generated by PDF Rescue (a ')
            run1.font.size = Pt(9)
            run1.italic = True
            
            # Add hyperlink
            add_hyperlink(branding_para, 'Supervertaler', 'https://supervertaler.com/')
            
            # Add text after hyperlink
            run2 = branding_para.add_run(' module)')
            run2.font.size = Pt(9)
            run2.italic = True
            
            # Add spacing
            doc.add_paragraph()
            
            # Add extracted text in order
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    # Page header
                    heading = doc.add_heading(f'Page {i}: {os.path.basename(file)}', level=2)
                    heading.runs[0].font.size = Pt(12)
                    
                    # Text content with formatting
                    text = self.extracted_texts[file]
                    if self.preserve_formatting_check.isChecked():
                        self._add_formatted_text(doc, text)
                    else:
                        # Split by double newlines to preserve paragraph breaks
                        # Replace single newlines with spaces to remove extraneous line breaks
                        paragraphs = re.split(r'\n\s*\n', text)
                        for para_text in paragraphs:
                            if para_text.strip():
                                # Replace single newlines within paragraph with spaces
                                para_text = para_text.replace('\n', ' ').strip()
                                para = doc.add_paragraph(para_text)
                                para.paragraph_format.line_spacing = 1.15
                                para.paragraph_format.space_after = Pt(6)
                    
                    # Page break except for last
                    if i < len(self.image_files):
                        doc.add_page_break()
            
            doc.save(output_file)
            
            self.log_message(f"Successfully saved {len(self.extracted_texts)} pages to: {Path(output_file).name}")
            self.status_label.setText(f"‚úì Saved to {os.path.basename(output_file)}")
            
            reply = QMessageBox.question(
                None,
                "Success",
                f"Document saved successfully!\n\n"
                f"{len(self.extracted_texts)} pages of text extracted\n\n"
                "Open the document now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                os.startfile(output_file)
        
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Failed to save document:\n\n{str(e)}")
    
    def _copy_all_text(self):
        """Copy all extracted text to clipboard"""
        if not self.extracted_texts:
            QMessageBox.warning(None, "No Text", "No extracted text to copy")
            return
        
        all_text = []
        for i, file in enumerate(self.image_files, 1):
            if file in self.extracted_texts:
                all_text.append(f"=== Page {i}: {os.path.basename(file)} ===\n")
                all_text.append(self.extracted_texts[file])
                all_text.append("\n\n")
        
        combined = "".join(all_text)
        
        # Use QApplication clipboard
        from PyQt6.QtGui import QClipboard
        app = QApplication.instance()
        if app:
            clipboard = app.clipboard()
            clipboard.setText(combined)
        
        self.status_label.setText(f"‚úì Copied {len(self.extracted_texts)} pages to clipboard")
        QMessageBox.information(None, "Copied", f"Copied text from {len(self.extracted_texts)} pages to clipboard!")


# === Standalone Application ===

if __name__ == "__main__":
    """Run PDF Rescue as a standalone application"""
    import sys
    from PyQt6.QtWidgets import QApplication, QMainWindow, QTextEdit, QLabel
    from PyQt6.QtCore import Qt
    
    class StandaloneApp(QMainWindow):
        """Minimal parent app for standalone mode"""
        def __init__(self):
            super().__init__()
            self.setWindowTitle("PDF Rescue - AI-Powered OCR Tool")
            self.setGeometry(100, 100, 1200, 800)
            
            # Central widget
            central_widget = QWidget()
            self.setCentralWidget(central_widget)
            
            # Main layout
            main_layout = QVBoxLayout(central_widget)
            main_layout.setContentsMargins(10, 10, 10, 10)
            
            # Title
            title = QLabel("PDF Rescue - AI-Powered OCR Tool")
            title.setFont(QFont("Segoe UI", 14, QFont.Weight.Bold))
            title.setAlignment(Qt.AlignmentFlag.AlignCenter)
            main_layout.addWidget(title)
            
            # Load API keys
            self.api_keys = {}
            api_file = Path("api_keys.txt")
            if not api_file.exists():
                # Try user_data folder
                api_file = Path("user_data_private" if os.path.exists(".supervertaler.local") else "user_data") / "api_keys.txt"
            
            if api_file.exists():
                with open(api_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#') and '=' in line:
                            key, value = line.split('=', 1)
                            if 'openai' in key.lower():
                                self.api_keys['openai'] = value.strip()
            
            if not self.api_keys.get('openai'):
                QMessageBox.critical(
                    self,
                    "API Key Missing",
                    "Could not find OpenAI API key in api_keys.txt\n\n"
                    "Please add a line like:\nOPENAI_API_KEY=your-key-here\n\n"
                    "Or place api_keys.txt in the user_data folder."
                )
                # Still create UI but warn user
            
            # Create PDF Rescue instance (standalone mode)
            self.pdf_rescue = PDFRescueQt(self, standalone=True)
            pdf_rescue_widget = QWidget()
            self.pdf_rescue.create_tab(pdf_rescue_widget)
            main_layout.addWidget(pdf_rescue_widget)
            
            # Add log at bottom
            log_group = QGroupBox("Activity Log")
            log_layout = QVBoxLayout(log_group)
            
            self.log_text = QTextEdit()
            self.log_text.setReadOnly(True)
            self.log_text.setMaximumHeight(150)
            self.log_text.setFont(QFont("Consolas", 9))
            log_layout.addWidget(self.log_text)
            
            main_layout.addWidget(log_group)
        
        def log(self, message: str):
            """Add message to log"""
            from datetime import datetime
            timestamp = datetime.now().strftime("%H:%M:%S")
            formatted_message = f"[{timestamp}] {message}"
            self.log_text.append(formatted_message)
        
        def load_api_keys(self):
            """Load API keys for compatibility"""
            return self.api_keys
    
    # Create and run standalone app
    app = QApplication(sys.argv)
    window = StandaloneApp()
    window.show()
    sys.exit(app.exec())

