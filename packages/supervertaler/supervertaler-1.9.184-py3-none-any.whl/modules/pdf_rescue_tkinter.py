"""
PDF Rescue Module
Embeddable version of the AI-powered OCR tool for extracting text from poorly formatted PDFs
Uses OpenAI's GPT-4 Vision API

This module can be embedded in the main Supervertaler application as a tab.
"""

import os
import base64
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from openai import OpenAI
from docx import Document
from docx.shared import Pt
import fitz  # PyMuPDF
import re


class PDFRescue:
    """
    PDF Rescue feature - extract text from images using AI OCR
    Can be embedded in any tkinter application as a tab or panel
    """
    
    def __init__(self, parent_app):
        """
        Initialize PDF Rescue module
        
        Args:
            parent_app: Reference to the main application (needs .api_keys attribute)
        """
        self.parent_app = parent_app
        self.client = None
        self.image_files = []
        self.extracted_texts = {}
        
        # Initialize OpenAI client
        api_key = None
        if hasattr(parent_app, 'api_keys'):
            api_key = parent_app.api_keys.get('openai')
        elif hasattr(parent_app, 'api_key'):
            api_key = parent_app.api_key
            
        if api_key:
            try:
                self.client = OpenAI(api_key=api_key)
            except Exception as e:
                print(f"Failed to initialize OpenAI client: {e}")
    
    def log_message(self, message: str):
        """Log a message to the parent app's log if available"""
        if hasattr(self.parent_app, 'log'):
            self.parent_app.log(f"[PDF Rescue] {message}")
        else:
            print(f"[PDF Rescue] {message}")
    
    def create_tab(self, parent):
        """
        Create the PDF Rescue tab UI
        
        Args:
            parent: The parent widget (notebook tab or frame)
        """
        # Save current state before recreating UI
        saved_files = self.image_files.copy() if hasattr(self, 'image_files') else []
        saved_texts = self.extracted_texts.copy() if hasattr(self, 'extracted_texts') else {}
        
        # Header
        header_frame = tk.Frame(parent, bg='#e3f2fd', relief='solid', borderwidth=1)
        header_frame.pack(fill='x', padx=5, pady=5)
        
        tk.Label(header_frame, text="üîç PDF Rescue - AI-Powered OCR", 
                font=('Segoe UI', 10, 'bold'), bg='#e3f2fd').pack(side='left', padx=10, pady=5)
        
        tk.Label(header_frame, text="Extract text from poorly formatted PDF screenshots",
                font=('Segoe UI', 9), bg='#e3f2fd', fg='#666').pack(side='left', padx=(0, 10), pady=5)
        
        # Split view: Files on left, Preview on right
        paned = ttk.PanedWindow(parent, orient='horizontal')
        paned.pack(fill='both', expand=True, padx=5, pady=5)
        
        # LEFT: File list
        left_frame = tk.Frame(paned)
        paned.add(left_frame, weight=1)
        
        tk.Label(left_frame, text="Images to Process", 
                font=('Segoe UI', 9, 'bold')).pack(anchor='w', pady=(0, 5))
        
        # File list with scrollbar
        list_container = tk.Frame(left_frame)
        list_container.pack(fill='both', expand=True)
        
        scroll = tk.Scrollbar(list_container, orient='vertical')
        scroll.pack(side='right', fill='y')
        
        self.file_listbox = tk.Listbox(list_container, yscrollcommand=scroll.set,
                                       font=('Consolas', 9))
        self.file_listbox.pack(fill='both', expand=True)
        scroll.config(command=self.file_listbox.yview)
        self.file_listbox.bind('<<ListboxSelect>>', self._on_file_select)
        
        # Buttons
        btn_frame = tk.Frame(left_frame)
        btn_frame.pack(fill='x', pady=(10, 0))
        
        tk.Button(btn_frame, text="üìÑ PDF", command=self._import_from_pdf,
                 bg='#9C27B0', fg='white', font=('Segoe UI', 8, 'bold'),
                 padx=8, pady=4).pack(side='left', padx=(0, 3))
        
        tk.Button(btn_frame, text="‚ûï Add Files", command=self._add_files,
                 bg='#2196F3', fg='white', font=('Segoe UI', 8, 'bold'),
                 padx=8, pady=4).pack(side='left', padx=3)
        
        tk.Button(btn_frame, text="üìÇ Folder", command=self._add_folder,
                 bg='#2196F3', fg='white', font=('Segoe UI', 8, 'bold'),
                 padx=8, pady=4).pack(side='left', padx=3)
        
        tk.Button(btn_frame, text="Clear", command=self._clear_list,
                 bg='#9E9E9E', fg='white', font=('Segoe UI', 8),
                 padx=8, pady=4).pack(side='left', padx=3)
        
        # RIGHT: Text preview
        right_frame = tk.Frame(paned)
        paned.add(right_frame, weight=2)
        
        tk.Label(right_frame, text="Extracted Text Preview", 
                font=('Segoe UI', 9, 'bold')).pack(anchor='w', pady=(0, 5))
        
        self.preview_text = scrolledtext.ScrolledText(right_frame, wrap='word',
                                                      font=('Segoe UI', 9),
                                                      height=15)
        self.preview_text.pack(fill='both', expand=True)
        
        # Processing options
        options_frame = tk.LabelFrame(parent, text="Processing Options", 
                                     padx=10, pady=10)
        options_frame.pack(fill='x', padx=5, pady=(0, 10))
        
        # Model selection and formatting option
        model_frame = tk.Frame(options_frame)
        model_frame.pack(fill='x', pady=(0, 5))
        
        tk.Label(model_frame, text="Model:", font=('Segoe UI', 9)).pack(side='left', padx=(0, 5))
        self.model_var = tk.StringVar(value="gpt-4o")
        models = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"]
        ttk.Combobox(model_frame, textvariable=self.model_var, values=models,
                    width=20, state='readonly').pack(side='left')
        
        # Formatting option
        self.preserve_formatting_var = tk.BooleanVar(value=True)
        tk.Checkbutton(model_frame, text="Preserve formatting (bold/italic/underline)",
                      variable=self.preserve_formatting_var,
                      font=('Segoe UI', 9)).pack(side='left', padx=(20, 0))
        
        # Custom instructions
        instructions_header = tk.Frame(options_frame)
        instructions_header.pack(fill='x', pady=(5, 2))
        
        tk.Label(instructions_header, text="Extraction Instructions:", 
                font=('Segoe UI', 9)).pack(side='left')
        
        tk.Button(instructions_header, text="üëÅÔ∏è Show Prompt", 
                 command=self._show_full_prompt,
                 bg='#9C27B0', fg='white', font=('Segoe UI', 8),
                 padx=8, pady=2).pack(side='right')
        
        self.instructions_text = scrolledtext.ScrolledText(options_frame, wrap='word',
                                                          font=('Segoe UI', 9),
                                                          height=3)
        self.instructions_text.pack(fill='x')
        
        default_instructions = """Extract all text from this image. The image is a screenshot from a poorly formatted PDF.
Please:
- Extract all visible text accurately
- Fix any obvious OCR errors or formatting issues
- Remove extraneous line breaks within paragraphs
- Preserve intentional paragraph breaks
- Maintain the logical flow and structure of the content
- For redacted/blacked-out text: insert a descriptive placeholder in square brackets in the document's language (e.g., [naam] for Dutch names, [name] for English names, [bedrag] for amounts, etc.)
- For stamps, signatures, or images: insert a descriptive placeholder in square brackets in the document's language (e.g., [handtekening], [stempel], [signature], [stamp], etc.)
- For any non-text elements that would normally appear: describe them briefly in square brackets
- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__
- Output clean, readable text only (no commentary)"""
        
        self.instructions_text.insert('1.0', default_instructions)
        
        # Action buttons
        action_frame = tk.Frame(parent, bg='white')
        action_frame.pack(fill='x', padx=5, pady=(0, 10))
        
        tk.Button(action_frame, text="üîç Process Selected", 
                 command=self._process_selected,
                 bg='#FF9800', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=(0, 5))
        
        tk.Button(action_frame, text="‚ö° Process ALL", 
                 command=self._process_all,
                 bg='#4CAF50', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        tk.Button(action_frame, text="üíæ Save DOCX", 
                 command=self._save_to_docx,
                 bg='#2196F3', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        tk.Button(action_frame, text="üìã Copy All", 
                 command=self._copy_all_text,
                 bg='#607D8B', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        tk.Button(action_frame, text="üìä Session Report", 
                 command=self._save_session_report,
                 bg='#795548', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=15, pady=6).pack(side='left', padx=5)
        
        # Status
        self.status_label = tk.Label(parent, text="Ready - Add images to begin", 
                                     font=('Segoe UI', 9), fg='#666', bg='white')
        self.status_label.pack(pady=(0, 5))
        
        # Progress bar
        self.progress = ttk.Progressbar(parent, mode='determinate')
        self.progress.pack(fill='x', padx=5, pady=(0, 5))
        
        # Restore state after UI recreation
        self.image_files = saved_files
        self.extracted_texts = saved_texts
        if self.image_files:
            self._update_listbox()
    
    # === File Management Methods ===
    
    def _import_from_pdf(self):
        """Import images directly from a PDF file"""
        pdf_file = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        )
        
        if not pdf_file:
            return
        
        try:
            # Open PDF
            doc = fitz.open(pdf_file)
            total_pages = len(doc)
            
            if total_pages == 0:
                messagebox.showwarning("Empty PDF", "The selected PDF has no pages.")
                return
            
            # Create folder for extracted images next to the PDF
            pdf_path = Path(pdf_file)
            pdf_name = pdf_path.stem
            images_folder = pdf_path.parent / f"{pdf_name}_images"
            
            # Create folder if it doesn't exist
            images_folder.mkdir(exist_ok=True)
            temp_dir = str(images_folder)
            
            # Log start
            if hasattr(self, 'log_message'):
                self.log_message(f"Starting PDF import: {Path(pdf_file).name}")
                self.log_message(f"Total pages: {total_pages}")
            
            # Extract each page as an image
            extracted_count = 0
            self.status_label.config(text=f"Extracting pages from PDF...")
            self.parent_app.root.update_idletasks()
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # Render page to pixmap (image) at 2x resolution for better quality
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                # Save as PNG
                img_filename = f"{pdf_name}_page_{page_num + 1:03d}.png"
                img_path = os.path.join(temp_dir, img_filename)
                pix.save(img_path)
                
                # Add to image list
                if img_path not in self.image_files:
                    self.image_files.append(img_path)
                    extracted_count += 1
                    
                    # Log each page
                    if hasattr(self, 'log_message'):
                        self.log_message(f"  Page {page_num + 1}/{total_pages} extracted: {img_filename}")
                
                # Update progress
                self.status_label.config(
                    text=f"Extracting page {page_num + 1}/{total_pages}..."
                )
                self.parent_app.root.update_idletasks()
            
            doc.close()
            
            # Update UI
            self._update_listbox()
            self.status_label.config(
                text=f"Imported {extracted_count} page(s) from PDF"
            )
            
            # Log completion
            if hasattr(self, 'log_message'):
                self.log_message(f"PDF import complete: {extracted_count} pages extracted")
                self.log_message(f"Temporary folder: {temp_dir}")
            
            messagebox.showinfo(
                "PDF Import Complete",
                f"Successfully extracted {extracted_count} page(s) from:\n{Path(pdf_file).name}\n\n"
                f"Images saved to folder:\n{temp_dir}\n\n"
                f"These images are kept for your reference and can be useful for the end client.\n\n"
                f"You can now process these pages with AI OCR."
            )
            
        except Exception as e:
            messagebox.showerror("PDF Import Error", f"Failed to import PDF:\n{str(e)}")
            self.status_label.config(text="PDF import failed")
    
    def _add_files(self):
        """Add individual image files"""
        files = filedialog.askopenfilenames(
            title="Select Image Files",
            filetypes=[
                ("Image files", "*.jpg *.jpeg *.png *.bmp *.gif *.tiff"),
                ("All files", "*.*")
            ]
        )
        
        if files:
            for file in files:
                if file not in self.image_files:
                    self.image_files.append(file)
            self._update_listbox()
            self.status_label.config(text=f"Added {len(files)} file(s)")
    
    def _add_folder(self):
        """Add all images from a folder"""
        folder = filedialog.askdirectory(title="Select Folder with Images")
        
        if folder:
            image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}
            files = []
            
            for file in sorted(os.listdir(folder)):
                file_path = os.path.join(folder, file)
                if os.path.isfile(file_path):
                    ext = os.path.splitext(file)[1].lower()
                    if ext in image_extensions and file_path not in self.image_files:
                        files.append(file_path)
            
            self.image_files.extend(files)
            self._update_listbox()
            self.status_label.config(text=f"Added {len(files)} file(s) from folder")
    
    def _clear_list(self):
        """Clear all files"""
        if self.image_files and messagebox.askyesno("Clear", "Remove all files?"):
            self.image_files = []
            self.extracted_texts = {}
            self._update_listbox()
            self.preview_text.delete('1.0', tk.END)
            self.status_label.config(text="List cleared")
    
    def _update_listbox(self):
        """Update file listbox"""
        self.file_listbox.delete(0, tk.END)
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            status = "‚úì " if file in self.extracted_texts else ""
            self.file_listbox.insert(tk.END, f"{status}{i:2d}. {filename}")
    
    def _on_file_select(self, event):
        """Show extracted text when file is selected"""
        selection = self.file_listbox.curselection()
        if not selection:
            return
        
        idx = selection[0]
        if idx < len(self.image_files):
            file = self.image_files[idx]
            if file in self.extracted_texts:
                self.preview_text.delete('1.0', tk.END)
                self.preview_text.insert('1.0', self.extracted_texts[file])
    
    def _show_full_prompt(self):
        """Show the exact prompt that will be sent to the AI"""
        instructions = self.instructions_text.get('1.0', 'end-1c').strip()
        
        # Apply formatting modifications like in _extract_text_from_image
        if self.preserve_formatting_var.get():
            if "markdown for text formatting" not in instructions:
                instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
        else:
            instructions = instructions.replace(
                "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
            ).replace(
                "- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
            )
        
        # Create popup window
        popup = tk.Toplevel()
        popup.title("Full Prompt Preview")
        popup.geometry("700x600")
        
        # Main frame
        main_frame = tk.Frame(popup, padx=15, pady=15)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title = tk.Label(main_frame, text="Exact Prompt Sent to OpenAI API",
                        font=('Segoe UI', 12, 'bold'))
        title.pack(pady=(0, 10))
        
        # Info frame
        info_frame = tk.LabelFrame(main_frame, text="Configuration", padx=10, pady=10)
        info_frame.pack(fill='x', pady=(0, 10))
        
        tk.Label(info_frame, text=f"Model: {self.model_var.get()}", 
                font=('Segoe UI', 9, 'bold')).pack(anchor='w')
        
        formatting_status = "‚úì Enabled" if self.preserve_formatting_var.get() else "‚úó Disabled"
        tk.Label(info_frame, text=f"Formatting Preservation: {formatting_status}", 
                font=('Segoe UI', 9)).pack(anchor='w')
        
        tk.Label(info_frame, text=f"Max Tokens: 4000", 
                font=('Segoe UI', 9)).pack(anchor='w')
        
        # Prompt text
        prompt_frame = tk.LabelFrame(main_frame, text="Full Instructions Text", 
                                     padx=10, pady=10)
        prompt_frame.pack(fill='both', expand=True, pady=(0, 10))
        
        prompt_text = scrolledtext.ScrolledText(prompt_frame, wrap='word',
                                               font=('Consolas', 9))
        prompt_text.pack(fill='both', expand=True)
        prompt_text.insert('1.0', instructions)
        prompt_text.config(state='disabled')
        
        # Note
        note = tk.Label(main_frame, 
                       text="Note: The image is sent as base64-encoded data along with these instructions.",
                       font=('Segoe UI', 8), fg='#666')
        note.pack(pady=(0, 5))
        
        # Close button
        tk.Button(main_frame, text="Close", command=popup.destroy,
                 bg='#607D8B', fg='white', font=('Segoe UI', 9, 'bold'),
                 padx=20, pady=6).pack()
    
    def _save_session_report(self):
        """Generate and save a session report in markdown format"""
        if not self.extracted_texts:
            messagebox.showwarning("No Data", "No OCR processing has been performed yet.\n\n"
                                 "Process some images first to generate a session report.")
            return
        
        # Ask for save location
        output_file = filedialog.asksaveasfilename(
            title="Save Session Report",
            defaultextension=".md",
            filetypes=[("Markdown files", "*.md"), ("Text files", "*.txt"), ("All files", "*.*")],
            initialfile="PDF_Rescue_SessionReport.md"
        )
        
        if not output_file:
            return
        
        try:
            from datetime import datetime
            
            # Generate report content
            report_lines = []
            report_lines.append("# PDF Rescue - Session Report\n")
            report_lines.append("**Generated by [Supervertaler](https://supervertaler.com/) ‚Ä¢ by Michael Beijer**\n\n")
            report_lines.append(f"**Date**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            report_lines.append("---\n\n")
            
            # Configuration section
            report_lines.append("## Configuration\n\n")
            report_lines.append(f"- **Model**: {self.model_var.get()}\n")
            formatting_status = "Enabled ‚úì" if self.preserve_formatting_var.get() else "Disabled ‚úó"
            report_lines.append(f"- **Formatting Preservation**: {formatting_status}\n")
            report_lines.append(f"- **Total Images Processed**: {len(self.extracted_texts)}\n")
            report_lines.append(f"- **Total Images in List**: {len(self.image_files)}\n\n")
            
            # Instructions used
            report_lines.append("## Extraction Instructions\n\n")
            report_lines.append("```\n")
            instructions = self.instructions_text.get('1.0', 'end-1c').strip()
            if self.preserve_formatting_var.get():
                if "markdown for text formatting" not in instructions:
                    instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
            report_lines.append(instructions)
            report_lines.append("\n```\n\n")
            
            # Processing summary
            report_lines.append("## Processing Summary\n\n")
            report_lines.append("| # | Image File | Status |\n")
            report_lines.append("|---|------------|--------|\n")
            
            for i, file in enumerate(self.image_files, 1):
                filename = os.path.basename(file)
                status = "‚úì Processed" if file in self.extracted_texts else "‚ßó Pending"
                report_lines.append(f"| {i} | {filename} | {status} |\n")
            
            report_lines.append("\n---\n\n")
            
            # Extracted text for each image
            report_lines.append("## Extracted Text\n\n")
            
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    filename = os.path.basename(file)
                    report_lines.append(f"### Page {i}: {filename}\n\n")
                    report_lines.append("```\n")
                    report_lines.append(self.extracted_texts[file])
                    report_lines.append("\n```\n\n")
                    report_lines.append("---\n\n")
            
            # Statistics
            report_lines.append("## Statistics\n\n")
            total_chars = sum(len(text) for text in self.extracted_texts.values())
            total_words = sum(len(text.split()) for text in self.extracted_texts.values())
            report_lines.append(f"- **Total Characters Extracted**: {total_chars:,}\n")
            report_lines.append(f"- **Total Words Extracted**: {total_words:,}\n")
            report_lines.append(f"- **Average Characters per Page**: {total_chars // len(self.extracted_texts) if self.extracted_texts else 0:,}\n")
            report_lines.append(f"- **Average Words per Page**: {total_words // len(self.extracted_texts) if self.extracted_texts else 0:,}\n\n")
            
            # Footer
            report_lines.append("---\n\n")
            report_lines.append("*Report generated by **PDF Rescue** - AI-Powered OCR Tool*\n\n")
            report_lines.append("*Part of [**Supervertaler**](https://supervertaler.com/) ‚Ä¢ by Michael Beijer*\n")
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.writelines(report_lines)
            
            self.log_message(f"Session report saved: {Path(output_file).name}")
            self.status_label.config(text=f"‚úì Report saved to {os.path.basename(output_file)}")
            
            if messagebox.askyesno("Success", 
                                  f"Session report saved successfully!\n\n"
                                  f"File: {Path(output_file).name}\n\n"
                                  "Open the report now?"):
                os.startfile(output_file)
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save report:\n\n{str(e)}")
    
    # === OCR Processing Methods ===
    
    def _encode_image(self, image_path):
        """Encode image to base64"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def _extract_text_from_image(self, image_path):
        """Use GPT-4 Vision to extract text from image"""
        if not self.client:
            return "[ERROR: OpenAI client not initialized. Check API key.]"
        
        try:
            base64_image = self._encode_image(image_path)
            instructions = self.instructions_text.get('1.0', 'end-1c').strip()
            
            # Add or remove formatting instruction based on checkbox
            if self.preserve_formatting_var.get():
                if "markdown for text formatting" not in instructions:
                    instructions += "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__"
            else:
                # Remove markdown instruction if present
                instructions = instructions.replace(
                    "\n- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
                ).replace(
                    "- Use markdown for text formatting: **bold text**, *italic text*, __underlined text__", ""
                )
            
            response = self.client.chat.completions.create(
                model=self.model_var.get(),
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": instructions
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=4000
            )
            
            return response.choices[0].message.content
        
        except Exception as e:
            return f"[ERROR extracting text: {str(e)}]"
    
    def _process_selected(self):
        """Process currently selected image"""
        selection = self.file_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select an image to process")
            return
        
        idx = selection[0]
        if idx >= len(self.image_files):
            return
        
        file = self.image_files[idx]
        filename = os.path.basename(file)
        
        self.log_message(f"Processing selected image: {filename}")
        self.status_label.config(text=f"Processing {filename}...")
        if hasattr(self.parent_app, 'root'):
            self.parent_app.root.update()
        
        text = self._extract_text_from_image(file)
        self.extracted_texts[file] = text
        
        self.preview_text.delete('1.0', tk.END)
        self.preview_text.insert('1.0', text)
        
        self._update_listbox()
        self.log_message(f"Successfully processed: {filename}")
        self.status_label.config(text=f"‚úì Processed {filename}")
    
    def _process_all(self):
        """Process all images in the list"""
        if not self.image_files:
            messagebox.showwarning("No Files", "Please add images first")
            return
        
        if not messagebox.askyesno("Process All", 
                                   f"Process all {len(self.image_files)} images?\n\n"
                                   "This will use API credits and may take several minutes."):
            return
        
        self.log_message(f"Starting batch processing: {len(self.image_files)} images")
        self.progress['maximum'] = len(self.image_files)
        self.progress['value'] = 0
        
        for i, file in enumerate(self.image_files, 1):
            filename = os.path.basename(file)
            self.status_label.config(text=f"Processing {i}/{len(self.image_files)}: {filename}...")
            if hasattr(self.parent_app, 'root'):
                self.parent_app.root.update()
            
            if file not in self.extracted_texts:
                text = self._extract_text_from_image(file)
                self.extracted_texts[file] = text
                self.log_message(f"  [{i}/{len(self.image_files)}] Processed: {filename}")
            else:
                self.log_message(f"  [{i}/{len(self.image_files)}] Skipped (already processed): {filename}")
            
            self.progress['value'] = i
            self._update_listbox()
        
        self.log_message(f"Batch processing complete: {len(self.image_files)} images processed")
        self.status_label.config(text=f"‚úì Processed all {len(self.image_files)} images!")
        messagebox.showinfo("Complete", 
                          f"Successfully processed {len(self.image_files)} images!\n\n"
                          "Click 'Save DOCX' to export the text.")
    
    # === Export Methods ===
    
    def _add_formatted_text(self, doc, text):
        """
        Add text to document with markdown formatting parsed
        Supports: **bold**, *italic*, __underline__
        """
        # Split text into paragraphs
        paragraphs = text.split('\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.15
            para.paragraph_format.space_after = Pt(12)
            
            # Parse markdown formatting using regex
            # Pattern matches: **bold**, *italic*, __underline__
            # We need to handle nested/overlapping formatting carefully
            position = 0
            
            # Combined pattern to find all formatting markers
            pattern = r'(\*\*.*?\*\*|\*.*?\*|__.*?__|.+?(?=\*\*|\*|__|$)|.)'
            
            # Simple approach: process sequentially
            remaining = para_text
            
            while remaining:
                # Check for bold (**text**)
                bold_match = re.match(r'\*\*(.*?)\*\*', remaining)
                if bold_match:
                    run = para.add_run(bold_match.group(1))
                    run.bold = True
                    remaining = remaining[bold_match.end():]
                    continue
                
                # Check for underline (__text__)
                underline_match = re.match(r'__(.*?)__', remaining)
                if underline_match:
                    run = para.add_run(underline_match.group(1))
                    run.underline = True
                    remaining = remaining[underline_match.end():]
                    continue
                
                # Check for italic (*text*)
                italic_match = re.match(r'\*(.*?)\*', remaining)
                if italic_match:
                    run = para.add_run(italic_match.group(1))
                    run.italic = True
                    remaining = remaining[italic_match.end():]
                    continue
                
                # No formatting - add plain text until next marker or end
                next_marker = len(remaining)
                for marker in ['**', '*', '__']:
                    pos = remaining.find(marker)
                    if pos != -1 and pos < next_marker:
                        next_marker = pos
                
                if next_marker == 0:
                    # Edge case: marker at start but no match (e.g., single * or **)
                    para.add_run(remaining[0])
                    remaining = remaining[1:]
                else:
                    plain_text = remaining[:next_marker] if next_marker < len(remaining) else remaining
                    if plain_text:
                        para.add_run(plain_text)
                    remaining = remaining[next_marker:]
    
    def _save_to_docx(self):
        """Save all extracted text to a Word document"""
        if not self.extracted_texts:
            messagebox.showwarning("No Text", "No extracted text to save.\n\n"
                                 "Process images first.")
            return
        
        output_file = filedialog.asksaveasfilename(
            title="Save Extracted Text",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            initialfile="extracted_text.docx"
        )
        
        if not output_file:
            return
        
        self.log_message(f"Saving extracted text to DOCX: {Path(output_file).name}")
        
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Extracted Text from Images', 0)
            title.runs[0].font.size = Pt(16)
            
            # Add extracted text in order
            for i, file in enumerate(self.image_files, 1):
                if file in self.extracted_texts:
                    # Page header
                    heading = doc.add_heading(f'Page {i}: {os.path.basename(file)}', level=2)
                    heading.runs[0].font.size = Pt(12)
                    
                    # Text content with formatting
                    text = self.extracted_texts[file]
                    if self.preserve_formatting_var.get():
                        self._add_formatted_text(doc, text)
                    else:
                        para = doc.add_paragraph(text)
                        para.paragraph_format.line_spacing = 1.15
                        para.paragraph_format.space_after = Pt(12)
                    
                    # Page break except for last
                    if i < len(self.image_files):
                        doc.add_page_break()
            
            doc.save(output_file)
            
            self.log_message(f"Successfully saved {len(self.extracted_texts)} pages to: {Path(output_file).name}")
            self.status_label.config(text=f"‚úì Saved to {os.path.basename(output_file)}")
            
            if messagebox.askyesno("Success", 
                                  f"Document saved successfully!\n\n"
                                  f"{len(self.extracted_texts)} pages of text extracted\n\n"
                                  "Open the document now?"):
                os.startfile(output_file)
        
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save document:\n\n{str(e)}")
    
    def _copy_all_text(self):
        """Copy all extracted text to clipboard"""
        if not self.extracted_texts:
            messagebox.showwarning("No Text", "No extracted text to copy")
            return
        
        all_text = []
        for i, file in enumerate(self.image_files, 1):
            if file in self.extracted_texts:
                all_text.append(f"=== Page {i}: {os.path.basename(file)} ===\n")
                all_text.append(self.extracted_texts[file])
                all_text.append("\n\n")
        
        combined = "".join(all_text)
        
        # Get root window from parent app or use clipboard differently
        if hasattr(self.parent_app, 'root'):
            self.parent_app.root.clipboard_clear()
            self.parent_app.root.clipboard_append(combined)
        
        self.status_label.config(text=f"‚úì Copied {len(self.extracted_texts)} pages to clipboard")
        messagebox.showinfo("Copied", f"Copied text from {len(self.extracted_texts)} pages to clipboard!")


# === Standalone Application ===

if __name__ == "__main__":
    """Run PDF Rescue as a standalone application"""
    
    class StandaloneApp:
        """Minimal parent app for standalone mode"""
        def __init__(self):
            self.root = tk.Tk()
            self.root.title("PDF Rescue - AI-Powered OCR Tool")
            self.root.geometry("1000x700")
            
            # Load API key from api_keys.txt
            self.api_keys = {}
            api_file = Path("api_keys.txt")
            if api_file.exists():
                with open(api_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#') and '=' in line:
                            key, value = line.split('=', 1)
                            if 'openai' in key.lower():
                                self.api_keys['openai'] = value.strip()
            
            if not self.api_keys.get('openai'):
                messagebox.showerror(
                    "API Key Missing",
                    "Could not find OpenAI API key in api_keys.txt\n\n"
                    "Please add a line like:\nOPENAI_API_KEY=your-key-here"
                )
                self.root.destroy()
                return
            
            # Create main container
            main_frame = tk.Frame(self.root)
            main_frame.pack(fill='both', expand=True, padx=10, pady=10)
            
            # Add title
            title = tk.Label(main_frame, text="PDF Rescue - AI-Powered OCR Tool",
                           font=('Segoe UI', 14, 'bold'))
            title.pack(pady=(0, 10))
            
            # Create PDF Rescue instance
            self.pdf_rescue = PDFRescue(self)
            self.pdf_rescue.create_tab(main_frame)
            
            # Add log at bottom
            log_frame = tk.LabelFrame(self.root, text="Activity Log", padx=5, pady=5)
            log_frame.pack(fill='x', padx=10, pady=(0, 10))
            
            self.log_text = scrolledtext.ScrolledText(log_frame, height=4, wrap='word',
                                                     font=('Consolas', 9))
            self.log_text.pack(fill='both', expand=True)
            self.log_text.config(state='disabled')
        
        def log(self, message: str):
            """Add message to log"""
            from datetime import datetime
            timestamp = datetime.now().strftime("%H:%M:%S")
            formatted_message = f"[{timestamp}] {message}\n"
            
            self.log_text.config(state='normal')
            self.log_text.insert('end', formatted_message)
            self.log_text.see('end')
            self.log_text.config(state='disabled')
        
        def run(self):
            """Start the application"""
            self.root.mainloop()
    
    # Create and run standalone app
    app = StandaloneApp()
    if hasattr(app, 'root') and app.root.winfo_exists():
        app.run()
