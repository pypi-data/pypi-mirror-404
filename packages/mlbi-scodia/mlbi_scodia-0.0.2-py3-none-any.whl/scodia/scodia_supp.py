from __future__ import annotations

import re, json
from dataclasses import is_dataclass, asdict
from typing import Any, Dict, List, Tuple
from IPython.display import display, Markdown, HTML
import pandas as pd

###########################################
### Display functions
###########################################

def display_img_base64( imgs, html = False ):
    if isinstance(imgs, str):
        if html:
            html_string = f'<img src="data:image/png;base64,{imgs}">'
            display(HTML(html_string))
        else:
            display(Image(data=base64.b64decode(imgs)))
    elif isinstance(imgs, list):
        for i in imgs:
            if html:
                html_string = f'<img src="data:image/png;base64,{i}">'
                display(HTML(html_string))
            else:
                display(Image(data=base64.b64decode(i)))
    elif isinstance(imgs, dict):
        for k in imgs.keys():
            if html:
                html_string = f'<img src="data:image/png;base64,{imgs[k]}">'
                display(HTML(html_string))
            else:
                display(Image(data=base64.b64decode(imgs[k])))
    return


def _val_to_str(v: Any, maxlen: int = 140) -> str:
    """Value를 사람이 보기 좋게 문자열화."""
    if v is None:
        s = "None"
    elif is_dataclass(v):
        s = str(v)  # UMAPEmbedConfig(...)처럼 보이게
    elif isinstance(v, (list, tuple, set)):
        s = str(list(v))
    elif isinstance(v, dict):
        # 너무 길어질 수 있어서 요약
        s = str(v)
    else:
        s = str(v)

    if len(s) > maxlen:
        s = s[: maxlen - 3] + "..."
    return s


def _flatten_args(
    args: Dict[str, Any],
    *,
    category_root: str = "General",
    sep: str = ".",
) -> List[Tuple[str, str, Any]]:
    """
    args를 (Category, Parameter, Value) 튜플 리스트로 flatten.
    - dict는 category를 하위로 내림
    - dataclass는 dict(asdict)로 풀어서 하위로 내림
    """
    rows: List[Tuple[str, str, Any]] = []

    def rec(obj: Any, cat: str, prefix: str = ""):
        if is_dataclass(obj):
            obj = asdict(obj)

        if isinstance(obj, dict):
            for k, v in obj.items():
                if is_dataclass(v) or isinstance(v, dict):
                    # 하위 category로 분리
                    sub_cat = k if cat == category_root and prefix == "" else f"{cat}{sep}{k}"
                    rec(v, sub_cat, "")
                else:
                    pname = k if prefix == "" else f"{prefix}{sep}{k}"
                    rows.append((cat, pname, v))
        else:
            # dict가 아닌데 여기까지 온 케이스는 거의 없음
            rows.append((cat, prefix or "value", obj))

    # 최상위 키는 General에 두되, dict/dataclass면 category로 내려보냄
    for k, v in (args or {}).items():
        if is_dataclass(v) or isinstance(v, dict):
            rec(v, k, "")  # embed_cfg / plot_cfg 같은 건 category로
        else:
            rows.append((category_root, k, v))

    return rows


def display_analysis_parameters(
    tool_name: str,
    full_args: Dict[str, Any],
    *,
    title: str | None = None,
    max_value_len: int = 160,
    show: bool = True,
):
    """
    Jupyter-friendly pretty table renderer for tool parameters.

    Parameters
    ----------
    tool_name : str
        Tool name, e.g., "plot_umap"
    full_args : dict
        Materialized full args (defaults + AI patch merged).
    title : str, optional
        Title shown above the table.
    max_value_len : int
        Truncate long values.
    show : bool
        If True, display in Jupyter (HTML). If False, return DataFrame/Styler.
    """
    
    rows = _flatten_args(full_args, category_root="General")
    df = pd.DataFrame(rows, columns=["Category", "Parameter", "Value"])
    df["Value"] = df["Value"].map(lambda x: _val_to_str(x, maxlen=max_value_len))

    # 정렬: General 먼저, 그 다음 category 알파벳
    df["__cat_rank"] = (df["Category"] != "General").astype(int)
    df = df.sort_values(["__cat_rank", "Category", "Parameter"]).drop(columns="__cat_rank")

    # Jupyter에서 예쁘게
    styler = (
        df.style
          .hide(axis="index")
          .set_caption(title or f"{tool_name} parameters")
          .set_table_styles([
              {"selector": "caption", "props": [("font-size", "16px"),
                                                ("font-weight", "600"),
                                                ("text-align", "left"),
                                                ("margin-bottom", "2px")]},
              {"selector": "th", "props": [("font-size", "14px"),
                                           ("text-align", "left"),
                                           ("border-bottom", "1px solid #ddd")]},
              {"selector": "td", "props": [("font-size", "14px"),
                                           ("border-bottom", "1px solid #f0f0f0"),
                                           ("vertical-align", "top")]},
          ])
          .set_properties(subset=["Category"], **{"font-weight": "600"})
    )

    if show:
        try:
            from IPython.display import display
            display(styler)
        except Exception:
            # 터미널 등에서는 그냥 DF 출력
            print(title or f"{tool_name} parameters")
            display(df) if "display" in globals() else print(df.to_string(index=False))
        return None

    return df, styler



def render_scoda_info_260128(json_input, lang: str = 'ko'):
    """
    JSON 데이터(dict 또는 str)를 받아서 주피터 노트북에 세련된 리포트로 출력합니다.
    """
    # 1. 입력 데이터 타입 정규화 (str -> dict)
    if isinstance(json_input, str):
        try:
            data = json.loads(json_input)
        except json.JSONDecodeError:
            # 혹시 싱글 쿼테이션(') 등이 섞인 불완전한 문자열일 경우를 대비
            try:
                import ast
                data = ast.literal_eval(json_input)
            except:
                return print("❌ 데이터 형식에 오류가 있어 리포트를 생성할 수 없습니다.")
    elif isinstance(json_input, dict):
        data = json_input
    else:
        return print(f"❌ 지원하지 않는 데이터 타입입니다: {type(json_input)}")

    # 2. 리포트 생성 (Markdown)
    report = []
    report.append("# 📊 SCODA Dataset Insight Report")
    report.append("> 해당 데이터셋에 대해 AI가 분석한 요약 및 활용 가이드입니다.\n")
    
    # 요약 섹션
    if 'summary' in data:
        if lang == 'ko':
            report.append("### 💡 데이터셋 요약")
        else:
            report.append("### 💡 Dataset summary")
        for item in data['summary']:
            report.append(f"- {item}")
        report.append("\n" + "---" * 10)

    # 분석 도구 섹션 (Table 활용)
    if 'functions' in data:
        if lang == 'ko':
            report.append("### 🛠️ 제공되는 분석 함수")
            report.append("| 분석 도구 | 상세 설명 |")
            report.append("| :--- | :--- |")
        else:
            report.append("### 🛠️ Functions/Tools provided")
            report.append("| Tool | Description |")
            report.append("| :--- | :--- |")
            
        for func in data['functions']:
            if ":" in func:
                name, desc = func.split(":", 1)
                report.append(f"| **{name.strip()}** | {desc.strip()} |")
            else:
                report.append(f"| - | {func} |")
        report.append("\n" + "---" * 10)

    # 분석 역량 섹션
    if 'capabilities' in data:
        if lang == 'ko':
            report.append("### 🚀 분석 가능 범위")
        else:
            report.append("### 🚀 Capabilities")
            
        for cap in data['capabilities']:
            report.append(f"- ✅ {cap}")
        report.append("\n" + "---" * 10)

    # 질문 예시 섹션
    eq = 'example questions'
    if ('example_questions' in data):
        eq = 'example_questions'
    
    if (eq in data):
        if lang == 'ko':
            report.append("### ❓ 질문 예시 (Copy & Paste)")
            report.append("AI에게 아래와 같이 질문하여 즉시 분석을 시작할 수 있습니다.\n")
        else:
            report.append("### ❓ Example Questions (Copy & Paste)")
            report.append("You can start your analysis immediately by asking the AI the following questions.\n")
            
        for q in data[eq]:
            # 말풍선 느낌의 스타일링
            # report.append(f"💬 `\"{q}\"`\n")
            # clean_q = q.strip()
            # report.append(f"💬 `\"{clean_q}\"`")
            report.append(f"- {q}")
        report.append("\n" + "---" * 10)
    
    # 3. 최종 출력
    display(Markdown("\n".join(report)))


def render_scoda_info(input_data, lang: str = 'ko'):
    """
    Render dataset info.
    - dict: structured JSON report
    - str (JSON-like): parse & render
    - str (plain text): render as narrative report
    """
    # 1. Normalize input
    if isinstance(input_data, dict):
        data = input_data
        mode = "json"

    elif isinstance(input_data, str):
        s = input_data.strip()

        # JSON-like string
        if s.startswith("{") and s.endswith("}"):
            try:
                data = json.loads(s)
                mode = "json"
            except Exception:
                mode = "text"
        else:
            mode = "text"

    else:
        print(f"❌ Unsupported data type: {type(input_data)}")
        return

    # 2. Render
    if mode == "text":
        title = "# 📊 SCODA Dataset Insight"
        subtitle = (
            "> AI-generated overview of the dataset and its capabilities.\n"
            if lang == "en"
            else "> 해당 데이터셋에 대한 AI 기반 요약 및 활용 가이드입니다.\n"
        )
        display(Markdown("\n".join([title, subtitle, input_data])))
        return

    # ---- 기존 JSON 기반 렌더링 (거의 그대로 유지) ----
    report = []
    report.append("# 📊 SCODA Dataset Insight Report")
    report.append("> 해당 데이터셋에 대해 AI가 분석한 요약 및 활용 가이드입니다.\n")

    if 'summary' in data:
        report.append("### 💡 Dataset Summary" if lang == "en" else "### 💡 데이터셋 요약")
        for item in data['summary']:
            report.append(f"- {item}")
        report.append("\n" + "---" * 10)

    if 'functions' in data:
        report.append("### 🛠️ Functions / Tools" if lang == "en" else "### 🛠️ 제공되는 분석 함수")
        for func in data['functions']:
            report.append(f"- {func}")
        report.append("\n" + "---" * 10)

    if 'capabilities' in data:
        report.append("### 🚀 Capabilities" if lang == "en" else "### 🚀 분석 가능 범위")
        for cap in data['capabilities']:
            report.append(f"- ✅ {cap}")
        report.append("\n" + "---" * 10)

    if 'example_questions' in data:
        report.append("### ❓ Example Questions" if lang == "en" else "### ❓ 질문 예시")
        for q in data['example_questions']:
            report.append(f"- {q}")
        report.append("\n" + "---" * 10)

    display(Markdown("\n".join(report)))


def render_scoda_explanation(ai_text):
    """
    HTML을 완전히 제거하고 순수 마크다운 문법만 사용하여 
    출력 오류를 완벽하게 방지합니다.
    """
    from IPython.display import display, Markdown
    
    # 1. 뱃지 대신 마크다운 강조와 구분선을 이용한 헤더
    # --- 는 가로 구분선을 만들어 레이아웃을 분리해줍니다.
    header_md = "### 💡 **Scoda AI Analysis Report**\n---"
    
    # 2. 본문은 인용구(>) 문법을 사용하여 왼쪽 테두리 효과를 줍니다.
    # 각 줄의 시작에 > 를 추가합니다.
    quoted_text = "\n".join([f"> {line}" if line.strip() else ">" for line in ai_text.strip().split("\n")])
    
    # 전체를 하나의 마크다운으로 합쳐서 출력
    full_markdown = f"{header_md}\n\n{quoted_text}"
    
    display(Markdown(full_markdown))
    return full_markdown



def render_gemini_questions(text_data) -> None:
    """
    HTML 태그를 완전히 제거하고 순수 마크다운 문법만 사용하여
    Gemini의 텍스트를 Jupyter 환경에서 안정적이고 세련되게 렌더링합니다.
    """
    from IPython.display import display, Markdown
    
    if text_data is None:
        return

    # 문자열인 경우 줄바꿈 기준으로 리스트화
    if isinstance(text_data, str):
        # 빈 줄은 필터링하여 깔끔하게 정리
        text_data = [line.strip() for line in text_data.split('\n') if line.strip()]
    
    if len(text_data) == 0:
        return

    # 1. 헤더 구성 (HTML 뱃지 대신 마크다운 강조와 아이콘 사용)
    # ### 는 적절한 크기의 제목을 만들고, 그 아래 구분선(---)을 배치합니다.
    header_md = "### 🤖 **Scoda AI Assistant Question**\n---"
    
    # 2. 본문 구성 (인용구 '>' 문법 활용)
    formatted_body = ""
    if isinstance(text_data, list):
        for i, q in enumerate(text_data):
            # 질문 번호를 굵게 표시하고 인용구 처리
            formatted_body += f"> **Q{i+1}.** {q}\n>\n"
    else:
        formatted_body = f"> {text_data}"

    # 3. 전체 마크다운 결합 및 출력
    # 마지막에 구분선을 한 번 더 넣어 영역을 확실히 구분합니다.
    full_markdown = f"{header_md}\n{formatted_body}\n---"
    
    display(Markdown(full_markdown))


###########################################
### Document Generation
###########################################

import io, re
import base64
import docx # 구분선 XML 처리를 위해 필요
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from IPython.display import display, HTML

def add_styled_text(paragraph, text, font_name):
    """
    문장 내의 **텍스트**를 찾아 Bold 처리하고 기호는 삭제하는 보조 함수
    """
    # **텍스트** 패턴 찾기
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # ** 제거하고 진하게 설정
            run = paragraph.add_run(part.replace('**', ''))
            run.bold = True
        else:
            # 일반 텍스트
            run = paragraph.add_run(part)
        
        # 폰트 설정 유지
        run.font.name = font_name
        run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font_name)


def create_doc_with_style(ai_text, img_base64=None):
    doc = Document()
    font_name = 'Malgun Gothic'

    # --- 1. 헤더 (Scoda AI Report Title) ---
    heading = doc.add_heading("💡 Scoda AI Analysis Report", level=1)
    for run in heading.runs:
        run.font.name = font_name
        run.font.color.rgb = RGBColor(44, 62, 80) # 진한 남색 계열
        run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), font_name)

    # --- 2. 가로 구분선 추가 (Markdown의 --- 효과) ---
    p = doc.add_paragraph()
    p_border = p._element.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 선 두께
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D3D3D3') # 연한 회색
    pb.append(bottom)
    p_border.append(pb)

    # --- 3. 이미지 삽입 (제목 바로 아래) ---
    if img_base64:
        try:
            raw_base64 = img_base64.split(",")[-1]
            img_bytes = base64.b64decode(raw_base64)
            img_buffer = io.BytesIO(img_bytes)
            
            # 이미지 너비 설정 (Inches 단위)
            doc.add_picture(img_buffer, width=Inches(5.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = 1 # Center
            
            # 캡션 추가
            caption = doc.add_paragraph("[Analysis Visualization Results]")
            caption.alignment = 1
            doc.add_paragraph() # 여백
        except Exception as e:
            print(f"이미지 추가 중 오류: {e}")

    # --- 4. 본문 내용 (인용구 스타일 및 폰트 적용) ---
    lines = ai_text.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        
        if stripped.startswith('>'):
            clean_text = stripped.replace('>', '', 1).strip()
            p = doc.add_paragraph(style='Quote')
            add_styled_text(p, clean_text, font_name) # 스타일 적용 함수 호출
        elif stripped.startswith('#'):
            level = stripped.count('#')
            p = doc.add_heading('', level=min(level, 9)) # 빈 제목 생성 후
            add_styled_text(p, stripped.replace('#', '').strip(), font_name) # 내용 추가
        else:
            p = doc.add_paragraph()
            add_styled_text(p, stripped, font_name)

    # --- 5. 버퍼로 저장 및 반환 (요청하신 부분) ---
    final_buffer = io.BytesIO()
    doc.save(final_buffer)
    final_buffer.seek(0)
    
    return final_buffer

def creat_download_link(buffer, filename):
    b64 = base64.b64encode(buffer.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" \
             download="{filename}">📂 여기를 클릭하여 Word 파일 다운로드</a>'



###########################################
### Code snippet generation
###########################################

import re, json
from typing import Any, Dict, Optional
from dataclasses import is_dataclass, fields

def _py_literal_pretty(x: Any, indent: int = 0) -> str:
    """Pretty Python literal serializer with indentation."""
    sp = " " * indent

    if is_dataclass(x):
        cls = x.__class__.__name__
        items = []
        for f in fields(x):
            v = getattr(x, f.name)
            items.append(f"{f.name}={_py_literal_pretty(v, indent + 4)}")
        inner = ",\n".join(" " * (indent + 4) + it for it in items)
        return f"{cls}(\n{inner}\n{sp})"

    if isinstance(x, dict):
        if not x:
            return "{}"
        items = []
        for k, v in x.items():
            items.append(
                f"{_py_literal_pretty(k)}: {_py_literal_pretty(v, indent + 4)}"
            )
        inner = ",\n".join(" " * (indent + 4) + it for it in items)
        return "{\n" + inner + "\n" + sp + "}"

    if isinstance(x, (list, tuple)):
        if not x:
            return "[]" if isinstance(x, list) else "()"
        inner = ",\n".join(
            " " * (indent + 4) + _py_literal_pretty(v, indent + 4) for v in x
        )
        if isinstance(x, tuple):
            return "(\n" + inner + ("\n" + sp + ")" if len(x) > 1 else ",\n" + sp + ")")
        return "[\n" + inner + "\n" + sp + "]"

    if isinstance(x, str):
        return json.dumps(x, ensure_ascii=False)
    if x is None:
        return "None"
    if isinstance(x, bool):
        return "True" if x else "False"
    if isinstance(x, (int, float)):
        return repr(x)

    return repr(x)


def build_repro_code_snippet_clean(
    func_name: str,
    exec_args: Dict[str, Any],
    *,
    adata_var_name: str = "adata",
    result_var_name: str = "result",
    assume_imported: bool = True,
) -> str:
    """
    Generate clean, runnable Python code (no markdown fences).
    Assumes plot function and config dataclasses are already imported.
    """

    args_literal = _py_literal_pretty(exec_args, indent=4)

    header = ""
    if assume_imported:
        header = (
            "# NOTE: Assumes the plot function and its config dataclasses\n"
            "# are already imported in this notebook.\n\n"
        )

    code = (
        f"{header}"
        f"{result_var_name} = {func_name}(\n"
        f"    {adata_var_name},\n"
        f"    **{args_literal}\n"
        f")"
    )
    return code


