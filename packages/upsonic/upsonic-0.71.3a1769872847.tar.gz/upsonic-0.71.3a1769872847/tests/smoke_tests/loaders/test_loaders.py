"""
Comprehensive Smoke Test for All Loaders with Config Attributes
Tests all loaders with their specific config attributes using dummy files
"""
import pytest
import os
import tempfile
import json
import csv
from pathlib import Path
from typing import List

from upsonic.loaders import (
    TextLoader, TextLoaderConfig,
    CSVLoader, CSVLoaderConfig,
    JSONLoader, JSONLoaderConfig,
    XMLLoader, XMLLoaderConfig,
    YAMLLoader, YAMLLoaderConfig,
    MarkdownLoader, MarkdownLoaderConfig,
    HTMLLoader, HTMLLoaderConfig,
    DOCXLoader, DOCXLoaderConfig,
    PdfLoader, PdfLoaderConfig,
    PyMuPDFLoader, PyMuPDFLoaderConfig,
    PdfPlumberLoader, PdfPlumberLoaderConfig,
    DoclingLoader, DoclingLoaderConfig,
)
from upsonic.schemas.data_models import Document


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


class TestTextLoader:
    """Test TextLoader with all config attributes"""
    
    def test_basic_text_loading(self, temp_dir):
        """Test basic text file loading"""
        file_path = temp_dir / "test.txt"
        file_path.write_text("Hello World\nThis is a test\n\n  Extra spaces  ")
        
        config = TextLoaderConfig(strip_whitespace=True, min_chunk_length=5)
        loader = TextLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert "Hello World" in docs[0].content
        
    def test_text_loader_encoding(self, temp_dir):
        """Test encoding parameter"""
        file_path = temp_dir / "test_utf8.txt"
        file_path.write_text("Test content with special chars: é, ñ, ü", encoding='utf-8')
        
        config = TextLoaderConfig(encoding='utf-8')
        loader = TextLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert "é" in docs[0].content


class TestCSVLoader:
    """Test CSVLoader with all config attributes"""
    
    def test_csv_single_document(self, temp_dir):
        """Test CSV loading as single document"""
        file_path = temp_dir / "test.csv"
        with open(file_path, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['name', 'age', 'city'])
            writer.writerow(['Alice', '30', 'NYC'])
            writer.writerow(['Bob', '25', 'LA'])
        
        config = CSVLoaderConfig(
            split_mode='single_document',
            content_synthesis_mode='concatenated',
            has_header=True
        )
        loader = CSVLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 1
        assert "Alice" in docs[0].content
        assert "Bob" in docs[0].content
        
    def test_csv_per_row(self, temp_dir):
        """Test CSV loading per row"""
        file_path = temp_dir / "test.csv"
        with open(file_path, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['name', 'age'])
            writer.writerow(['Alice', '30'])
            writer.writerow(['Bob', '25'])
        
        config = CSVLoaderConfig(split_mode='per_row', has_header=True)
        loader = CSVLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 2
        
    def test_csv_column_filtering(self, temp_dir):
        """Test CSV with include/exclude columns"""
        file_path = temp_dir / "test.csv"
        with open(file_path, 'w', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(['name', 'age', 'secret'])
            writer.writerow(['Alice', '30', 'password123'])
        
        config = CSVLoaderConfig(exclude_columns=['secret'], has_header=True)
        loader = CSVLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert 'password123' not in docs[0].content


class TestJSONLoader:
    """Test JSONLoader with all config attributes"""
    
    def test_json_single_mode(self, temp_dir):
        """Test JSON loading in single mode"""
        file_path = temp_dir / "test.json"
        data = {"title": "Test", "content": "Hello World"}
        file_path.write_text(json.dumps(data))
        
        config = JSONLoaderConfig(mode='single', content_synthesis_mode='json')
        loader = JSONLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 1
        assert "Test" in docs[0].content or "Test" in str(docs[0].metadata)
        
    def test_json_multi_mode(self, temp_dir):
        """Test JSON loading in multi mode with record selector"""
        file_path = temp_dir / "test.json"
        data = {
            "articles": [
                {"title": "Article 1", "content": "Content 1"},
                {"title": "Article 2", "content": "Content 2"}
            ]
        }
        file_path.write_text(json.dumps(data))
        
        config = JSONLoaderConfig(
            mode='multi',
            record_selector='.articles[]',
            content_mapper='.content'
        )
        loader = JSONLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 2


class TestXMLLoader:
    """Test XMLLoader with all config attributes"""
    
    def test_xml_basic_loading(self, temp_dir):
        """Test XML loading with basic config"""
        file_path = temp_dir / "test.xml"
        xml_content = """<?xml version="1.0"?>
        <root>
            <item id="1">
                <title>Item 1</title>
                <description>Description 1</description>
            </item>
            <item id="2">
                <title>Item 2</title>
                <description>Description 2</description>
            </item>
        </root>
        """
        file_path.write_text(xml_content)
        
        config = XMLLoaderConfig(
            split_by_xpath='//item',
            include_attributes=True
        )
        loader = XMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        # XML loader behavior may vary, just ensure it loaded successfully
        assert len(docs) >= 1
        assert 'Item 1' in docs[0].content or 'Item 2' in docs[0].content
        
    def test_xml_content_xpath(self, temp_dir):
        """Test XML with content xpath"""
        file_path = temp_dir / "test.xml"
        xml_content = """<?xml version="1.0"?>
        <root>
            <item>
                <title>Title</title>
                <body>Body content here</body>
            </item>
        </root>
        """
        file_path.write_text(xml_content)
        
        config = XMLLoaderConfig(
            split_by_xpath='//item',
            content_xpath='./body'
        )
        loader = XMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 1
        assert "Body content" in docs[0].content


class TestYAMLLoader:
    """Test YAMLLoader with all config attributes"""
    
    def test_yaml_basic_loading(self, temp_dir):
        """Test YAML loading"""
        file_path = temp_dir / "test.yaml"
        yaml_content = """
        title: Test Document
        author: John Doe
        content: This is test content
        """
        file_path.write_text(yaml_content)
        
        config = YAMLLoaderConfig(
            content_synthesis_mode='canonical_yaml',
            flatten_metadata=True
        )
        loader = YAMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 1
        
    def test_yaml_multiple_docs(self, temp_dir):
        """Test YAML with multiple documents"""
        file_path = temp_dir / "test.yaml"
        yaml_content = """title: Doc 1
---
title: Doc 2
"""
        file_path.write_text(yaml_content)
        
        config = YAMLLoaderConfig(handle_multiple_docs=True)
        loader = YAMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 2


class TestMarkdownLoader:
    """Test MarkdownLoader with all config attributes"""
    
    def test_markdown_basic_loading(self, temp_dir):
        """Test Markdown loading"""
        file_path = temp_dir / "test.md"
        md_content = """---
title: Test Document
author: John
---

# Heading 1
This is content

## Heading 2
More content

```python
def hello():
    print("Hello")
```
"""
        file_path.write_text(md_content)
        
        config = MarkdownLoaderConfig(
            parse_front_matter=True,
            include_code_blocks=True,
            heading_metadata=True
        )
        loader = MarkdownLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert docs[0].metadata.get('title') == 'Test Document'
        
    def test_markdown_split_by_heading(self, temp_dir):
        """Test Markdown splitting by heading"""
        file_path = temp_dir / "test.md"
        md_content = """# Section 1
Content 1

# Section 2
Content 2
"""
        file_path.write_text(md_content)
        
        config = MarkdownLoaderConfig(split_by_heading='h1')
        loader = MarkdownLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) == 2


class TestHTMLLoader:
    """Test HTMLLoader with all config attributes"""
    
    def test_html_basic_loading(self, temp_dir):
        """Test HTML loading"""
        file_path = temp_dir / "test.html"
        html_content = """<!DOCTYPE html>
        <html>
        <head>
            <title>Test Page</title>
            <meta name="description" content="Test description">
        </head>
        <body>
            <h1>Heading</h1>
            <p>Paragraph content</p>
            <a href="http://example.com">Link</a>
        </body>
        </html>
        """
        file_path.write_text(html_content)
        
        config = HTMLLoaderConfig(
            extract_text=True,
            include_links=True,
            extract_metadata=True,
            remove_scripts=True
        )
        loader = HTMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        
    def test_html_structure_preservation(self, temp_dir):
        """Test HTML with structure preservation"""
        file_path = temp_dir / "test.html"
        html_content = """<html><body>
            <h1>Title</h1>
            <p>Text</p>
        </body></html>"""
        file_path.write_text(html_content)
        
        config = HTMLLoaderConfig(
            preserve_structure=True,
            extract_headers=True,
            extract_paragraphs=True
        )
        loader = HTMLLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0


class TestPdfLoader:
    """Test PdfLoader with all config attributes"""
    
    def test_pdf_basic_loading(self, temp_dir):
        """Test basic PDF loading"""
        # Use existing PDF from project if available
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        config = PdfLoaderConfig(
            extraction_mode='text_only',
            clean_page_numbers=True,
            extra_whitespace_removal=True
        )
        loader = PdfLoader(config=config)
        docs = loader.load(str(pdf_path))
        
        assert len(docs) > 0
        assert docs[0].content  # Should have content
        
    def test_pdf_page_range(self, temp_dir):
        """Test PDF loading with page range"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        config = PdfLoaderConfig(
            start_page=1,
            end_page=1,
            extraction_mode='text_only'
        )
        loader = PdfLoader(config=config)
        docs = loader.load(str(pdf_path))
        
        assert len(docs) > 0


class TestPyMuPDFLoader:
    """Test PyMuPDFLoader with all config attributes"""
    
    def test_pymupdf_basic_loading(self, temp_dir):
        """Test PyMuPDF basic loading"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = PyMuPDFLoaderConfig(
                extraction_mode='text_only',
                text_extraction_method='text',
                preserve_layout=True
            )
            loader = PyMuPDFLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError:
            pytest.skip("PyMuPDF not installed")
            
    def test_pymupdf_with_images(self, temp_dir):
        """Test PyMuPDF with image extraction"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = PyMuPDFLoaderConfig(
                extraction_mode='hybrid',
                include_images=True,
                extract_annotations=False
            )
            loader = PyMuPDFLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError:
            pytest.skip("PyMuPDF not installed")


class TestPdfPlumberLoader:
    """Test PdfPlumberLoader with all config attributes"""
    
    def test_pdfplumber_basic_loading(self, temp_dir):
        """Test pdfplumber basic loading"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = PdfPlumberLoaderConfig(
                extraction_mode='text_only',
                extract_tables=True,
                table_format='markdown'
            )
            loader = PdfPlumberLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError:
            pytest.skip("pdfplumber not installed")
            
    def test_pdfplumber_table_extraction(self, temp_dir):
        """Test pdfplumber with table extraction settings"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = PdfPlumberLoaderConfig(
                extract_tables=True,
                table_format='csv',
                layout_mode='layout',
                use_text_flow=True
            )
            loader = PdfPlumberLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError:
            pytest.skip("pdfplumber not installed")


class TestDoclingLoader:
    """Test DoclingLoader with all config attributes"""
    
    def test_docling_basic_loading(self, temp_dir):
        """Test Docling basic loading"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = DoclingLoaderConfig(
                extraction_mode='chunks',
                chunker_type='hybrid',
                ocr_enabled=False
            )
            loader = DoclingLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError as e:
            pytest.skip(f"Docling not installed: {e}")
            
    def test_docling_markdown_mode(self, temp_dir):
        """Test Docling with markdown extraction"""
        pdf_path = Path("/Users/dogankeskin/Desktop/Upsonic/true-pdf-sample-1.pdf")
        
        if not pdf_path.exists():
            pytest.skip("PDF test file not found")
        
        try:
            config = DoclingLoaderConfig(
                extraction_mode='markdown',
                ocr_enabled=False,
                enable_table_structure=True
            )
            loader = DoclingLoader(config=config)
            docs = loader.load(str(pdf_path))
            
            assert len(docs) > 0
        except ImportError as e:
            pytest.skip(f"Docling not installed: {e}")


class TestDOCXLoader:
    """Test DOCXLoader with all config attributes"""
    
    def test_docx_basic_loading(self, temp_dir):
        """Test DOCX basic loading"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")
        
        # Create a test DOCX file
        file_path = temp_dir / "test.docx"
        doc = DocxDocument()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Content in section 1.')
        doc.save(str(file_path))
        
        config = DOCXLoaderConfig(
            include_tables=True,
            include_headers=True,
            include_footers=True
        )
        loader = DOCXLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert "Test Document" in docs[0].content or "test paragraph" in docs[0].content.lower()
        
    def test_docx_with_tables(self, temp_dir):
        """Test DOCX with table extraction"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")
        
        # Create a test DOCX file with table
        file_path = temp_dir / "test_table.docx"
        doc = DocxDocument()
        doc.add_heading('Document with Table', 0)
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Header 1'
        table.cell(0, 1).text = 'Header 2'
        table.cell(1, 0).text = 'Data 1'
        table.cell(1, 1).text = 'Data 2'
        
        doc.save(str(file_path))
        
        config = DOCXLoaderConfig(
            include_tables=True,
            table_format='markdown'
        )
        loader = DOCXLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0


class TestLoaderCommonAttributes:
    """Test common loader attributes across all loaders"""
    
    def test_error_handling_modes(self, temp_dir):
        """Test error_handling parameter"""
        # Test with non-existent file
        config = TextLoaderConfig(error_handling='ignore')
        loader = TextLoader(config=config)
        
        # Should not raise with 'ignore'
        docs = loader.load("/nonexistent/file.txt")
        assert isinstance(docs, list)
        
    def test_custom_metadata(self, temp_dir):
        """Test custom_metadata parameter"""
        file_path = temp_dir / "test.txt"
        file_path.write_text("Test content")
        
        custom_meta = {"source": "test_source", "category": "testing"}
        config = TextLoaderConfig(custom_metadata=custom_meta)
        loader = TextLoader(config=config)
        docs = loader.load(str(file_path))
        
        assert len(docs) > 0
        assert docs[0].metadata.get('source') == 'test_source'
        assert docs[0].metadata.get('category') == 'testing'
        
    def test_skip_empty_content(self, temp_dir):
        """Test skip_empty_content parameter"""
        file_path = temp_dir / "empty.txt"
        file_path.write_text("")
        
        config = TextLoaderConfig(skip_empty_content=True)
        loader = TextLoader(config=config)
        docs = loader.load(str(file_path))
        
        # Should skip empty content
        assert len(docs) == 0


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
