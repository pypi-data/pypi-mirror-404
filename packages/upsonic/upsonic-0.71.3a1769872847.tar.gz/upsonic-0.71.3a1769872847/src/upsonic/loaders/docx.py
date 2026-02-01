import asyncio
from pathlib import Path
from typing import List, Optional, Union

try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    _DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DocxDocument = None
    DocxTable = None
    DocxParagraph = None
    _DOCX_AVAILABLE = False


from upsonic.schemas.data_models import Document
from upsonic.loaders.base import BaseLoader
from upsonic.loaders.config import DOCXLoaderConfig


class DOCXLoader(BaseLoader):
    """
    A loader for Microsoft Word (.docx) files.

    This loader extracts text from paragraphs, tables, headers, and footers.
    Each .docx file is loaded as a single Document object.
    """

    def __init__(self, config: Optional[DOCXLoaderConfig] = None):
        """Initializes the DOCXLoader with its specific configuration."""
        if config is None:
            config = DOCXLoaderConfig()
        if not _DOCX_AVAILABLE:
            from upsonic.utils.printing import import_error
            import_error(
                package_name="python-docx",
                install_command='pip install "upsonic[docx-loader]"',
                feature_name="DOCX loader"
            )
        super().__init__(config)
        self.config: DOCXLoaderConfig = config

    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Gets the list of supported file extensions."""
        return [".docx"]

    def _format_table(self, table: DocxTable) -> str:
        """Formats a table object into a string based on the config."""
        if self.config.table_format == "html":
            html = ["<table>"]
            for row in table.rows:
                html.append("  <tr>")
                for cell in row.cells:
                    html.append(f"    <td>{cell.text}</td>")
                html.append("  </tr>")
            html.append("</table>")
            return "\n".join(html)

        elif self.config.table_format == "markdown":
            markdown = []
            header = [cell.text for cell in table.rows[0].cells]
            markdown.append("| " + " | ".join(header) + " |")
            markdown.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table.rows[1:]:
                row_text = [cell.text.replace("\n", " ") for cell in row.cells]
                markdown.append("| " + " | ".join(row_text) + " |")
            return "\n".join(markdown)
        
        text = []
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text.append("\t".join(row_text))
        return "\n".join(text)

    def _load_single_file(self, file_path: Path) -> List[Document]:
        """Helper method to load a single .docx file into one or zero Documents."""
        try:
            document_id = self._generate_document_id(file_path)
            if document_id in self._processed_document_ids:
                raise FileExistsError(
                    f"Source file '{file_path.resolve()}' has already been processed by this loader instance."
                )
            self._processed_document_ids.add(document_id)

            doc: DocxDocument = docx.Document(file_path)
            content_parts = []

            if self.config.include_headers:
                for section in doc.sections:
                    for header in [section.header, section.first_page_header, section.even_page_header]:
                        if header:
                            for paragraph in header.paragraphs:
                                content_parts.append(paragraph.text)

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            if self.config.include_tables:
                for table in doc.tables:
                    content_parts.append(self._format_table(table))
            
            if self.config.include_footers:
                for section in doc.sections:
                    for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                        if footer:
                            for paragraph in footer.paragraphs:
                                content_parts.append(paragraph.text)
            
            full_content = "\n\n".join(part for part in content_parts if part and part.strip())

            if self.config.skip_empty_content and not full_content.strip():
                return []

            metadata = self._create_metadata(file_path)
            core_props = doc.core_properties
            if self.config.include_metadata:
                metadata.update({
                    "author": core_props.author,
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "title": core_props.title,
                    "subject": core_props.subject,
                    "created": core_props.created,
                    "modified": core_props.modified,
                })
            
            doc_obj = Document(document_id=document_id, content=full_content, metadata=metadata)
            return [doc_obj]

        except Exception as e:
            return self._handle_loading_error(str(file_path), e)

    def load(self, source: Union[str, Path, List[Union[str, Path]]]) -> List[Document]:
        """Loads documents from the given .docx source(s) synchronously."""
        files_to_process = self._resolve_sources(source)
        all_documents = []
        for file_path in files_to_process:
            if not self._check_file_size(file_path):
                continue
            all_documents.extend(self._load_single_file(file_path))
        return all_documents

    async def aload(self, source: Union[str, Path, List[Union[str, Path]]]) -> List[Document]:
        """
        Loads documents from the given .docx source(s) asynchronously.
        """
        files_to_process = self._resolve_sources(source)
        
        valid_files = []
        for file_path in files_to_process:
            if not self._check_file_size(file_path):
                continue
            valid_files.append(file_path)

        tasks = [asyncio.to_thread(self._load_single_file, file) for file in valid_files]
        results = await asyncio.gather(*tasks)
        return [doc for sublist in results for doc in sublist]

    def batch(self, sources: List[Union[str, Path]]) -> List[Document]:
        """A synchronous batch load implementation."""
        return self.load(sources)

    async def abatch(self, sources: List[Union[str, Path]]) -> List[Document]:
        """An efficient asynchronous batch load implementation using asyncio.gather."""
        return await self.aload(sources)