"""Tests for accepting and rejecting tracked changes."""

from __future__ import annotations

from pathlib import Path

from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionDocument, RevisionParagraph


class DescribeAcceptReject_individual:
    """Accept/reject individual tracked changes."""

    def it_accepts_an_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Before ")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_insertion(text="inserted", author="A", revision_id=1)

        tracked.accept()

        assert rp.has_track_changes is False
        assert "inserted" in rp.text

    def it_rejects_an_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Before ")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_insertion(text="inserted", author="A", revision_id=1)

        tracked.reject()

        assert rp.has_track_changes is False
        assert "inserted" not in rp.text

    def it_accepts_a_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_deletion(start=0, end=5, author="A", revision_id=1)

        tracked.accept()

        assert rp.has_track_changes is False
        # After accepting deletion, "Hello" should be gone
        assert "Hello" not in rp.text

    def it_rejects_a_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)
        tracked = rp.add_tracked_deletion(start=0, end=5, author="A", revision_id=1)

        tracked.reject()

        assert rp.has_track_changes is False
        # After rejecting deletion, "Hello" should be restored as normal text
        assert "Hello" in rp.text


class DescribeAcceptReject_document_level:
    """Accept/reject all changes in a document."""

    def it_accepts_all_changes(self):
        doc = Document()
        p1 = doc.add_paragraph("Start ")
        rp1 = RevisionParagraph.from_paragraph(p1)
        rp1.add_tracked_insertion(text="added", author="A", revision_id=1)

        p2 = doc.add_paragraph("Keep this")
        rp2 = RevisionParagraph.from_paragraph(p2)
        rp2.add_tracked_deletion(start=0, end=4, author="A", revision_id=2)

        rdoc = RevisionDocument(doc)
        rdoc.accept_all()

        # After accepting all: insertions kept, deletions removed
        assert len(rdoc.track_changes) == 0

    def it_rejects_all_changes(self):
        doc = Document()
        p1 = doc.add_paragraph("Start ")
        rp1 = RevisionParagraph.from_paragraph(p1)
        rp1.add_tracked_insertion(text="added", author="A", revision_id=1)

        rdoc = RevisionDocument(doc)
        rdoc.reject_all()

        assert len(rdoc.track_changes) == 0


class DescribeAcceptReject_from_docx:
    """Accept/reject using a real docx file."""

    def it_can_accept_all_on_test_docx(self, oxml_docx_path: Path):
        rdoc = RevisionDocument(str(oxml_docx_path))

        # Verify there are track changes first
        assert len(rdoc.track_changes) > 0

        rdoc.accept_all()

        assert len(rdoc.track_changes) == 0
