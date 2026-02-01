"""Tests for replace_tracked and replace_tracked_at."""

from __future__ import annotations

import pytest
from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionDocument, RevisionParagraph


class DescribeRevisionParagraph_replace_tracked:
    """Tests for RevisionParagraph.replace_tracked."""

    def it_replaces_within_single_run(self):
        doc = Document()
        para = doc.add_paragraph("Hello Unisys World")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 1
        assert rp.has_track_changes is True
        # The deletions should contain "Unisys"
        assert any(d.text == "Unisys" for d in rp.deletions)
        # The insertions should contain "test"
        assert any(i.text == "test" for i in rp.insertions)

    def it_handles_multiple_occurrences(self):
        doc = Document()
        para = doc.add_paragraph("Unisys and Unisys again")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 2
        assert len(rp.deletions) == 2
        assert len(rp.insertions) == 2

    def it_returns_zero_when_no_match(self):
        doc = Document()
        para = doc.add_paragraph("No match here")
        rp = RevisionParagraph.from_paragraph(para)

        count = rp.replace_tracked("Unisys", "test", author="Tester")

        assert count == 0
        assert rp.has_track_changes is False

    def it_preserves_surrounding_text(self):
        doc = Document()
        para = doc.add_paragraph("Before Unisys After")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked("Unisys", "test", author="Tester")

        # accepted_text should have "Before test After"
        assert "Before" in rp.accepted_text
        assert "After" in rp.accepted_text
        assert "test" in rp.accepted_text


class DescribeRevisionParagraph_replace_tracked_at:
    """Tests for RevisionParagraph.replace_tracked_at."""

    def it_replaces_within_single_run(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked_at(start=6, end=11, replace_text="Universe", author="Tester")

        assert rp.has_track_changes is True
        assert any(d.text == "World" for d in rp.deletions)
        assert any(i.text == "Universe" for i in rp.insertions)

    def it_replaces_across_multiple_runs(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("Hello ")
        para.add_run("World")
        rp = RevisionParagraph.from_paragraph(para)

        # Replace "o Wor" (spans two runs: "Hello " and "World")
        rp.replace_tracked_at(start=4, end=9, replace_text="X", author="Tester")

        assert rp.has_track_changes is True
        assert any(d.text == "o Wor" for d in rp.deletions)
        assert any(i.text == "X" for i in rp.insertions)

    def it_preserves_before_and_after_text(self):
        doc = Document()
        para = doc.add_paragraph("")
        para.add_run("Hello ")
        para.add_run("World")
        rp = RevisionParagraph.from_paragraph(para)

        rp.replace_tracked_at(start=4, end=9, replace_text="X", author="Tester")

        accepted = rp.accepted_text
        assert "Hell" in accepted
        assert "X" in accepted
        assert "ld" in accepted

    def it_raises_on_invalid_offsets(self):
        doc = Document()
        para = doc.add_paragraph("Hello")
        rp = RevisionParagraph.from_paragraph(para)

        with pytest.raises(ValueError, match="Invalid offsets"):
            rp.replace_tracked_at(start=10, end=15, replace_text="test", author="Tester")

    def it_raises_on_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("")
        rp = RevisionParagraph.from_paragraph(para)

        with pytest.raises(ValueError, match="Invalid offsets"):
            rp.replace_tracked_at(start=0, end=5, replace_text="test", author="Tester")


class DescribeRevisionDocument_find_and_replace_tracked:
    """Tests for RevisionDocument.find_and_replace_tracked."""

    def it_replaces_across_entire_document(self):
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("World is great")
        doc.add_paragraph("No match")
        rdoc = RevisionDocument(doc)

        count = rdoc.find_and_replace_tracked("World", "Universe", author="Bot")

        assert count == 2
        changes = rdoc.track_changes
        assert len(changes) > 0

    def it_returns_zero_for_no_matches(self):
        doc = Document()
        doc.add_paragraph("Hello")
        rdoc = RevisionDocument(doc)

        count = rdoc.find_and_replace_tracked("zzz", "yyy", author="Bot")

        assert count == 0
