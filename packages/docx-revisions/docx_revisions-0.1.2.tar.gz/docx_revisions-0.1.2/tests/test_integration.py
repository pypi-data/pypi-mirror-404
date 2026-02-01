"""Integration tests — round-trip and real docx file tests."""

from __future__ import annotations

from pathlib import Path

from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionDocument, RevisionParagraph


class DescribeRoundTrip:
    """Create → save → reopen → verify."""

    def it_round_trips_tracked_changes(self, tmp_path: Path):
        # Create
        doc = Document()
        para = doc.add_paragraph("Original text. ")
        rp = RevisionParagraph.from_paragraph(para)
        rp.add_tracked_insertion(text="Added text.", author="Author1", revision_id=1)

        # Save
        path = tmp_path / "test_roundtrip.docx"
        doc.save(str(path))

        # Reopen
        rdoc = RevisionDocument(str(path))
        paras = rdoc.paragraphs
        # Find the paragraph with track changes (skip default empty paragraph)
        tc_paras = [p for p in paras if p.has_track_changes]
        assert len(tc_paras) >= 1

        rp2 = tc_paras[0]
        assert len(rp2.insertions) == 1
        assert rp2.insertions[0].text == "Added text."
        assert rp2.insertions[0].author == "Author1"

    def it_round_trips_tracked_deletions(self, tmp_path: Path):
        doc = Document()
        para = doc.add_paragraph("Hello World Goodbye")
        rp = RevisionParagraph.from_paragraph(para)
        rp.add_tracked_deletion(start=6, end=12, author="Deleter", revision_id=1)

        path = tmp_path / "test_del_roundtrip.docx"
        doc.save(str(path))

        rdoc = RevisionDocument(str(path))
        tc_paras = [p for p in rdoc.paragraphs if p.has_track_changes]
        assert len(tc_paras) >= 1
        assert len(tc_paras[0].deletions) == 1
        assert tc_paras[0].deletions[0].text == "World "

    def it_round_trips_replace_tracked(self, tmp_path: Path):
        doc = Document()
        doc.add_paragraph("Replace Foo here and Foo there")
        rdoc = RevisionDocument(doc)
        count = rdoc.find_and_replace_tracked("Foo", "Bar", author="Bot")
        assert count == 2

        path = tmp_path / "test_replace_roundtrip.docx"
        rdoc.save(path)

        rdoc2 = RevisionDocument(str(path))
        changes = rdoc2.track_changes
        assert len(changes) >= 2  # At least 2 del + 2 ins


class DescribeRealDocx:
    """Tests using real docx files."""

    def it_reads_all_changes_from_test_docx(self, oxml_docx_path: Path):
        rdoc = RevisionDocument(str(oxml_docx_path))
        changes = rdoc.track_changes
        # Should have multiple tracked changes
        assert len(changes) > 0

    def it_accepts_all_and_produces_clean_text(self, oxml_docx_path: Path):
        rdoc = RevisionDocument(str(oxml_docx_path))
        rdoc.accept_all()

        # After accepting all, no track changes should remain
        assert len(rdoc.track_changes) == 0

    def it_reports_tricky_docx_has_no_paragraph_level_changes(self, tricky_docx_path: Path):
        rdoc = RevisionDocument(str(tricky_docx_path))
        for para in rdoc.paragraphs:
            assert para.has_track_changes is False


class DescribeRevisionDocument_save:
    """Tests for the save method."""

    def it_saves_to_path(self, tmp_path: Path):
        rdoc = RevisionDocument()
        path = tmp_path / "test_save.docx"
        rdoc.save(path)
        assert path.exists()

    def it_saves_after_modifications(self, tmp_path: Path):
        doc = Document()
        doc.add_paragraph("Hello")
        rdoc = RevisionDocument(doc)
        rdoc.find_and_replace_tracked("Hello", "World", author="Bot")

        path = tmp_path / "test_modified.docx"
        rdoc.save(path)

        # Verify the saved file can be reopened
        rdoc2 = RevisionDocument(str(path))
        assert len(rdoc2.paragraphs) > 0


class DescribeRevisionDocument_construction:
    """Tests for RevisionDocument construction."""

    def it_creates_blank_document_with_None(self):
        rdoc = RevisionDocument()
        assert rdoc.document is not None

    def it_wraps_existing_document(self):
        doc = Document()
        rdoc = RevisionDocument(doc)
        assert rdoc.document is doc

    def it_opens_from_path(self, oxml_docx_path: Path):
        rdoc = RevisionDocument(str(oxml_docx_path))
        assert len(rdoc.paragraphs) > 0
