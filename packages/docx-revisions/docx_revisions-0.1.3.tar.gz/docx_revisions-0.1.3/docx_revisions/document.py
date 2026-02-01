"""RevisionDocument — the main entry point for working with tracked changes.

Opens a ``.docx`` file (or wraps an existing ``Document``) and exposes
paragraphs as ``RevisionParagraph`` objects, plus document-level operations
for accepting, rejecting, and performing find-and-replace with tracking.
"""

from __future__ import annotations

from pathlib import Path
from typing import IO, List

from docx import Document as _new_document
from docx.document import Document as _DocumentClass

from docx_revisions.paragraph import RevisionParagraph
from docx_revisions.revision import TrackedChange


class RevisionDocument:
    """Entry point for reading and writing tracked changes in a docx file.

    Example:
        ```python
        from docx_revisions import RevisionDocument

        rdoc = RevisionDocument("contract.docx")
        for para in rdoc.paragraphs:
            if para.has_track_changes:
                print(para.accepted_text)

        rdoc.accept_all()
        rdoc.save("contract_clean.docx")
        ```
    """

    def __init__(self, path_or_doc: str | Path | IO[bytes] | _DocumentClass | None = None):
        """Open a docx file or wrap an existing ``Document``.

        Args:
            path_or_doc: A file path, file-like object, or an existing
                ``Document`` instance.  Pass ``None`` to create a new
                blank document.
        """
        if isinstance(path_or_doc, _DocumentClass):
            self._document = path_or_doc
        else:
            self._document = _new_document(path_or_doc)

    @property
    def document(self) -> _DocumentClass:
        """The underlying python-docx ``Document`` object."""
        return self._document

    @property
    def paragraphs(self) -> List[RevisionParagraph]:
        """All body paragraphs as ``RevisionParagraph`` objects."""
        return [RevisionParagraph.from_paragraph(p) for p in self._document.paragraphs]

    @property
    def track_changes(self) -> List[TrackedChange]:
        """All tracked changes across the entire document body."""
        changes: List[TrackedChange] = []
        for para in self.paragraphs:
            changes.extend(para.track_changes)
        return changes

    def accept_all(self) -> None:
        """Accept every tracked change in the document.

        Insertions are kept (wrapper removed), deletions are removed entirely.
        """
        for para in self.paragraphs:
            for change in list(para.track_changes):
                change.accept()

    def reject_all(self) -> None:
        """Reject every tracked change in the document.

        Insertions are removed entirely, deletions are kept (wrapper removed,
        ``w:delText`` converted back to ``w:t``).
        """
        for para in self.paragraphs:
            for change in list(para.track_changes):
                change.reject()

    def find_and_replace_tracked(
        self, search_text: str, replace_text: str, author: str = "", comment: str | None = None
    ) -> int:
        """Find and replace across the whole document with track changes.

        Searches all paragraphs in the document body and tables.

        Args:
            search_text: Text to find.
            replace_text: Replacement text.
            author: Author name for the revisions.
            comment: Optional comment text (requires python-docx comment
                support).

        Returns:
            Total number of replacements made.

        Example:
            ```python
            rdoc = RevisionDocument("doc.docx")
            count = rdoc.find_and_replace_tracked(
                "Acme Corp", "NewCo Inc", author="Legal"
            )
            print(f"Replaced {count} occurrences")
            rdoc.save("doc_revised.docx")
            ```
        """
        total_count = 0

        for para in self.paragraphs:
            total_count += para.replace_tracked(search_text, replace_text, author=author, comment=comment)

        for table in self._document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        rp = RevisionParagraph.from_paragraph(p)
                        total_count += rp.replace_tracked(search_text, replace_text, author=author, comment=comment)

        return total_count

    def save(self, path: str | Path) -> None:
        """Save the document to *path*.

        Args:
            path: Destination file path.
        """
        self._document.save(str(path))
