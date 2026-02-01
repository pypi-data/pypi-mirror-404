"""Tests for RevisionParagraph — reading and writing track changes."""

from pathlib import Path

import pytest
from docx import Document

import docx_revisions  # noqa: F401
from docx_revisions import RevisionParagraph

# -- Reading from real docx ------------------------------------------------


class DescribeRevisionParagraph_reading:
    """Reading track changes from OXML_TrackChanges_Test.docx."""

    def it_detects_paragraphs_with_track_changes(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        paras = [RevisionParagraph.from_paragraph(p) for p in doc.paragraphs]
        # Paras 2,3,4,5,7,8 have track changes; others don't
        assert paras[0].has_track_changes is False
        assert paras[1].has_track_changes is False
        assert paras[2].has_track_changes is True
        assert paras[3].has_track_changes is True
        assert paras[4].has_track_changes is True
        assert paras[5].has_track_changes is True
        assert paras[6].has_track_changes is False
        assert paras[7].has_track_changes is True
        assert paras[8].has_track_changes is True

    def it_reads_insertion_from_para_2(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[2])
        assert len(rp.insertions) == 1
        assert rp.insertions[0].author == "Test Author"
        assert rp.insertions[0].text == "inserted"
        assert len(rp.deletions) == 0

    def it_reads_deletion_from_para_3(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[3])
        assert len(rp.deletions) == 1
        assert rp.deletions[0].author == "Test Author"
        assert rp.deletions[0].text == "old "
        assert len(rp.insertions) == 0

    def it_reads_ins_and_del_from_para_4(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[4])
        assert len(rp.insertions) == 1
        assert len(rp.deletions) == 1
        assert rp.insertions[0].author == "Editor A"
        assert rp.insertions[0].text == "color"
        assert rp.deletions[0].author == "Editor A"
        assert rp.deletions[0].text == "colour"

    def it_reads_deletion_from_para_5(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[5])
        assert len(rp.deletions) == 1
        assert rp.deletions[0].author == "Reviewer"
        assert rp.deletions[0].text == "reenter"

    def it_reads_insertion_with_em_dash_from_para_7(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[7])
        assert len(rp.insertions) == 1
        assert rp.insertions[0].author == "Reviewer B"
        assert "—" in rp.insertions[0].text  # em-dash

    def it_reads_conflicting_ins_del_from_para_8(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[8])
        assert len(rp.insertions) == 1
        assert len(rp.deletions) == 1
        assert rp.insertions[0].author == "Author A"
        assert rp.deletions[0].author == "Author B"

    def it_provides_accepted_text(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        # Para 3: deletion of "old " — accepted text starts with "This sentence remains."
        rp3 = RevisionParagraph.from_paragraph(doc.paragraphs[3])
        assert rp3.accepted_text.startswith("This sentence remains.")

    def it_provides_original_text(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        # Para 3: deletion of "old " — original text starts with "This old sentence remains."
        rp3 = RevisionParagraph.from_paragraph(doc.paragraphs[3])
        assert rp3.original_text.startswith("This old sentence remains.")

    def it_returns_track_changes_in_order(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[4])
        changes = rp.track_changes
        # Para 4 has del(colour) then ins(color) in document order
        assert len(changes) == 2


# -- Writing ---------------------------------------------------------------


class DescribeRevisionParagraph_writing:
    """Writing track changes via RevisionParagraph."""

    def it_adds_tracked_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Existing text. ")
        rp = RevisionParagraph.from_paragraph(para)

        tracked = rp.add_tracked_insertion(text="New text.", author="TestAuthor", revision_id=1)

        assert tracked.author == "TestAuthor"
        assert tracked.revision_id == 1
        assert tracked.text == "New text."
        assert rp.has_track_changes is True

    def it_auto_generates_revision_id(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        rp = RevisionParagraph.from_paragraph(para)

        t1 = rp.add_tracked_insertion(text="first", author="A")
        t2 = rp.add_tracked_insertion(text="second", author="A")

        assert t2.revision_id > t1.revision_id

    def it_adds_tracked_deletion(self):
        doc = Document()
        para = doc.add_paragraph("Hello World")
        rp = RevisionParagraph.from_paragraph(para)

        tracked = rp.add_tracked_deletion(start=0, end=5, author="Deleter", revision_id=10)

        assert tracked.author == "Deleter"
        assert tracked.text == "Hello"
        assert rp.has_track_changes is True

    def it_raises_on_invalid_deletion_offsets(self):
        doc = Document()
        para = doc.add_paragraph("Hello")
        rp = RevisionParagraph.from_paragraph(para)

        with pytest.raises(ValueError, match="Invalid offsets"):
            rp.add_tracked_deletion(start=10, end=15, author="A")

    def it_sets_date_on_insertion(self):
        doc = Document()
        para = doc.add_paragraph("Text")
        rp = RevisionParagraph.from_paragraph(para)

        tracked = rp.add_tracked_insertion(text="new", author="A")
        assert tracked.date is not None


# -- iter_inner_content ----------------------------------------------------


class DescribeRevisionParagraph_iter_inner_content:
    """Tests for iter_inner_content with include_revisions."""

    def it_excludes_revisions_by_default(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[2])
        items = list(rp.iter_inner_content())
        # Should only yield Run/Hyperlink, not TrackedInsertion
        from docx_revisions.revision import TrackedInsertion

        assert not any(isinstance(i, TrackedInsertion) for i in items)

    def it_includes_revisions_when_asked(self, oxml_docx_path: Path):
        doc = Document(str(oxml_docx_path))
        rp = RevisionParagraph.from_paragraph(doc.paragraphs[2])
        items = list(rp.iter_inner_content(include_revisions=True))
        from docx_revisions.revision import TrackedInsertion

        assert any(isinstance(i, TrackedInsertion) for i in items)
