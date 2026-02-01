from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from functools import partial 
def tableDOC(doc, data, row, col, title):
    row_labels = row
    col_labels = col
    doc.add_paragraph(title)
    table = doc.add_table(rows=1 + len(data), cols=1 + len(col_labels))
    table.style = 'Table Grid'
    header_cell = table.cell(0, 0)
    header_cell.text = "实况\预报"
    header_cell.paragraphs[0].runs[0].font.bold = False
    for idx, label in enumerate(col_labels):
        cell = table.cell(0, idx+1)
        cell.text = label
        cell.paragraphs[0].runs[0].font.bold = False
    for row_idx, (label, row) in enumerate(zip(row_labels, data)):
        table.cell(row_idx + 1, 0).text = label
        table.cell(row_idx + 1, 0).paragraphs[0].runs[0].font.bold = False
        for col_idx, value in enumerate(row):
            table.cell(row_idx + 1, col_idx + 1).text = str(value) 
def add_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    heading.style.font.size = Pt(18)  # 设置字体大小为18磅
def add_text_to_doc(doc, text):
    para = doc.add_paragraph(text)
    para.style = doc.styles['Body Text']
def add_image_to_doc(doc, image_path):
    doc.add_picture(image_path, width=Inches(5.0)) 
def partialFN(tabledata):
    return partial(tableDOC,row=tabledata["labels"], col= tabledata["columns"],title=tabledata["title"])

dataCY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云10min晴雨混淆矩阵"
    } 
dataWTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信10min晴雨混淆矩阵"
       } 
data2CY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云20min晴雨混淆矩阵"
    } 
data2WTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信20min晴雨混淆矩阵"
       } 

# {0: ['0.016', '0.976', '6157.69%'], 1: ['0.177', '0.892', '405.15%']} dict_values([['0.386', '0.921', '138.33%'], ['0.399', '0.819', '105.52%']])  生成word表格输入数据

# from docx import Document
# doc = Document() 
# doc.save("./test6.docx")

"""

import pandas as pd  
from shancx.Dsalgor.matrixLibJU import TS ,ACC ,F1 ,FAR ,PO ,sun_rain_Matrix ,pre1h_Matrix
from shancx.DOCJU import partialFN,add_text_to_doc,add_image_to_doc,add_heading
import numpy as np

from docx import Document


dataCY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云1H晴雨混淆矩阵"
    } 
dataWTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信1H晴雨混淆矩阵"
       } 
data2CY ={
    "labels": ["晴", "雨"],
    "columns": ["晴", "雨"],
    "title": f"彩云2H晴雨混淆矩阵"
    } 
data2WTX ={ 
       "labels": ["晴", "雨"],
       "columns": ["晴", "雨"],
       "title": f"维天信2H晴雨混淆矩阵"
       } 

TSVD= {
       "labels": ["1h", "2h"],
       "columns":["彩云", "维天信", "提升百分比"],
       "title": "TS评分"
      }
F1VD = {
        "labels": ["1h", "2h"],
        "columns":["彩云", "维天信", "提升百分比"],
        "title": "F1评分"
      }
ACCVD= {
    "labels": ["1h", "2h"],
    "columns":["彩云", "维天信", "提升百分比"],
    "title": "准确率评分"
    }
POVD = {
    "labels": ["1h", "2h"],
    "columns":["彩云", "维天信", "提升百分比"],
    "title": "漏报率评分 "                        
    }
FARVD = {
    "labels": ["1h", "2h"],
    "columns":["彩云", "维天信", "提升百分比"],
    "title": "空报率评分 "
     }

w_1h_bigD = {
           "labels": ["晴", "小雨", "中雨", "大雨", "暴雨"],
           "columns": ["晴", "小雨", "中雨", "大雨", "暴雨"],
           "title": f"维天信1小时预报"
           }
w_2h_bigD = {
           "labels": ["晴", "小雨", "中雨", "大雨", "暴雨"],
           "columns": ["晴", "小雨", "中雨", "大雨", "暴雨"],
           "title": f"维天信2小时预报"
           }
  
dataCY = partialFN(dataCY)
dataWTX = partialFN(dataWTX)
data2CY = partialFN(data2CY)
data2WTX = partialFN(data2WTX)                                          
TSVD = partialFN(TSVD)
F1VD = partialFN(F1VD)
ACCVD = partialFN(ACCVD)
POVD = partialFN(POVD)
FARVD = partialFN(FARVD)
w_1h_bigD = partialFN(w_1h_bigD)
w_2h_bigD = partialFN(w_2h_bigD)


basepathcsv = "/mnt/wtx_weather_forecast/scx/EXAMDATACSV10min/10min_202412240000_202412250000.csv"
df = pd.read_csv(basepathcsv) 
TSV = TS(df["PRE1_r"], df["PRE1_w"], thresholdR=0.1, thresholdF=0.1) 

TSV = {}  
for i in range(1,3):
    F1h =TS( df[f"PRE{i}_r"],df[f"PRE{i}_c"],thresholdR=0.1, thresholdF=0.031)
    F1hm =TS(df[f"PRE{i}_r"],df[f"PRE{i}_w"],thresholdR=0.1, thresholdF=0.1)
    TSV[i]=[np.round(F1h,3),np.round(F1hm,3),f"{np.round((F1hm-F1h)/F1h*100,2)}%"]
    print(f"{i}h {np.round(F1h,3)} {np.round(F1hm,3)} {np.round((F1hm-F1h)/F1h*100,2)}%")
print("F1")
print("time", "CY", "WTX")
F1V = {}
for i in range(1,3):
    F1h =F1(df[f"PRE{i}_r"], df[f"PRE{i}_c"],thresholdR=0.1, thresholdF=0.031)
    F1hm =F1(df[f"PRE{i}_r"], df[f"PRE{i}_w"],thresholdR=0.1, thresholdF=0.1)
    F1V[i] = [np.round(F1h, 3), np.round(F1hm, 3),f"{np.round((F1hm-F1h)/F1h*100,2)}%"]
    print(f"{i}h {np.round(F1h,3)} {np.round(F1hm,3)} {np.round((F1hm-F1h)/F1h*100,2)}%")
print("ACC")
print("time", "CY", "WTX")
ACCV = {}
for i in range(1, 3):
    F1h = ACC(df[f"PRE{i}_r"], df[f"PRE{i}_c"],thresholdR=0.1, thresholdF=0.031)
    F1hm = ACC(df[f"PRE{i}_r"], df[f"PRE{i}_w"],thresholdR=0.1, thresholdF=0.1)
    ACCV[i] = [np.round(F1h, 3), np.round(F1hm, 3),f"{np.round((F1hm-F1h)/F1h*100,2)}%"]
    print(f"{i}h {np.round(F1h, 3)} {np.round(F1hm, 3)} {np.round((F1hm - F1h) / F1h * 100, 2)}%")
print("PO")
print("time", "CY", "WTX")
POV = {}
for i in range(1,3):
    F1h =PO(df[f"PRE{i}_r"], df[f"PRE{i}_c"],thresholdR=0.1, thresholdF=0.031)
    F1hm =PO(df[f"PRE{i}_r"], df[f"PRE{i}_w"],thresholdR=0.1, thresholdF=0.1)
    POV[i] = [np.round(F1h, 3), np.round(F1hm, 3),f"{np.round((F1hm-F1h)/F1h*100,2)*-1}%"]
    print(f"{i}h {np.round(F1h,3)} {np.round(F1hm,3)} {np.round((F1hm-F1h)/F1h*100,2)*-1}%")
print("FAR")
print("time", "CY", "WTX")
FARV = {}
for i in range(1, 3):
    F1h = FAR(df[f"PRE{i}_r"], df[f"PRE{i}_c"],thresholdR=0.1, thresholdF=0.031)
    F1hm = FAR(df[f"PRE{i}_r"], df[f"PRE{i}_w"],thresholdR=0.1, thresholdF=0.1)
    FARV[i] = [np.round(F1h, 3), np.round(F1hm, 3),f"{np.round((F1hm-F1h)/F1h*100,2)*-1}%"]
    print(f"{i}h {np.round(F1h, 3)}  {np.round(F1hm, 3)} {np.round((F1hm-F1h)/F1h*100,2)*-1}%")

cm1_C = sun_rain_Matrix(df["PRE1_r"].values, df["PRE1_c"].values,0.031)
cm1_W = sun_rain_Matrix(df["PRE1_r"].values, df["PRE1_w"].values,0.1)
cm2_C = sun_rain_Matrix(df["PRE2_r"].values, df["PRE2_c"].values,0.031)
cm2_W = sun_rain_Matrix(df["PRE2_r"].values, df["PRE2_w"].values,0.1)

cm1_W_pre1h = pre1h_Matrix(df["PRE1_r"].values, df["PRE1_w"].values,"WTX")
cm2_W_pre1h = pre1h_Matrix(df["PRE2_r"].values, df["PRE2_w"].values,"WTX")

doc = Document()    #CY_fn(doc,data =
add_heading(doc,"测试报告样例")
dataCY   (doc,data = cm1_C)
dataWTX  (doc,data =cm1_W )
data2CY  (doc,data = cm2_C)
data2WTX (doc,data = cm2_W)
TSVD     (doc,data =TSV.values() )
F1VD     (doc,data =F1V.values() )
ACCVD    (doc,data = ACCV.values())
add_image_to_doc(doc,"./TSF1.png")
add_text_to_doc(doc,"TS评分、F1和准确率ACC对比")
add_image_to_doc(doc,"./POFAR.png")
add_text_to_doc(doc,"漏报PO、空报FAR对比")
POVD     (doc,data = POV.values())
FARVD    (doc,data = FARV.values())
w_1h_bigD(doc,data = cm1_W_pre1h)   
w_2h_bigD(doc,data = cm2_W_pre1h)
add_image_to_doc(doc,"bigsmall.png")
doc.save("./makedoc_test.docx")
 

"""
