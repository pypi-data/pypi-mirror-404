# Built-In imports
import re

# Third-party imports
from .pydantic_models import WordTableModel
from docx import Document
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def cell_shader(cell: _Cell, hex_color: tuple[int, int, int]) -> None:
    shading_elm = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), hex_color.strip("#").upper())
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def delete_paragraph(paragraph: Paragraph) -> None:
    """
    Function to delete a given paragraph in a Word document.
    Requires as input the doxc paragraph object.
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def set_col_width(table: Table, column: int, width: int | float) -> None:
    width = Inches(float(width))
    for row in table.rows:
        row.cells[column].width = width


def find_para_by_string(doc_obj: _Document, search: str) -> int:
    """
    Function to find a string in a Word document and return the paragraph
    number. Receives as input the Word document object (using docx module) and
    a search string.
    """
    for i, p in enumerate(doc_obj.paragraphs):
        if re.match(search, p.text):
            return i


def get_para_by_string(doc_obj: _Document, search: str) -> Paragraph:
    """
    Function to find a string in a Word document and return the paragraph.
    Receives as input the Word document object (using docx module) and
    a search string.
    """
    for p in doc_obj.paragraphs:
        if re.match(search, p.text):
            return p


def replace_placeholder_with_table(
    doc_obj: _Document, placeholder: str, table: Table
) -> None:
    """
    Function to relocate a Word table object to immediately follow a given
    reference paragraph identified by the placeholder. Receives as input the
    placeholder string and the Word table object (using docx module).
    After moving the Word table after the placeholder paragraph, delete the
    placeholder paragraph.
    """
    # Locate the paragraph from the supplied placeholder text
    paragraph: Paragraph = get_para_by_string(doc_obj, placeholder)
    if not paragraph:
        print(f'WARNING: Could not locate placeholder "{placeholder}"')
    else:
        # Move the Word table to a new paragraph immediately after the placeholder paragraph
        paragraph._p.addnext(table._tbl)
        # Delete the placeholder paragraph
        delete_paragraph(paragraph)


def build_table(
    docx_obj: _Document | Table, table_dict: dict, remove_leading_para: bool = True
) -> Table:
    """
    Convert a WordTableModel-style dictionary into a Word table object.
    Supports nested tables and merged cells within a row.
    Automatically sets nested table widths to fill merged cells.
    """
    raw_table = WordTableModel.model_validate(table_dict)

    # Create table if docx_obj is _Document
    if isinstance(docx_obj, _Document):
        table = docx_obj.add_table(
            rows=len(raw_table.rows),
            cols=len(raw_table.rows[0].cells) if raw_table.rows else 0,
        )
    elif isinstance(docx_obj, Table):
        table = docx_obj
    else:
        raise ValueError(f"docx_obj must be _Document or Table, got {type(docx_obj)}")

    # Apply table style
    if raw_table.style:
        table.style = raw_table.style

    for i, row in enumerate(raw_table.rows):
        table_row = table.rows[i]
        col_idx = 0

        while col_idx < len(row.cells):
            cell = row.cells[col_idx]
            table_cell = table_row.cells[col_idx]

            # --- Handle merged cells ---
            if isinstance(cell, str) and cell.lower() == "merge":
                left_cell = table_row.cells[col_idx - 1]
                left_cell.merge(table_cell)
                col_idx += 1
                continue

            # --- Regular cell formatting ---
            if cell.background_color:
                cell_shader(table_cell, cell.background_color)
            if cell.width:
                set_col_width(table, col_idx, cell.width)

            # --- Paragraphs ---
            for para_idx, para in enumerate(cell.paragraphs):
                if para_idx >= len(table_cell.paragraphs):
                    p = table_cell.add_paragraph()
                else:
                    p = table_cell.paragraphs[para_idx]

                if isinstance(para.text, str):
                    p.text = para.text
                elif isinstance(para.text, list):
                    # Multi-line text with breaks
                    p.text = ""
                    for idx, run_text in enumerate(para.text):
                        run = p.add_run(run_text)
                        if idx < len(para.text) - 1:
                            run.add_break()

                # Paragraph alignment
                if para.alignment:
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                    }
                    p.alignment = alignment_map.get(para.alignment.value.lower())

                # Paragraph style
                if para.style:
                    p.style = para.style

            # --- Nested table ---
            if cell.table:
                # Compute width of merged cell
                merged_cols = 1
                temp_idx = col_idx + 1
                while (
                    temp_idx < len(row.cells)
                    and isinstance(row.cells[temp_idx], str)
                    and row.cells[temp_idx].lower() == "merge"
                ):
                    merged_cols += 1
                    temp_idx += 1

                # Sum widths of merged columns, or use default
                parent_width = sum(
                    table.columns[col_idx + k].width or Cm(2.5)
                    for k in range(merged_cols)
                )

                nested_rows = len(cell.table.rows)
                nested_cols = len(cell.table.rows[0].cells) if cell.table.rows else 0

                nested_table = table_cell.add_table(rows=nested_rows, cols=nested_cols)
                nested_table.allow_autofit = False  # We'll set widths manually

                # Assign column widths proportionally
                nested_col_width = parent_width / nested_cols
                for col in nested_table.columns:
                    for nested_cell in col.cells:
                        nested_cell.width = nested_col_width

                # Recursively build nested table
                build_table(nested_table, cell.table.model_dump())

                # Remove leading paragraph in merged/nested cell
                if remove_leading_para and table_cell.paragraphs:
                    delete_paragraph(table_cell.paragraphs[0])

            col_idx += 1

    # Remove leading empty paragraph in document
    if isinstance(docx_obj, _Document) and remove_leading_para and docx_obj.paragraphs:
        if not docx_obj.paragraphs[0].text.strip():
            delete_paragraph(docx_obj.paragraphs[0])

    return table


def insert_paragraph_after(paragraph: Paragraph, text: str = None, style: str = None):
    """
    Insert a new paragraph after the given paragraph.
    """

    # Create a new empty <w:p> element
    new_p = OxmlElement("w:p")

    # Insert the new <w:p> after the given paragraph’s <w:p>
    paragraph._p.addnext(new_p)

    # Wrap the XML element as a python-docx Paragraph
    new_para = Paragraph(new_p, paragraph._parent)

    # Set text and style
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style


# TABLE (DICTIONARY) SHORTCUTS
def insert_text_into_row(cell_text: list) -> dict:
    """
    Generate a table row dictionary from a list of text strings.
    Each string in the list becomes a cell in the row.
    Assumes no styling.

    :param cell_text: A list of text strings for each cell in the row. Use "merge" to indicate merged cells.
    :type cell_text: list
    :return: A dictionary representing the table row.
    :rtype: dict
    """

    row: dict = {"cells": []}
    for text in cell_text:
        if text.lower() == "merge":
            row["cells"].append("merge")
            continue

        row["cells"].append(
            {
                "paragraphs": [
                    {
                        "text": [text],
                    },
                ],
            }
        )

    return row


def insert_text_by_table_coords(table: dict, row: int, col: int, text: str) -> dict:
    table["rows"][row]["cells"][col]["paragraphs"][0]["text"] = [text]
    return table


def generate_table(
    num_rows: int, num_cols: int, header_row: list, style: str = None
) -> dict:
    """
    Generate a basic table dictionary with specified number of rows and columns.
    Each cell contains empty text.

    :param num_rows: Number of rows in the table.
    :type num_rows: int
    :param num_cols: Number of columns in the table.
    :type num_cols: int
    :param header_row: A list of text strings for the header row.
    :type header_row: list
    :param style: The style to apply to the table.
    :type style: str
    :return: A dictionary representing the table.
    :rtype: dict
    """

    table: dict = {"rows": []}

    if style:
        table["style"] = style

    for _ in range(num_rows):
        row: dict = {"cells": []}
        for _ in range(num_cols):
            row["cells"].append(
                {
                    "paragraphs": [
                        {
                            "text": [""],
                        },
                    ],
                }
            )
        table["rows"].append(row)

    if header_row:
        for col_idx, header_text in enumerate(header_row):
            table["rows"][0]["cells"][col_idx]["paragraphs"][0]["text"] = [header_text]

    return table
