import os
import os.path
import glob
from datetime import datetime
import uuid

from typing import Union

import docx
from pypdf import PdfReader

from sentence_transformers import SentenceTransformer
from sklearn.cluster import KMeans
from typing import List
import re

def format_pdf_date(date_str):
    """
    Convert PDF date string to readable format.
    PDF dates are typically in format "D:YYYYMMDDHHmmSS" or "D:YYYYMMDDHHmmSS+HH'mm'"
    
    Args:
        date_str (str): PDF date string
        
    Returns:
        str: Formatted date string or None if conversion fails
    """
    if not date_str:
        return None
        
    try:
        # Remove 'D:' prefix if present
        if date_str.startswith('D:'):
            date_str = date_str[2:]
        
        # Basic format: YYYYMMDDHHmmSS
        if len(date_str) >= 14:
            dt = datetime.strptime(date_str[:14], '%Y%m%d%H%M%S')
            return dt.strftime('%Y-%m-%d %H:%M:%S')
        return None
    except:
        return None


def get_word_text(filename:str, metadata:dict) -> list[dict]:
    """
    Get all text from a Word file, using the package 'docx'. No images are currently OCR'd.
    
    Args:
        filename (str): the full filename to use

        metadata (dict): the metadata for the file, currently we only use the 'id' value.
    
    Returns:
        list[dict]: A list of dictionary, where every dictionary is the text for one of the paragraphs
    """
    doc = docx.Document(filename)
    jsons = []
    for para in doc.paragraphs:
        try:
            line = {
                "text": para.text,
                "id": metadata["id"],
            }
            
            jsons.append(line)
        except:
            print(f"File could not be processed: {filename}")

    return jsons


def get_pdf_text(filename:str, metadata:dict) -> list[dict]:
    """
    Get all text from a PDF file, using the package 'pymupdf'. No images or image PDFs are currently OCR'd.
    
    Args:
        filename (str): the full filename to use

        metadata (dict): the metadata for the file, currently we only use the 'id' value.
    
    Returns:
        list[dict]: A list of dictionary, where every dictionary is the text for one of the pages
    """
    #doc = fitz.open(filename)
    reader = PdfReader(filename)
    jsons = []
    #for page in doc:
    for page in reader.pages:
        try:
            line = {
                #"text": page.get_text(),
                "text": page.extract_text(),
                "id": metadata["id"]
            }
            
            jsons.append(line)
        except:
            print(f"File could not be processed: {filename}")

    return jsons

def get_filename_metadata(filename:str, metadata_id:[str, int]):
    """
    Get the filename metadata.
    
    Args:
        filename (str):    The input filename to extract the metadata from. 
                 
        metadata_id ([str, int]): The 'id' to be used for the metadata
    
    Returns:
        metadata = {
            "id": metadata_id,
            "filename": "the_actual_filename"
        }
    """


    actual_filename= os.path.basename(filename)
    metadata = {
        "id": metadata_id,
        "filename": actual_filename
    }
        
    return metadata

def process_filename(filename:str, metadata_id:Union[str, int]) -> tuple[list[dict], dict]:
    """
    Get all texts and metadata from a PDF/Word file.
    
    Args:
        file (str): the full filename to use

        metadata_id ([str, int]): the metadata_id to be used
    
    Returns:
        - list[dict]: A list of dicts with text and and the metadata id. 
        - dict: The created metadata
    """
    try:
        metadata = get_filename_metadata(filename, metadata_id)
       
        read_inputs = ""
        
        if filename.endswith(".docx") or filename.endswith(".doc"):
            #print("word doc")
            doc = docx.Document(filename)
            core_properties = doc.core_properties
            metadata['created'] = str(core_properties.created)
            metadata['modified'] = str(core_properties.modified)
            metadata['author'] = core_properties.author
            read_inputs = get_word_text(filename, metadata)
       
        elif filename.endswith(".pdf") :
            #print("pdf doc")
            #doc = fitz.open(filename)
            reader = PdfReader(filename)
            pdf_metadata = reader.metadata
            metadata['created'] =  str(format_pdf_date(pdf_metadata.get('/CreationDate')))
            metadata['modified'] = str(format_pdf_date(pdf_metadata.get('/ModDate', None)))
            metadata['author'] = pdf_metadata.author
            read_inputs = get_pdf_text(filename, metadata)
    except Exception as e:
        print(f"Error handling file: {filename}")
        print(e)
        return "", metadata
        
    return read_inputs, metadata

def process_documents(folder:str, print_info:bool=True, exclude_extensions:List=[]):
    docs_count = 0
    corpus_json = []
    all_metadata = {}
    for filename in glob.glob(f"{folder}/**", recursive=True):
        extension = os.path.splitext(filename)[1]  
        if not os.path.isdir(filename) and not "~" in filename and extension not in exclude_extensions:
            metadata_id = uuid.uuid4()
            if print_info:
                docs_count += 1

                if docs_count % 50 == 0:
                    print(f"Processed: {docs_count}")

            read_inputs, metadata = process_filename(filename, metadata_id)
            #print(metadata)

            if read_inputs and read_inputs != "":
                    
                corpus_json.extend(read_inputs)
                
                all_metadata[metadata["id"]] = metadata
            
    if print_info:
        print(f"{docs_count} docs processed")
    
    return corpus_json, all_metadata


def download_nltk_data(package="punkt"):
    import os
    from nltk.data import find
    try:
        # Try to find the package
        find(f'tokenizers/{package}')
    except LookupError:
        # Package not found, so download it
        print(f"Downloading {package} package...")
        nltk.download(package)
        print("Download complete.")
    else:
        print(f"{package} package is already downloaded.")

def basic_sentence_chunks(text: str, chunk_size: int = 3) -> List[str]:
    import nltk
    from nltk.tokenize import sent_tokenize 

    nltk.download('punkt_tab', quiet=True)

    """Split text into chunks of N sentences."""
    sentences = sent_tokenize(text)
    chunks = []
    
    for i in range(0, len(sentences), chunk_size):
        chunk = ' '.join(sentences[i:i + chunk_size])
        chunks.append(chunk)
        
    return chunks

@DeprecationWarning
def sliding_window_chunks(text: str, window_size: int = 512, stride: int = 256) -> List[str]:
    """Split text using sliding window with overlap."""
    chunks = []
    
    for i in range(0, len(text), stride):
        chunk = text[i:i + window_size]
        # Only add if chunk is substantial
        if len(chunk) > window_size // 2:
            chunks.append(chunk)
            
    return chunks

def semantic_clustering_chunks(text: str, num_clusters: int = 3) -> List[str]:
    """Split text using sentence embeddings and clustering."""
    from nltk.tokenize import sent_tokenize
    sentences = sent_tokenize(text)
    if len(sentences) < num_clusters:
        return [text]
    
    sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        
    # Generate embeddings
    embeddings = sentence_model.encode(sentences)
    
    # Cluster sentences
    kmeans = KMeans(n_clusters=num_clusters, random_state=42)
    clusters = kmeans.fit_predict(embeddings)
    
    # Group sentences by cluster
    cluster_chunks = [[] for _ in range(num_clusters)]
    for sentence, cluster_id in zip(sentences, clusters):
        cluster_chunks[cluster_id].append(sentence)
        
    # Join sentences within each cluster
    chunks = [' '.join(chunk) for chunk in cluster_chunks if chunk]
    return chunks

def smart_paragraph_chunks(text: str, max_chunk_size: int = 1000) -> List[str]:
    """Split text into chunks based on paragraph breaks while respecting size limits."""
    from nltk.tokenize import sent_tokenize
    # Split on paragraph breaks
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current_chunk = []
    current_size = 0
    
    for para in paragraphs:
        para = para.strip()
        para_size = len(para)
        
        if current_size + para_size > max_chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_size = 0
            
        if para_size > max_chunk_size:
            # If paragraph itself is too large, split it
            sentences = sent_tokenize(para)
            temp_chunk = []
            temp_size = 0
            
            for sent in sentences:
                sent_size = len(sent)
                if temp_size + sent_size > max_chunk_size and temp_chunk:
                    chunks.append(' '.join(temp_chunk))
                    temp_chunk = []
                    temp_size = 0
                temp_chunk.append(sent)
                temp_size += sent_size
                
            if temp_chunk:
                chunks.append(' '.join(temp_chunk))
        else:
            current_chunk.append(para)
            current_size += para_size
            
    if current_chunk:
        chunks.append(' '.join(current_chunk))
        
    return chunks
