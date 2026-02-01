"""
MS Word Tool migrated to use AbstractDocumentTool framework.
"""
from typing import Any, Dict, List, Optional, Union
import re
import tempfile
import os
from pathlib import Path
import io
from urllib.parse import urlparse
import aiohttp
import aiofiles
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from jinja2 import Environment, FileSystemLoader
from pydantic import BaseModel, Field, field_validator
import mammoth
import markdown
from bs4 import BeautifulSoup, NavigableString
from markdownify import markdownify as md
from .document import AbstractDocumentTool, DocumentGenerationArgs


class MSWordArgs(DocumentGenerationArgs):
    """Arguments schema for MS Word Document generation."""

    template_name: Optional[str] = Field(
        None,
        description="Name of the HTML template (e.g., 'report.html') to render before conversion"
    )
    template_vars: Optional[Dict[str, Any]] = Field(
        None,
        description="Variables to pass to the HTML template (e.g., title, author, date)"
    )
    docx_template: Optional[str] = Field(
        None,
        description="Path to a DOCX template file to use as base document"
    )
    style_config: Optional[Dict[str, Any]] = Field(
        None,
        description="Custom styling configuration for the document"
    )
    page_margins: Optional[Dict[str, float]] = Field(
        None,
        description="Page margins in inches (top, bottom, left, right)"
    )

    @field_validator('template_name')
    @classmethod
    def validate_template_name(cls, v):
        if v and not v.endswith('.html'):
            v = f"{v}.html"
        return v


class MSWordTool(AbstractDocumentTool):
    """
    Microsoft Word Document Generation Tool.

    This tool converts text content (including Markdown and HTML) into professionally
    formatted Word documents (.docx). It supports custom templates, styling, and
    advanced document formatting features.

    Features:
    - Markdown to Word conversion with proper formatting
    - HTML to Word conversion support
    - Custom DOCX template support
    - Jinja2 HTML template processing
    - Configurable styling and page setup
    - Table, list, and heading support
    - Professional document formatting
    """

    name = "msword_generator"
    description = (
        "Generate Microsoft Word documents from text, Markdown, or HTML content. "
        "Supports custom templates, styling, and professional document formatting. "
        "Perfect for creating reports, documentation, and formatted documents."
    )
    args_schema = MSWordArgs

    # Document type configuration
    document_type = "document"
    default_extension = "docx"
    supported_extensions = [".docx", ".dotx"]

    def __init__(
        self,
        templates_dir: Optional[Path] = None,
        default_html_template: Optional[str] = None,
        **kwargs
    ):
        """
        Initialize the MS Word Tool.

        Args:
            templates_dir: Directory containing HTML and DOCX templates
            default_html_template: Default HTML template for content processing
            **kwargs: Additional arguments for AbstractDocumentTool
        """
        super().__init__(templates_dir=templates_dir, **kwargs)

        self.default_html_template = default_html_template

        # Initialize Jinja2 environment for HTML templates
        if self.templates_dir:
            self.html_env = Environment(
                loader=FileSystemLoader(str(self.templates_dir)),
                autoescape=True
            )
        else:
            self.html_env = None

    def _detect_content_type(self, text: str) -> str:
        """Detect if content is HTML, Markdown, or plain text."""
        text_stripped = text.strip()

        # Simple HTML detection
        if (text_stripped.startswith('<') and text_stripped.endswith('>')) or \
           any(tag in text_stripped.lower() for tag in ['<html', '<div', '<p', '<h1']):
            return 'html'

        # Markdown detection
        markdown_patterns = [
            r'^#{1,6}\s',           # Headers
            r'^\*\s',               # Bullet points
            r'^\d+\.\s',            # Numbered lists
            r'\*\*.*?\*\*',         # Bold
            r'\*.*?\*',             # Italic
            r'`.*?`',               # Code
            r'\[.*?\]\(.*?\)',      # Links
        ]

        for pattern in markdown_patterns:
            if re.search(pattern, text_stripped, re.MULTILINE):
                return 'markdown'

        return 'markdown'  # Default to markdown for processing

    def _render_html_template(
        self,
        content: str,
        template_name: Optional[str],
        template_vars: Optional[Dict[str, Any]]
    ) -> str:
        """Render content through Jinja2 HTML template if provided."""
        if not template_name or not self.html_env:
            return content

        try:
            template = self.html_env.get_template(template_name)
            vars_dict = template_vars or {}

            # Add default variables
            vars_dict.setdefault('content', content)
            vars_dict.setdefault('date', self._get_current_date())
            vars_dict.setdefault('timestamp', self._get_current_timestamp())

            rendered = template.render(**vars_dict)
            self.logger.info(
                f"Rendered content through HTML template: {template_name}"
            )
            return rendered

        except Exception as e:
            self.logger.error(f"HTML template rendering failed: {e}")
            return content

    def _preprocess_markdown(self, text: str) -> str:
        """Preprocess markdown to handle common issues."""
        # Replace placeholder variables with empty strings
        text = re.sub(r'\{[a-zA-Z0-9_]+\}', '', text)

        # Handle f-strings that weren't evaluated
        text = re.sub(r'f"""(.*?)"""', r'\1', text, flags=re.DOTALL)
        text = re.sub(r"f'''(.*?)'''", r'\1', text, flags=re.DOTALL)

        # Remove triple backticks and language indicators
        text = re.sub(r'```[a-zA-Z]*\n', '', text)
        text = re.sub(r'```', '', text)

        # Fix heading issues (ensure space after #)
        text = re.sub(r'(#+)([^ \n])', r'\1 \2', text)

        # Fix escaped newlines if any
        text = text.replace('\\n', '\n')

        return text

    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert markdown to HTML."""
        try:
            html = markdown.markdown(
                markdown_text,
                extensions=['extra', 'codehilite', 'tables']  # Removed 'toc' to avoid issues
            )
            return html
        except Exception as e:
            self.logger.error(f"Markdown conversion failed: {e}")
            # Fallback: wrap in paragraphs
            paragraphs = markdown_text.split('\n\n')
            html_paragraphs = [f'<p>{p.replace(chr(10), "<br>")}</p>' for p in paragraphs if p.strip()]
            return '\n'.join(html_paragraphs)

    def _create_document(self, template_path: Optional[str] = None) -> Document:
        """Create or load DOCX document."""
        if template_path:
            template_file = self._get_template_path(template_path)
            if template_file and template_file.exists():
                self.logger.info(f"Loading DOCX template: {template_file}")
                return Document(str(template_file))

        # Create new document with basic styling
        doc = Document()
        self._setup_document_styles(doc)
        return doc

    def _setup_document_styles(self, doc: Document) -> None:
        """Set up basic document styles."""
        try:
            styles = doc.styles

            # Configure Normal style
            if 'Normal' in styles:
                normal = styles['Normal']
                normal.font.name = 'Calibri'
                normal.font.size = Pt(11)

            # Configure heading styles
            for i in range(1, 7):
                heading_name = f'Heading {i}'
                if heading_name in styles:
                    heading = styles[heading_name]
                    heading.font.name = 'Calibri'
                    heading.font.size = Pt(18 - i * 2)

            self.logger.debug("Document styles configured successfully")

        except Exception as e:
            self.logger.error(f"Style setup failed: {e}")

    def _apply_page_margins(self, doc: Document, margins: Dict[str, float]) -> None:
        """Apply custom page margins to the document."""
        try:
            section = doc.sections[0]

            if 'top' in margins:
                section.top_margin = Inches(margins['top'])
            if 'bottom' in margins:
                section.bottom_margin = Inches(margins['bottom'])
            if 'left' in margins:
                section.left_margin = Inches(margins['left'])
            if 'right' in margins:
                section.right_margin = Inches(margins['right'])

            self.logger.debug(f"Applied page margins: {margins}")

        except Exception as e:
            self.logger.error(f"Failed to apply page margins: {e}")

    def _html_to_docx(self, html_content: str, doc: Document) -> None:
        """Convert HTML content to DOCX document."""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')

            # Process each element in the HTML
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'ul', 'ol', 'table', 'br']):
                self._process_html_element(element, doc)

        except Exception as e:
            self.logger.error(f"HTML to DOCX conversion failed: {e}")
            # Fallback: add as plain text
            doc.add_paragraph(html_content)

    def _process_html_element(self, element, doc: Document) -> None:
        """Process individual HTML elements."""
        tag_name = element.name.lower()

        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            heading_text = self._get_text_content(element)
            if heading_text.strip():
                doc.add_heading(heading_text, level=level)

        elif tag_name in ['p', 'div']:
            text = self._get_text_content(element)
            if text.strip():
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, element)

        elif tag_name == 'table':
            self._process_table(element, doc)

        elif tag_name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                text = self._get_text_content(li)
                if text.strip():
                    list_style = 'List Bullet' if tag_name == 'ul' else 'List Number'
                    doc.add_paragraph(text, style=list_style)

        elif tag_name == 'br':
            doc.add_paragraph()

    def _get_text_content(self, element) -> str:
        """Extract text content from HTML element."""
        if isinstance(element, NavigableString):
            return str(element)

        text_parts = []
        for content in element.contents:
            if isinstance(content, NavigableString):
                text_parts.append(str(content))
            else:
                text_parts.append(self._get_text_content(content))

        return ''.join(text_parts).strip()

    def _process_table(self, table_element, doc: Document) -> None:
        """Process HTML table and convert to DOCX table."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Create table with appropriate dimensions
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=0, cols=max_cols)
        table.style = 'Table Grid'

        for row in rows:
            cells = row.find_all(['td', 'th'])
            table_row = table.add_row()

            for i, cell in enumerate(cells):
                if i < len(table_row.cells):
                    cell_text = self._get_text_content(cell)
                    table_row.cells[i].text = cell_text

                    # Make header cells bold
                    if cell.name == 'th':
                        for paragraph in table_row.cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_formatted_text(self, paragraph, element) -> None:
        """Add formatted text to paragraph maintaining basic formatting."""
        if isinstance(element, NavigableString):
            paragraph.add_run(str(element))
            return

        for content in element.contents:
            if isinstance(content, NavigableString):
                run = paragraph.add_run(str(content))
            else:
                text_content = self._get_text_content(content)
                run = paragraph.add_run(text_content)

                # Apply basic formatting based on HTML tags
                if hasattr(content, 'name'):
                    if content.name in ['strong', 'b']:
                        run.bold = True
                    elif content.name in ['em', 'i']:
                        run.italic = True
                    elif content.name == 'code':
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)

    async def _generate_document_content(self, content: str, **kwargs) -> bytes:
        """
        Generate Word document content from input.

        Args:
            content: Input content (text, markdown, or HTML)
            **kwargs: Additional arguments from MSWordArgs

        Returns:
            DOCX document as bytes
        """
        try:
            # Extract arguments
            template_name = kwargs.get('template_name')
            template_vars = kwargs.get('template_vars')
            docx_template = kwargs.get('docx_template')
            style_config = kwargs.get('style_config')
            page_margins = kwargs.get('page_margins')

            # Process content through HTML template if provided
            processed_content = self._render_html_template(content, template_name, template_vars)

            # Detect content type
            content_type = self._detect_content_type(processed_content)
            self.logger.info(f"Detected content type: {content_type}")

            # Create DOCX document
            doc = self._create_document(docx_template)

            # Apply page margins if specified
            if page_margins:
                self._apply_page_margins(doc, page_margins)

            # Convert content to DOCX based on type
            if content_type == 'html':
                self._html_to_docx(processed_content, doc)
            else:  # markdown or plain text
                # Preprocess and convert markdown to HTML
                cleaned_content = self._preprocess_markdown(processed_content)
                html_content = self._markdown_to_html(cleaned_content)
                self._html_to_docx(html_content, doc)

            # Save document to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            return doc_bytes.getvalue()

        except Exception as e:
            self.logger.error(f"Error generating Word document: {e}")
            raise

    async def _execute(
        self,
        content: str,
        output_filename: Optional[str] = None,
        file_prefix: str = "document",
        output_dir: Optional[str] = None,
        overwrite_existing: bool = False,
        template_name: Optional[str] = None,
        template_vars: Optional[Dict[str, Any]] = None,
        docx_template: Optional[str] = None,
        style_config: Optional[Dict[str, Any]] = None,
        page_margins: Optional[Dict[str, float]] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Execute Word document generation (AbstractTool interface).

        Args:
            content: Content to convert to Word document
            output_filename: Custom filename (without extension)
            file_prefix: Prefix for auto-generated filenames
            output_dir: Custom output directory
            overwrite_existing: Whether to overwrite existing files
            template_name: HTML template name for content processing
            template_vars: Variables for HTML template
            docx_template: DOCX template file path
            style_config: Custom styling configuration
            page_margins: Page margins configuration
            **kwargs: Additional arguments

        Returns:
            Dictionary with document generation results
        """
        try:
            self.logger.info(
                f"Starting Word document generation with {len(content)} characters of content"
            )

            # Use the safe document creation workflow
            result = await self._create_document_safely(
                content=content,
                output_filename=output_filename,
                file_prefix=file_prefix,
                output_dir=output_dir,
                overwrite_existing=overwrite_existing or self.overwrite_existing,
                extension="docx",
                template_name=template_name,
                template_vars=template_vars,
                docx_template=docx_template,
                style_config=style_config,
                page_margins=page_margins
            )

            if result['status'] == 'success':
                self.logger.info(
                    f"Word document created successfully: {result['metadata']['filename']}"
                )

            return result

        except Exception as e:
            self.logger.error(f"Error in Word document generation: {e}")
            raise


class WordToMarkdownTool(AbstractDocumentTool):
    """
    Tool for converting Word documents to Markdown format.

    This tool downloads Word documents from URLs and converts them to Markdown
    format for easier processing by LLMs and other text analysis tools.
    """

    name = "word_to_markdown"
    description = (
        "Convert Word documents to Markdown format from URLs. "
        "Downloads Word documents and converts them to clean Markdown text. "
        "Useful for processing and analyzing Word documents."
    )

    # Document type configuration
    document_type = "conversion"
    default_extension = "md"
    supported_extensions = [".md", ".txt"]

    def __init__(self, **kwargs):
        """Initialize the Word to Markdown tool."""
        super().__init__(**kwargs)
        self._temp_dir = None

    async def _download_file(self, url: str) -> str:
        """Download Word document from URL to temporary file."""
        # Create temporary directory if needed
        if not self._temp_dir:
            self._temp_dir = tempfile.mkdtemp()

        # Generate filename from URL
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        if not filename.endswith(('.docx', '.doc')):
            filename += '.docx'

        file_path = os.path.join(self._temp_dir, filename)

        # Download file
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(f"Download failed with status {response.status}")

                async with aiofiles.open(file_path, 'wb') as f:
                    await f.write(await response.read())

        self.logger.info(f"Downloaded Word document: {filename}")
        return file_path

    async def _convert_to_markdown(self, file_path: str) -> str:
        """Convert Word document to Markdown using mammoth."""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                markdown_text = md(html)

                # Add conversion warnings as comments
                if result.messages:
                    warnings = "\n".join([f"<!-- Warning: {msg} -->" for msg in result.messages])
                    markdown_text = f"{warnings}\n\n{markdown_text}"

                return markdown_text

        except Exception as e:
            self.logger.error(f"Conversion to markdown failed: {e}")
            raise

    async def _cleanup_temp_files(self, file_path: Optional[str] = None) -> None:
        """Clean up temporary files and directory."""
        try:
            if file_path and os.path.exists(file_path):
                os.remove(file_path)

            if self._temp_dir and os.path.exists(self._temp_dir):
                if not os.listdir(self._temp_dir):  # Only remove if empty
                    os.rmdir(self._temp_dir)
                    self._temp_dir = None

        except Exception as e:
            self.logger.warning(f"Cleanup failed: {e}")

    async def convert_from_url(self, url: str, save_markdown: bool = False, **kwargs) -> Dict[str, Any]:
        """
        Convert Word document from URL to Markdown.

        Args:
            url: URL of the Word document
            save_markdown: Whether to save the markdown to a file
            **kwargs: Additional arguments for file saving

        Returns:
            Dictionary with conversion results
        """
        file_path = None
        try:
            # Download the file
            file_path = await self._download_file(url)

            # Convert to markdown
            markdown_content = await self._convert_to_markdown(file_path)

            result = {
                "status": "success",
                "markdown_content": markdown_content,
                "source_url": url,
                "content_length": len(markdown_content),
                "message": "Word document converted to Markdown successfully"
            }

            # Optionally save markdown to file
            if save_markdown:
                file_result = await self._create_document_safely(
                    content=markdown_content,
                    extension="md",
                    **kwargs
                )
                if file_result['status'] == 'success':
                    result.update({
                        "saved_file": file_result['metadata'],
                        "file_path": file_result['metadata']['file_path'],
                        "file_url": file_result['metadata']['file_url']
                    })

            return result

        except Exception as e:
            self.logger.error(f"Word to Markdown conversion failed: {e}")
            return {
                "status": "error",
                "error": str(e),
                "source_url": url,
                "message": f"Failed to convert Word document: {str(e)}"
            }

        finally:
            # Clean up temporary files
            await self._cleanup_temp_files(file_path)

    async def _generate_document_content(self, content: str, **kwargs) -> str:
        """Generate markdown content (implementation required by AbstractDocumentTool)."""
        # This tool is primarily for URL conversion, but we implement this for completeness
        return content

    async def _execute(self, url: str, save_markdown: bool = False, **kwargs) -> Dict[str, Any]:
        """Execute Word to Markdown conversion."""
        return await self.convert_from_url(url, save_markdown, **kwargs)
